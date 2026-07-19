import os
import json
import io
import re
import datetime
import base64
import numpy as np
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from google import genai
from google.genai import types
from models import AnalysisResponse
from dotenv import load_dotenv
from openai import OpenAI

# Load the .env that sits next to this module, so the app works regardless of the
# current working directory (e.g. when uvicorn is launched from the repo root with
# --app-dir). Falls back to default discovery; real OS env vars (e.g. on Render)
# still take precedence since load_dotenv does not override existing vars.
load_dotenv(os.path.join(os.path.dirname(__file__), ".env"))
load_dotenv()

# --- Gemini client (used by Letter Builder + ITA Strategy) ---
gemini_client = genai.Client()

# --- OpenAI client (used by NOC Finder + Auditor + OCR) ---
_openai_api_key = os.getenv("OPENAI_API_KEY", "")
openai_client = OpenAI(api_key=_openai_api_key) if _openai_api_key else None

# The Employment Letter Auditor grades every letter on a strong, deterministic model. gpt-4o-mini was
# retired from this path: it under-grades duty coverage and is non-deterministic (57%/57%/57%/71% on a
# letter whose accurate coverage is 85%), which showed users different numbers on the Finder vs the
# Auditor. Haiku 4.5 is accurate and stable (85%/85%); paid tiers escalate to Sonnet 4.6.
AUDIT_STANDARD_MODEL = "claude-haiku-4-5-20251001"

# Duty coverage: a strong OR partial match counts as a demonstrated duty (1), weak/missing count
# for nothing. Coverage % = demonstrated / total ESSENTIAL official duties. Non-essential duties
# (the ones the NOC lists as "May …" — occasional/optional) are excluded from the denominator so a
# letter isn't penalised for not showing duties that aren't core to the occupation.
def _is_essential_duty(duty_text: str) -> bool:
    return not str(duty_text or "").strip().lower().startswith("may ")

def _backfill_official_duties(result: dict, code: str) -> None:
    """Guarantee the duty-by-duty table covers EVERY official main duty of the NOC. Any duty the model
    omitted is graded semantically against the applicant's demonstrated evidence (so a genuinely-shown
    duty the model forgot is still credited) and appended; if it can't be graded, it's added as
    'missing' so the table is at least complete. Sub-occupation scoping happens later."""
    na = result.get("noc_analysis")
    if not isinstance(na, dict) or not code:
        return
    dm = na.get("duties_match") or []
    ent = NOC_CODE_TO_ENTRY.get(code) or {}
    official = [d for d in (ent.get("duties_flat") or ent.get("duties") or []) if d and d.strip()]
    if not official:
        return
    present = {(m.get("noc_duty", "") or "").strip().lower() for m in dm}
    missing = [d for d in official if d.strip().lower() not in present]
    if not missing:
        return
    # Grade the omitted duties against what the letter already evidenced (the model's matched quotes).
    profile = [(m.get("letter_evidence") or "").strip() for m in dm
               if m.get("match_strength") in ("strong", "partial") and (m.get("letter_evidence") or "").strip()]
    graded = {}
    if profile and openai_client:
        try:
            rp = openai_client.embeddings.create(model="text-embedding-3-small", input=profile)
            rd = openai_client.embeddings.create(model="text-embedding-3-small", input=missing)
            P = np.array([x.embedding for x in rp.data], dtype=np.float32)
            P /= (np.linalg.norm(P, axis=1, keepdims=True) + 1e-9)
            D = np.array([x.embedding for x in rd.data], dtype=np.float32)
            D /= (np.linalg.norm(D, axis=1, keepdims=True) + 1e-9)
            sims = (D @ P.T).max(axis=1)
            for d, s in zip(missing, sims):
                graded[d] = "strong" if s >= STRONG_DUTY_SIM else ("partial" if s >= 0.30 else "missing")
        except Exception as e:
            print(f"[Auditor] duty backfill grading failed: {e}")
    for d in missing:
        dm.append({"noc_duty": d, "letter_evidence": "", "match_strength": graded.get(d, "missing")})
    na["duties_match"] = dm
    print(f"[Auditor] Backfilled {len(missing)} official dut(ies) the model omitted "
          f"({sum(1 for d in missing if graded.get(d) in ('strong','partial'))} evidenced).")


def coverage_pct(pairs) -> int:
    """pairs = iterable of (duty_text, match_strength). Binary coverage over essential duties."""
    pairs = list(pairs)
    essential = [(d, s) for (d, s) in pairs if _is_essential_duty(d)]
    if not essential:                      # every duty was "May …" — fall back to all of them
        essential = pairs
    total = len(essential)
    if not total:
        return 0
    covered = sum(1 for (_d, s) in essential if str(s or "").lower() in ("strong", "partial"))
    return int(round(100 * covered / total))

# Load the NOC index once at startup
_noc_index_path = os.path.join(os.path.dirname(__file__), "noc_index.json")
with open(_noc_index_path, "r", encoding="utf-8") as f:
    NOC_INDEX = json.load(f)
print(f"Loaded NOC index: {len(NOC_INDEX)} unit groups")

# Pre-build fast lookup: NOC code -> title (used by sanitizer)
NOC_LOOKUP = {}
for _entry in NOC_INDEX.values():
    _code = _entry.get("code", "")
    if _code:
        NOC_LOOKUP[_code] = _entry.get("title", "")
print(f"Built NOC_LOOKUP: {len(NOC_LOOKUP)} codes")

# Pre-build reverse lookup: title (lowercase) -> code (used by sanitizer for hallucination fix)
NOC_TITLE_TO_CODE = {}
for _code, _title in NOC_LOOKUP.items():
    NOC_TITLE_TO_CODE[_title.lower().strip()] = _code
print(f"Built NOC_TITLE_TO_CODE: {len(NOC_TITLE_TO_CODE)} titles")

# Pre-build fast lookup: NOC code -> full index entry. NOC_INDEX is keyed by an
# internal index key, not the code, so without this map every code->entry lookup
# was an O(n) linear scan over all 516 entries (done several times per request).
NOC_CODE_TO_ENTRY = {}
for _entry in NOC_INDEX.values():
    _code = _entry.get("code", "")
    if _code:
        NOC_CODE_TO_ENTRY[_code] = _entry
print(f"Built NOC_CODE_TO_ENTRY: {len(NOC_CODE_TO_ENTRY)} entries")


def get_noc_entry(code: str) -> dict | None:
    """Return the full NOC index entry for a 5-digit code, or None. O(1)."""
    return NOC_CODE_TO_ENTRY.get(code)


# Matches employer-stated NOC references like "NOC 42201", "(NOC: 42201)",
# "NOC Code 42201", or "NOC 42201 - Social and community service workers".
_EMPLOYER_NOC_RE = re.compile(
    r'(?i)\(?\s*NOC\s*[:#]?\s*\d{4,5}\s*[-–—]?\s*[^\n)]{0,80}\)?'
)


def strip_employer_noc_references(text: str) -> tuple[str, bool]:
    """Redact employer-stated NOC codes from input text.

    Employers frequently state an (often incorrect) NOC code in employment
    letters; leaving it in the prompt anchors the model to that code. Returns
    (cleaned_text, was_modified). Shared by the NOC Finder, the multi-agent
    pipeline, and the Auditor so the redaction rule lives in one place.
    """
    cleaned = _EMPLOYER_NOC_RE.sub('[EMPLOYER NOC REFERENCE REDACTED]', text)
    return cleaned, cleaned != text

# Load NOC embeddings for RAG and pre-compute numpy matrix
NOC_EMBEDDINGS = {}
_NOC_EMB_MATRIX = None  # Pre-computed numpy matrix for fast similarity
_NOC_EMB_KEYS = []      # Ordered list of index keys matching matrix rows
_embeddings_path = os.path.join(os.path.dirname(__file__), "noc_embeddings.json")
if os.path.exists(_embeddings_path):
    with open(_embeddings_path, "r", encoding="utf-8") as f:
        NOC_EMBEDDINGS = json.load(f)
    # Pre-compute numpy matrix: one np.array call instead of 516 per request
    _NOC_EMB_KEYS = list(NOC_EMBEDDINGS.keys())
    _NOC_EMB_MATRIX = np.array([NOC_EMBEDDINGS[k] for k in _NOC_EMB_KEYS])
    print(f"Loaded NOC embeddings: {len(NOC_EMBEDDINGS)} vectors, matrix shape: {_NOC_EMB_MATRIX.shape}")
else:
    print("WARNING: noc_embeddings.json not found. RAG NOC Finder will fail.")

# Load per-duty embeddings for duty-level reranking (zero-cost at query time)
_DUTY_EMB_MATRIX = None   # (N_duties, 1536) pre-computed duty vectors
_DUTY_RANGES = {}         # NOC code -> (start_row, end_row) in the matrix
_duty_emb_path = os.path.join(os.path.dirname(__file__), "noc_duty_embeddings.npz")
_duty_idx_path = os.path.join(os.path.dirname(__file__), "noc_duty_index.json")
if os.path.exists(_duty_emb_path) and os.path.exists(_duty_idx_path):
    _DUTY_EMB_MATRIX = np.load(_duty_emb_path)["embeddings"]
    with open(_duty_idx_path, "r", encoding="utf-8") as f:
        _duty_idx = json.load(f)
    _DUTY_RANGES = {code: tuple(rng) for code, rng in _duty_idx["ranges"].items()}
    print(f"Loaded per-duty embeddings: {_DUTY_EMB_MATRIX.shape[0]} duties across {len(_DUTY_RANGES)} NOCs")
else:
    print("WARNING: Per-duty embeddings not found. Duty-level reranking disabled.")

# Load pre-computed lead statement embeddings for industry/employer-type matching
_LEAD_EMB_MATRIX = None   # (516, 1536) lead statement vectors
_LEAD_EMB_CODES = []      # NOC codes in matrix order
_lead_emb_path = os.path.join(os.path.dirname(__file__), "noc_lead_embeddings.npz")
_lead_idx_path = os.path.join(os.path.dirname(__file__), "noc_lead_index.json")
if os.path.exists(_lead_emb_path) and os.path.exists(_lead_idx_path):
    _LEAD_EMB_MATRIX = np.load(_lead_emb_path)["embeddings"]
    with open(_lead_idx_path, "r", encoding="utf-8") as f:
        _LEAD_EMB_CODES = json.load(f)["codes"]
    # Build code → row index mapping for fast lookup
    _LEAD_CODE_TO_IDX = {code: idx for idx, code in enumerate(_LEAD_EMB_CODES)}
    print(f"Loaded lead statement embeddings: {_LEAD_EMB_MATRIX.shape[0]} NOCs")
else:
    _LEAD_CODE_TO_IDX = {}
    print("WARNING: Lead statement embeddings not found. Lead-weighted reranking disabled.")

# MIME type mapping for images
IMAGE_MIME_TYPES = {
    '.jpg': 'image/jpeg',
    '.jpeg': 'image/jpeg',
    '.png': 'image/png',
    '.bmp': 'image/bmp',
    '.tiff': 'image/tiff',
    '.tif': 'image/tiff',
    '.webp': 'image/webp',
}

def extract_text_from_pdf(pdf_bytes: bytes) -> str:
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    text = ""
    for page in doc:
        text += page.get_text() + "\n"
    return text

def pdf_pages_to_images(pdf_bytes: bytes, max_pages: int = 5) -> list[tuple[bytes, str]]:
    """Convert each page of a PDF to a PNG image. Returns list of (image_bytes, mime_type). Limits to max_pages."""
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    images = []
    
    # Only process up to max_pages to prevent API rate limits / huge costs
    pages_to_process = min(len(doc), max_pages)
    
    for i in range(pages_to_process):
        page = doc[i]
        # Render at 2x resolution for better OCR quality
        pix = page.get_pixmap(dpi=200)
        img_bytes = pix.tobytes("png")
        images.append((img_bytes, "image/png"))
        
    print(f"Converted {len(images)} PDF page(s) out of {len(doc)} to images for vision processing")
    return images

def extract_text_from_docx(docx_bytes: bytes) -> str:
    doc = DocxDocument(io.BytesIO(docx_bytes))
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                paragraphs.append(" | ".join(row_text))
    return "\n".join(paragraphs)

def _build_prompt_text(noc_reference: str, target_noc: str = None) -> str:
    """Builds the system instruction for the skeptical immigration officer auditor prompt."""
    
    # NOTE: The NOC code is determined UPSTREAM by the NOC Finder (auto_detect_noc) and force-locked
    # in code after this prompt runs, so the Auditor does NOT re-derive the NOC from scratch — it
    # EVALUATES the applicant's duty coverage against the supplied NOC. `target_noc` is therefore set
    # on essentially every call; the no-target branch below is only a defensive fallback.
    multi_title_rule = """
MULTI-TITLE NOC GROUPS: Some codes cover several DISTINCT occupations (e.g. "Translators, terminologists and interpreters" — an interpreter is none of the other two). When the NOC is such a group, evaluate the applicant ONLY against the duties of the sub-occupation they actually perform. The other sub-occupations' duties are NOT APPLICABLE — exclude them entirely (do not list them in duties_match and do not count them in coverage). A perfect interpreter must not be penalized for not translating documents."""

    if target_noc:
        task_noc_matching = f"""
=== TASK 1 - NOC EVALUATION (against the pre-determined code) ===
The primary NOC for this document has already been determined: NOC {target_noc}. Your job is to
EVALUATE the applicant's duty alignment against it, not to search for a different code.
- Set `noc_analysis.detected_code` strictly to "{target_noc}".
- Fetch and set `noc_analysis.detected_title` to the exact title for {target_noc} from the database.
- In `noc_analysis.match_score`, give the objective duty-coverage percentage (0-100). Be honest — it is fine if it is low.
- In `noc_analysis.confidence`, rate how confident you are that {target_noc} is the right occupation (0-100).
- Map their duties strictly against the official duties of NOC {target_noc}; explain overlaps and note glaring gaps.
- You may list genuinely better fits in `alternative_nocs` (codes from the provided database ONLY).
{multi_title_rule}
"""
    else:
        # Defensive fallback only — auto-detection upstream failed. Pick by duty coverage from the
        # provided candidates; keep it lean since this path is rarely reached.
        task_noc_matching = f"""
=== TASK 1 - NOC SELECTION (fallback — pick best fit from the provided candidates) ===
Compare the letter's duties against the official "duties" of each candidate NOC and select the code
with the highest genuine duty coverage. Match on the SEMANTIC FUNCTION the person performs, not on
surface keywords or the job title (a frequent industry/product/setting word is not the function).
- `noc_analysis.detected_code`: the 5-digit code (from the provided database ONLY — never invent codes).
- `noc_analysis.detected_title`: the exact title from the database.
- `noc_analysis.match_score`: objective duty coverage out of 100.
- `noc_analysis.confidence`: how confident you are in the selection (0-100).
- `noc_analysis.alternative_nocs`: secondary matches >= 50% (database codes only).
{multi_title_rule}
"""

    return f"""You are an advanced AI system acting as a STRICT, SKEPTICAL, and FAIR Canadian Immigration Officer auditing employment letters for Express Entry (which includes the Canadian Experience Class, the Federal Skilled Worker Program, the Federal Skilled Trades Program, and Provincial Nominee streams).

SCOPE — READ CAREFULLY: Your ONLY job is to assess whether THIS LETTER credibly documents the applicant's duties and meets IRCC's employment-letter requirements for the target NOC. You are NOT assessing program eligibility. Foreign (non-Canadian) work experience is fully valid for the Federal Skilled Worker Program and counts toward Express Entry. NEVER treat work performed outside Canada as a risk, a gap, or a disqualifier, and NEVER state that the experience "does not qualify", "will not count", or "cannot support the application" because of where it took place. The applicant's location of experience is neutral metadata only.

**IMPORTANT: Today's date is {datetime.date.today().strftime('%B %d, %Y')}. Use this as the current date for any date-related analysis. Do NOT hallucinate or guess a different date.**

PRIMARY OBJECTIVE:
Determine whether the applicant's work experience, as described in the employment letter, would likely be:
1) ACCEPTED
2) FLAGGED FOR PROCEDURAL FAIRNESS LETTER (PFL)
3) REFUSED

You must think like an IRCC officer whose role is to VERIFY, NOT TRUST.
Do NOT assume the applicant qualifies. The burden of proof is on the applicant.

You have been given:
1. A PRE-FILTERED subset of the NOC 2021 database containing the most relevant unit groups for this document. Each entry has an official code, title, lead statement, and main duties.
2. A document uploaded by the user - either as extracted text or as an image.

**CRITICAL RULE: You may ONLY reference NOC codes and titles that appear in the provided database below. Do NOT invent, guess, or recall NOC codes from memory. If a NOC code is not in the database provided, you MUST NOT use it anywhere in your response — not as the primary match, not as an alternative, and not in any narrative text.**

---

=== CORE PRINCIPLES ===

1. BURDEN OF PROOF - Only consider what is explicitly stated or strongly implied in the letter. Do NOT assume missing duties were performed.
2. SKEPTICAL REVIEW - Actively look for weaknesses, ambiguity, and gaps that could justify refusal.
3. DUTY-BASED MATCHING - Match duties semantically, not by keywords. The applicant must demonstrate the MAJORITY (>=70%) of the MAIN DUTIES of a single NOC.
4. PARTIAL MATCH = RISK - If duties are vague, generic, or incomplete, increase refusal risk.

---

=== VALIDATION PIPELINE (PRE-CHECKS - Do this FIRST) ===

Classify each check as HARD_FAIL (stop analysis) or SOFT_FAIL (continue with warnings) or PASS.

CHECK 1 - READABLE CONTENT (HARD_FAIL if blank/corrupted):
- If the document is blank, empty, corrupted, or contains no readable text whatsoever, REJECT.
- Set `document_type` to "Blank / Unreadable Document".

CHECK 2 - DOCUMENT TYPE (HARD_FAIL if wrong type):
Only ACCEPTABLE:
  Employment / Reference / Experience letter issued BY an employer
  Job offer letter that includes duties/responsibilities
  A single letter covering multiple roles at the SAME company

NOT ACCEPTABLE (REJECT with clear explanation):
  Payslips, T4/T4A slips, ROE, tax/payroll documents
  Resumes or CVs (self-authored)
  Cover letters (applicant-authored)
  Job postings / advertisements
  LinkedIn profile screenshots
  Contracts without duties section
  Bank statements, invoices, receipts, ID documents
  Business cards, certificates, diplomas
  Unrelated photos or images

CHECK 3 - MULTIPLE EMPLOYERS (HARD_FAIL):
- If letters from TWO OR MORE DIFFERENT employers/companies are merged into one file, REJECT.
- A single letter covering multiple ROLES at the SAME company is perfectly valid.

CHECK 4 - LANGUAGE (SOFT_FAIL):
- If NOT in English or French, add a medium-severity risk. Still attempt analysis if duties are discernible.

CHECK 5 - DUTIES QUALITY (SOFT_FAIL):
- If fewer than 2 duties, flag as high-severity risk.
- If duties appear copy-pasted VERBATIM from the NOC website, flag as high-severity risk.

CHECK 6 - HEAVY REDACTION (SOFT_FAIL):
- If the APPLICANT redacted significant portions (e.g. blacked-out salary or duties), add a warning noting which elements could not be verified.
- CRITICAL EXCEPTION: our system automatically removes/masks any NOC code numbers the employer wrote in the letter (to stop those numbers biasing your NOC selection). This removal is intentional and expected. NEVER flag a missing, removed, blanked, or "redacted" NOC code/reference as redaction, tampering, or an authenticity concern — it is our doing, not the applicant's.

=== REJECTION OUTPUT FORMAT ===
If ANY of Checks 1-3 result in HARD_FAIL:
- Set `decision` to "REFUSE", `confidence_score` to 95+.
- Set `noc_analysis.applicable` to false, `noc_analysis.detected_code` to "", `noc_analysis.match_score` to 0.
- Write a clear `officer_narrative` explaining what was detected and what to upload instead.
- Populate `refusal_reasons` with specific grounds.
- STOP. Do NOT attempt NOC matching or compliance auditing.

If all checks pass (or only SOFT_FAIL), proceed with the full analysis.

---

{task_noc_matching}

---

=== TASK 2 - DUTY EVIDENCE MAPPING (CRITICAL - This drives the decision) ===

You MUST include an entry in `duties_match` for EVERY official main duty of the selected NOC listed
in the database below — never omit a duty. Include it even if it seems unrelated or belongs to a
different sub-occupation packed under the same code. Completeness is mandatory: if the target NOC
lists 13 main duties, `duties_match` must contain 13 entries. (The platform handles sub-occupation
scoping afterward — that is not your job.)
For each duty, set `match_strength`:
  - "strong" - clear semantic alignment, specific evidence quoted from the letter
  - "partial" - related language but vague or incomplete
  - "weak" - only tangentially related
  - "missing" - no evidence in the letter at all

For `missing_critical_duties`, list every NOC duty that received "missing" or "weak" match_strength.
Do NOT compute or state a coverage percentage anywhere — the platform computes the exact figure.

For `lead_statement_*`:
  - Quote the official lead statement from the NOC database
  - Quote the most relevant evidence from the letter
  - Explain how they align (or don't)

---

=== TASK 3 - IRCC COMPLIANCE AUDIT ===

Evaluate the employment letter against the official IRCC requirements for CEC reference letters.

MANDATORY ELEMENTS (populate `mandatory_requirements` booleans - set TRUE only if verifiably present):
1. OFFICIAL COMPANY LETTERHEAD
2. APPLICANT'S FULL NAME
3. COMPANY CONTACT INFORMATION - Address, telephone, email
4. JOB TITLE(S)
5. DATES OF EMPLOYMENT - Specific start and end dates
6. HOURS WORKED PER WEEK - Must prove full-time (30+) or state part-time hours
7. SALARY / COMPENSATION - Any format acceptable (hourly, weekly, monthly, annual)
8. SIGNATORY - Name, title, and signature of supervisor OR HR officer (both valid)

HOW TO JUDGE ELEMENTS 1, 3 AND 8 (letterhead / contact information / signatory):
- Use BOTH the page images (when provided) and the extracted text. A header or footer block with
  the company name, logo, address, phone, email or website IS official letterhead.
- Word documents and plain-text extractions often reach you WITHOUT page images, and PDF text
  extraction frequently loses header/footer formatting. Judge from textual cues (company block,
  signature block with name and title). NEVER mark these elements false merely because no image
  was provided or because formatting was lost in extraction.
- Contact information counts wherever it appears: header, footer, signature block, or body.
- Mark FALSE only when the element is affirmatively absent from BOTH the text and the images.

For `compliance.score`: (count of true mandatory_requirements / 8) x 100
For `compliance.missing_elements`: List mandatory elements completely absent.
For `compliance.warnings`: List elements present but insufficient/ambiguous.

=== DO NOT FLAG (IMPORTANT - These are NOT issues) ===
- Work experience gained OUTSIDE Canada — foreign experience is valid for Express Entry (FSW etc.). Location is NEVER a risk or disqualifier.
- A removed / masked / "redacted" NOC code or reference — our system did that on purpose to prevent bias. Not an authenticity concern.
- Letter dated AFTER employment end date (normal for reference letters)
- Salary in any format (hourly, monthly, biweekly - all acceptable)
- HR signatory instead of supervisor (both valid per IRCC)
- Ongoing employment as "currently employed", "to date", "present" (all acceptable)
- Job title not matching NOC title exactly (duties override titles)
- Minor wording differences from official NOC descriptions
- Minor formatting inconsistencies

---

=== TASK 4 - RISK ASSESSMENT ===

`overall_risk`: low (strong case) / moderate (gaps but defensible) / high (likely refusal)
`pfl_likelihood`: low / medium / high - how likely is IRCC to issue a Procedural Fairness Letter?
`key_risks`: Every identified risk, ordered by severity (highest first). Each must be specific and actionable.

---

=== TASK 5 - DECISION (You MUST choose ONE. No neutrality allowed.) ===

ACCEPT if:
- >=75% duty coverage
- Clear, specific, verifiable duties
- All critical compliance elements present

PFL_RISK if:
- 50-75% duty coverage OR
- Ambiguity / vague duties OR
- Missing supporting clarity
- You MUST populate `refusal_reasons` with the specific grounds that would trigger a PFL.

REFUSE if:
- <50% duty coverage OR
- Missing key duties entirely OR
- Insufficient evidence to establish NOC alignment
- You MUST populate `refusal_reasons`.

---

=== TASK 6 - OFFICER NARRATIVE ===

Write `officer_narrative` in a realistic IRCC officer tone:
- Formal, concise, evidence-based, slightly skeptical
- 3-5 sentences referencing specific duties and gaps
- Example: "The duties described are insufficiently detailed to establish alignment with the claimed NOC. While the applicant references supervisory responsibilities, the letter lacks specificity regarding..."
- NEVER state, estimate, or cite a duty-coverage PERCENTAGE or fraction (e.g. "60%", "4 of 7 duties") anywhere in the narrative, action plan, risks, or refusal reasons. The platform computes and displays the exact coverage figure separately, and any number you invent would contradict it. Describe coverage only in words — "most core duties are clearly demonstrated", "several key duties are missing", etc.

---

=== TASK 7 - ACTION PLAN & SUGGESTED WORDING ===

`action_plan`: Priority-ordered, specific, actionable fixes. Most critical first. Each item must be directly tied to an identified risk or gap. Not generic advice.
`suggested_wording`: If duties are weak or missing, provide sample sentences the applicant can give to their employer to strengthen the letter's alignment with the NOC.

---

=== TASK 8 - LOCATION OF EXPERIENCE (neutral metadata — NOT a risk) ===

Classify the company address / geographic references for our records ONLY. This never affects the
decision, risk, or narrative, and outside-Canada experience is NOT a problem (see SCOPE above):
- Canadian address/postal code/province -> "canada"
- Non-Canadian address -> "outside_canada"
- Cannot determine -> "unknown"

---

=== NOC 2021 DATABASE (Pre-filtered subset — ONLY use codes listed here) ===
{noc_reference}

---

=== FINAL RULE ===
If evidence is weak or missing, you MUST penalize heavily.
You MUST NOT "help" the applicant pass.
Your role is to PROTECT the integrity of the immigration system while giving fair, actionable feedback.

Output your analysis strictly conforming to the requested JSON schema. Be precise and evidence-based.
"""

def build_noc_finder_prompt(noc_reference: str, target_noc: str = None, function_classification: dict = None) -> str:
    """Builds the NOC Finder prompt v2 — fast, reliable, IRCC-consistent NOC suggestion."""

    today = datetime.date.today().strftime('%B %d, %Y')

    # Build function classification injection block
    # When the Function Classifier agent has run, its output is injected as a hard
    # constraint into the NOC selection prompt. This prevents domain bias (picking a
    # NOC that matches the industry rather than what the person does) by
    # pre-determining the applicant's functional role.
    func_block = ""
    if function_classification and not target_noc:
        fc = function_classification
        pf = fc.get('primary_function', 'OTHER')
        conf = fc.get('confidence', 'low')
        verbs = ', '.join(fc.get('key_verbs', []))
        reasoning = fc.get('reasoning', 'N/A')
        secondary = fc.get('secondary_function', '') or 'None'
        func_block = f"""
=== FUNCTION CLASSIFICATION (Pre-determined by independent agent — DO NOT override) ===

An independent analysis of the applicant's duties — performed WITHOUT seeing any
NOC candidates — classified their primary work function as:

  PRIMARY FUNCTION: {pf}
  CONFIDENCE: {conf}
  KEY VERBS: {verbs}
  REASONING: {reasoning}
  SECONDARY FUNCTION: {secondary}

You MUST respect this classification when selecting a NOC:
- Your primary NOC MUST align with the {pf} functional category
- If your top-ranked candidates by _duty_match_rank belong to a DIFFERENT functional
  category (e.g., TRADES_PRODUCTION when the classification says INSPECTION_QC), you MUST
  reject them and select from candidates matching the classified function
- The ONLY exception: if the classification confidence is "low" AND you have strong
  evidence from the duties that a different function is more appropriate, you may
  override — but you MUST explain why in why_this_noc
"""
    elif not target_noc:
        # Fallback: self-classification when Function Classifier hasn't run
        func_block = """
=== FUNCTION VS. DOMAIN — MANDATORY CLASSIFICATION (DO THIS FIRST) ===

BEFORE selecting any NOC, you MUST classify the applicant's PRIMARY FUNCTION into exactly
one of these four categories based on their duties:

  (A) TRADES/PRODUCTION — They physically BUILD, FABRICATE, WELD, ASSEMBLE, or INSTALL products
  (B) INSPECTION/QC — They INSPECT, AUDIT, TEST, REVIEW CERTIFICATES, or ensure COMPLIANCE
  (C) SUPERVISION — They MANAGE SCHEDULES, HIRE STAFF, ASSIGN WORK, or COORDINATE workers
  (D) ENGINEERING — They DESIGN SYSTEMS, OPTIMIZE PROCESSES, or DEVELOP PROGRAMS

Key verb signals:
- "Responsible for" a process usually means OVERSEEING it, not personally performing it → (B) or (C).
- "Inspecting / testing / monitoring / verifying / witnessing" work → category (B), even when it
  happens in a trades or production setting. They check the work; they do not build the product.
- "Planning / scheduling / managing staff" → category (C), not (A).

After classifying, your selected NOC MUST belong to the SAME functional category.
If your top-ranked candidate belongs to a different category (e.g., a trades NOC when the applicant's
function is INSPECTION/QC), reject it and select from the candidates that match the function.
"""

    # --- Task 1 block varies for auto vs targeted evaluation ---
    if target_noc:
        task_1 = f"""
=== TASK 1 — TARGETED NOC EVALUATION ===

The user explicitly requested evaluation against NOC {target_noc}.
You MUST lock the primary match to NOC {target_noc}.

- Set `recommended_noc.code` strictly to "{target_noc}".
- Set `recommended_noc.title` to the exact title for {target_noc} from the database.

=== CONFIDENCE CALCULATION (MUST FOLLOW THIS EXACTLY) ===

To compute `recommended_noc.confidence`, you MUST:
1. Look up ALL main duties for NOC {target_noc} from the database.
2. For EACH main duty, classify the applicant's evidence:
   - "strong" = clear semantic alignment with specific evidence from the input
   - "partial" = related language but vague or incomplete
   - "missing" = no evidence in the input at all
3. Set `recommended_noc.duties_total` = total number of main duties for the NOC.
4. Set `recommended_noc.duties_matched` = count of "strong" + "partial".
5. Set `recommended_noc.confidence` = (duties_matched / duties_total) × 100.

These three numbers MUST be mathematically consistent. Do NOT estimate — count the duties.

- Classify `result_type` based on confidence: STRONG_MATCH (≥70%), MODERATE_MATCH (45-69%), NO_MATCH (<45%).
- `key_matches`: List the duties classified as "strong" (up to 5)
- `key_gaps`: List the duties classified as "missing" (up to 3)
- You may still list better fits in `alternatives`.
"""
    else:
        task_1 = f"""
=== TASK 1 — NOC MATCHING ===

=== PRE-COMPUTED DUTY COVERAGE (USE THIS DATA) ===

Each NOC entry in the database below includes machine-computed scores:
- `_duty_match_rank`: Position in duty-level matching (1 = strongest duty overlap)
- `_pre_computed_duty_coverage_pct`: % of this NOC's duties that semantically match the input
- `_pre_computed_duties_matched` / `_pre_computed_duties_total`: raw duty counts
- `_lead_statement_match`: How well the NOC's lead statement (employer type/industry) aligns
  with the user's described work (0.0-1.0). Higher = better industry match.
  CRITICAL: If two NOCs have similar duty scores but different lead_statement_match scores,
  ALWAYS prefer the one with the higher _lead_statement_match. This corrects cases where two
  occupations share duties but belong to different industries.

These scores were computed by comparing EACH individual NOC duty against the user's text 
using embedding similarity. They are OBJECTIVE and should be your PRIMARY signal for 
NOC selection. Do NOT override these scores based on job title alone.

START your evaluation with the top-ranked candidates (rank 1-5) and verify their 
alignment by reading the duties yourself. Only select a lower-ranked candidate if the
top candidates' duties genuinely do not match the user's described work AND their
_lead_statement_match is low.

Steps:

1. Read the user's duties carefully.
2. Start with the NOCs ranked #1-3 by `_duty_match_rank` — these have the highest
   objective duty overlap with the input.
3. For your top candidate, VERIFY the machine scores by reading each duty yourself.
4. Compare duties against candidate NOCs in the database:
   - Lead statement alignment
   - Main duties overlap (SEMANTIC, not keyword matching)
3. For your top candidate NOC, compute DUTY COVERAGE using the method below.
4. Select the BEST NOC based on priority order:
   a) Duty coverage %
   b) Lead statement alignment
   c) Specificity of duties
5. Also evaluate up to 2 ALTERNATIVE NOCs with honest scores.

=== CRITICAL: SEMANTIC MATCHING, NOT KEYWORD MATCHING ===
Focus on the SEMANTIC MEANING of what the person actually does day-to-day.
A word that appears frequently in the letter (an industry, material, or setting) is NOT necessarily
the person's function: a role that inspects for hazards and recommends controls is a safety function
regardless of the setting it occurs in. Always ask: "What is this person's core function?"

{func_block}

=== EMPLOYER INDUSTRY CROSS-CHECK (a tie-breaker, NOT the primary test) ===

Duties decide the NOC. Use the employer's industry only to break a tie when two NOCs fit the
duties equally well:
1. Identify what the employer's business actually is, from the letter's context.
2. Read the lead statement of each tied NOC — it names the employer types where the occupation
   is typically found.
3. When the duty match is genuinely equal, prefer the NOC whose lead statement matches the employer.

Two important cautions:
- A strong DUTY match is NOT overridden by an employer-type mismatch. Do not abandon the
  best duty-aligned NOC just because the employer setting differs from its lead statement.
- OUTSOURCING / STAFFING / MARKETING / PROMOTIONS / CALL-CENTRE / BPO agencies: the agency's OWN
  industry does NOT define the occupation — classify by the work actually performed, often for an
  end client (e.g. someone at a marketing agency selling a bank's products is doing financial-product
  sales; someone placed by a staffing agency onto a factory line is doing that factory job).

=== TIE-BREAKING RULE ===
If two NOC codes score within 5 points, select the NOC whose LEAD STATEMENT most accurately
describes the person's primary role. List the runner-up as the first alternative.

=== CONFIDENCE CALCULATION (MUST FOLLOW THIS EXACTLY) ===

To compute `recommended_noc.confidence`, you MUST:
1. Look up ALL main duties for the selected NOC from the database.
2. For EACH main duty, classify the applicant's evidence:
   - "strong" = clear semantic alignment with specific evidence from the input
   - "partial" = related language but vague or incomplete
   - "missing" = no evidence in the input at all
3. Set `recommended_noc.duties_total` = total number of main duties for the NOC.
4. Set `recommended_noc.duties_matched` = count of "strong" + "partial".
5. Set `recommended_noc.confidence` = (duties_matched / duties_total) × 100.

These three numbers MUST be mathematically consistent. Do NOT estimate — count the duties.

=== MATCH CLASSIFICATION ===
- ≥70% → result_type = "STRONG_MATCH"
- 45–69% → result_type = "MODERATE_MATCH"
- <45% → result_type = "NO_MATCH"

If NO NOC reaches ~45%, still return the best candidate but set result_type to NO_MATCH
and explain in `why_this_noc` that alignment is weak.

=== OUTPUT MAPPING ===
- `recommended_noc.code`: Best-matching 5-digit NOC code
- `recommended_noc.title`: Exact title from the database
- `recommended_noc.confidence`: Computed duty coverage percentage (0-100)
- `recommended_noc.duties_total`: Total main duties for this NOC in the database
- `recommended_noc.duties_matched`: How many duties have strong or partial evidence
- `why_this_noc`: 1-2 sentence explanation of selection + any key concerns
- `key_matches`: Duties classified as "strong" (up to 5, short strings)
- `key_gaps`: Duties classified as "missing" (up to 3, short strings)
- `alternatives`: Up to 2 alternative NOCs with their computed confidence scores.
  **CRITICAL: You may ONLY suggest alternative NOCs whose codes appear in the provided database. NEVER invent or recall NOC codes from memory.**
"""

    return f"""You are a Canadian immigration NOC (National Occupational Classification) expert specializing in NOC 2021.

Your role is to IDENTIFY the most likely NOC for a user's work experience, while ensuring consistency with IRCC evaluation standards.

---

PRIMARY OBJECTIVE:
Recommend the MOST LIKELY NOC based on the user's duties, BUT ONLY if the duties appear to demonstrate meaningful alignment with that NOC.

You are NOT allowed to confidently recommend a NOC if alignment is weak. Be honest.

---

IMPORTANT CONTEXT:
- Today's date: {today}
- The applicant must demonstrate a MAJORITY (~60-70%) of the main duties of a NOC to be considered a strong match.
- Your output must be CONSISTENT with a future audit. Do NOT recommend a NOC that would likely fail a detailed audit.

---

=== INPUT VALIDATION (RUN FIRST) ===

IF DOCUMENT INPUT:

1. READABILITY:
   - If blank, corrupted, or no readable text → Set document_valid=false, rejection_reason="Document is blank or unreadable."

2. DOCUMENT TYPE:
   ACCEPT: Employment/reference/experience letters, Job offer letters WITH duties, Single document with multiple roles at SAME company
   ALLOW WITH LOWER CONFIDENCE: Resume/CV, Job descriptions (set applicable=true but note lower reliability in `notes`)
   REJECT: Payslips, T4s, ID documents, contracts without duties

3. MULTIPLE EMPLOYERS:
   - If multiple DIFFERENT employers → REJECT

4. DUTIES:
   - If fewer than 2 meaningful duties → REJECT

IF MANUAL INPUT:
   - If no clear job title OR fewer than 2 concrete duties → REJECT

VALIDATION RESULT:
- If PASSES: set document_valid=true, rejection_reason=""
- If FAILS: set document_valid=false, write clear rejection_reason. Do NOT populate noc_analysis.

{task_1}

=== TASK 2 — LOCATION OF EXPERIENCE ===
- If the input clearly mentions a Canadian address/province → "canada"
- If it clearly mentions a location outside Canada → "outside_canada"
- If unclear → "unknown"

=== STRICT RULES ===
- NEVER assume missing duties were performed
- NEVER rely on job title alone — always evaluate actual duties
- IGNORE any NOC codes written in the document by the employer. Employers frequently choose the
  wrong NOC. Your job is to independently determine the best match based on DUTIES, not to
  confirm the employer's claim.
- Start your evaluation with the top 3 candidates by `_duty_match_rank`. If the applicant's
  core FUNCTION (e.g., inspecting vs. building vs. supervising vs. designing) does not match
  those candidates' duties, expand your evaluation to ALL provided candidates and briefly
  explain why you rejected the higher-ranked options.
- ALWAYS prefer accuracy over completeness
- KEEP notes/explanations short and precise (1-2 sentences)
- ENSURE consistency with IRCC-style evaluation logic
- Score alternative NOCs honestly as if doing a full dedicated evaluation against each

Your goal is to provide a FAST, RELIABLE, and TRUSTWORTHY NOC suggestion — not a full audit.

=== NOC 2021 DATABASE (Pre-filtered subset — ONLY use codes listed here) ===
{noc_reference}

Output your analysis strictly conforming to the requested JSON schema.
"""


def ocr_from_page_images(page_images: list[tuple[bytes, str]], max_pages: int = 5) -> str:
    """Use GPT-4o-mini vision to OCR text from scanned PDF page images.

    This is a fallback for scanned PDFs where pdfminer returns no text. The cap matches
    pdf_pages_to_images (5): a 3+ page scanned letter must not be silently truncated —
    the audit would otherwise judge duties/signatures it never read.
    """
    if not openai_client or not page_images:
        return ""
    
    content = [{"type": "text", "text": "Extract ALL text visible in this document image. Return ONLY the raw text, no commentary."}]
    for img_bytes, mime_type in page_images[:max_pages]:
        b64 = base64.b64encode(img_bytes).decode('utf-8')
        content.append({
            "type": "image_url",
            "image_url": {"url": f"data:{mime_type};base64,{b64}", "detail": "auto"}
        })
    
    try:
        resp = openai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": content}],
            temperature=0.0,
            max_tokens=2000
        )
        extracted = resp.choices[0].message.content or ""
        print(f"[OCR] Extracted {len(extracted)} chars from {min(len(page_images), max_pages)} page image(s)")
        return extracted
    except Exception as e:
        print(f"[OCR] Failed: {e}")
        return ""


def extract_document_content(doc_bytes: bytes, ext: str, is_image: bool) -> tuple[str, list]:
    """Extract text and page images from any supported document format.
    
    Returns (user_content, page_images) tuple ready for RAG search and AI processing.
    Handles scanned PDF OCR fallback automatically.
    """
    page_images = []
    user_content = ""
    
    if is_image:
        mime_type = IMAGE_MIME_TYPES.get(ext, 'image/jpeg')
        page_images.append((doc_bytes, mime_type))
        # OCR the image so RAG gets real text
        user_content = ocr_from_page_images(page_images)
        if not user_content.strip():
            user_content = "The user uploaded an image of their employment letter. Extract the job title and duties."
    elif ext == '.pdf':
        page_images = pdf_pages_to_images(doc_bytes)
        extracted_text = extract_text_from_pdf(doc_bytes)
        # Scanned PDF fallback: OCR page images if pdfminer returned nothing
        if len(extracted_text.strip()) < 50 and page_images:
            print("[Scanned PDF] Text extraction returned <50 chars, running OCR...")
            extracted_text = ocr_from_page_images(page_images)
        user_content = f"=== EXTRACTED PDF TEXT ===\n{extracted_text}"
    elif ext in ('.docx', '.doc'):
        user_content = f"=== EXTRACTED WORD TEXT ===\n{extract_text_from_docx(doc_bytes)}"
    else:
        user_content = f"=== EXTRACTED TEXT ===\n{doc_bytes.decode('utf-8', errors='replace')}"

    # French letters are valid IRCC documents, but the NOC index + embeddings are English:
    # translate once here so every downstream consumer (Finder, Auditor, re-evals via the
    # persisted text) works at full accuracy without ever re-paying for translation.
    user_content = translate_to_english_if_french(user_content)

    return user_content, page_images


def page_images_only(doc_bytes: bytes, ext: str, is_image: bool) -> list:
    """Render page images WITHOUT any text extraction or OCR. Used by re-evaluations that
    already have the persisted extracted text but still need the visual pages (letterhead /
    signature checks). Returns [] when the file bytes are unavailable (re-eval still proceeds
    on the persisted text)."""
    if not doc_bytes:
        return []
    if is_image:
        return [(doc_bytes, IMAGE_MIME_TYPES.get(ext, 'image/jpeg'))]
    if ext == '.pdf':
        return pdf_pages_to_images(doc_bytes)
    return []


# ── French letter support ───────────────────────────────────────────────────────
# IRCC accepts documents in both official languages; our NOC index and embeddings are
# English-only, so French letters are translated once at extraction time.
_FRENCH_MARKERS = (
    " le ", " la ", " les ", " des ", " une ", " être ", " été ", " avec ", " pour ",
    " dans ", " nous ", " vous ", " ainsi ", " tâches", " fonctions", " emploi ",
    " travail ", " entreprise ", " monsieur ", " madame ", " attestation ", " poste ",
    " responsabilités", " salaire ", " heures ", " semaine ", " depuis ", " chez ",
)
_ENGLISH_MARKERS = (
    " the ", " and ", " with ", " for ", " duties ", " responsibilities ", " employment ",
    " letter ", " company ", " salary ", " hours ", " week ", " since ", " position ",
)


def _looks_french(text: str) -> bool:
    """Cheap heuristic so English letters never pay for a language-detection LLM call."""
    t = f" {(text or '').lower()} "
    fr = sum(t.count(m) for m in _FRENCH_MARKERS)
    en = sum(t.count(m) for m in _ENGLISH_MARKERS)
    return fr >= 5 and fr > en * 1.5


def translate_to_english_if_french(text: str) -> str:
    """Translate a French document to English (one gpt-4o-mini call), preserving structure.
    Returns the text unchanged when it doesn't look French or translation is unavailable.
    The translated text is marked so the auditor knows the original was French (valid for IRCC)."""
    if not text or not openai_client or not _looks_french(text):
        return text
    try:
        resp = openai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": (
                    "Translate this French employment document to English. Preserve the structure, "
                    "line breaks, names, dates, numbers and job titles exactly. Return ONLY the translation.")},
                {"role": "user", "content": text},
            ],
            temperature=0.0,
        )
        translated = (resp.choices[0].message.content or "").strip()
        if translated:
            print(f"[French] Letter detected as French — translated to English ({len(translated)} chars)")
            return ("=== TRANSLATED FROM FRENCH (original letter is in French — a valid IRCC language; "
                    "any French page images correspond to this English translation) ===\n" + translated)
    except Exception as e:
        print(f"[French] Translation failed (continuing with original text): {e}")
    return text


def preprocess_duties_for_embedding(user_text: str) -> str:
    """Extract duty-focused text from employment letters/user input.
    
    Strips boilerplate (addresses, dates, signatory blocks) and isolates
    the duty section for embedding. Used by both RAG search and Function Classifier.
    
    Returns cleaned, duty-focused text (max 8000 chars).
    """
    import re
    text = user_text
    # Remove the "=== EXTRACTED ... ===" header
    text = re.sub(r'^===.*?===\s*', '', text)
    
    # --- Attempt 1: Extract just the duty section from employment letters ---
    # Look for common duty section markers in employment letters
    duty_section = None
    duty_markers = [
        # Allow up to 40 chars between "duties" and "included/are/were" to handle
        # patterns like "duties as a Supervisor included the following"
        r'(?i)(?:duties|responsibilities|job duties|main duties|key duties|'
        r'principal duties|role and responsibilities|scope of work|'
        r'duties and responsibilities).{0,40}?(?:included?|are|were|as follows|'
        r'but (?:were|are) not limited to|:)',
    ]
    for marker in duty_markers:
        match = re.search(marker, text)
        if match:
            # Extract from the marker to the end, then trim at common ending markers
            section = text[match.start():]
            # Trim at signatory/closing markers
            end_match = re.search(
                r'(?i)(?:^|\n)\s*(?:if you (?:require|need|have)|sincerely|regards|'
                r'yours truly|please (?:do not hesitate|feel free)|for verification|'
                r'should you (?:require|need)|we wish|authorized signatory|'
                r'managing director|human resource|HR manager)',
                section
            )
            if end_match:
                section = section[:end_match.start()]
            duty_section = section.strip()
            break
    
    # --- Attempt 1b: Bullet-point fallback for documents without headers ---
    # Some employment letters list duties as bullet points without a preamble like
    # "duties included:". Detect these as duty sections.
    if not duty_section or len(duty_section) <= 50:
        bullet_lines = re.findall(r'(?m)^[\s]*[•▪●–\-]\s*.{20,}', text)
        if len(bullet_lines) >= 3:
            duty_section = '\n'.join(bullet_lines)
            print(f"[RAG Preprocess] Extracted {len(bullet_lines)} bullet-point duties as fallback")
    
    # Also extract job title line if present (important context for embedding)
    title_line = ""
    title_match = re.search(
        r'(?i)(?:job title|position|capacity|role|designation)\s*(?:of|as|:|-|–)?\s*(.+)',
        text
    )
    if title_match:
        title_line = title_match.group(0).strip()[:100]
    
    if duty_section and len(duty_section) > 50:
        # Use the focused duty section with the job title
        text = f"{title_line}\n{duty_section}" if title_line else duty_section
        print(f"[RAG Preprocess] Extracted duty section: {len(duty_section)} chars "
              f"(from {len(user_text)} total)")
    else:
        # --- Attempt 2: Aggressive boilerplate stripping ---
        # Remove phone/email/fax lines
        text = re.sub(r'(?m)^.*?(PHONE|TOLL FREE|FAX|www\.|http|@).*$', '', text)
        text = re.sub(r'(?m)^.*?\d{3}[- ]\d{3}[- ]\d{4}.*$', '', text)
        # Remove postal/zip codes (Canadian and Indian PIN codes)
        text = re.sub(r'(?m)^.*?[A-Z]\d[A-Z]\s*\d[A-Z]\d.*$', '', text)
        text = re.sub(r'(?m)^.*?\d{6}.*$', '', text)  # 6-digit PIN codes
        # Remove GSTIN/tax ID lines
        text = re.sub(r'(?m)^.*?(?:GSTIN|GST|PAN|TIN|EIN)\s*[:#]?\s*\w+.*$', '', text)
        # Remove common letter boilerplate
        text = re.sub(r'(?im)^.*(?:to whom it may concern|this is to certify|'
                       r'ref\.?\s*no|dated\s*:|sincerely|regards|yours truly|'
                       r'managing director|authorized signatory|'
                       r'if you require any|please feel free|'
                       r'for verification purposes).*$', '', text)
        # Remove date lines (DD/MM/YYYY, MM.DD.YYYY, etc.)
        text = re.sub(r'(?m)^.*?\d{1,2}[./\-]\d{1,2}[./\-]\d{2,4}.*$', '', text)
        print(f"[RAG Preprocess] No duty section found, used aggressive stripping")
    
    text = re.sub(r'\s+', ' ', text).strip()
    
    text = text[:8000]  # Safe limit for embedding model
    if not text.strip():
        text = "General professional duties"
    
    return text


# Number of NOC candidates the semantic search returns before reranking. Wider =
# higher recall but larger prompts/cost. Tune via backend/tests/eval_noc.py.
RAG_TOP_K = 40



def semantic_search_nocs(user_text: str, top_k: int = RAG_TOP_K) -> dict:
    """Embed the user's text and find the top_k closest NOC codes.

    Uses pre-computed numpy embedding matrix for fast vectorized similarity.
    The candidate window (RAG_TOP_K) trades recall against prompt size/cost; it is
    tuned by the accuracy eval (backend/tests/eval_noc.py), not by individual cases.
    """
    if not openai_client or _NOC_EMB_MATRIX is None:
        raise ValueError("OpenAI client or NOC embeddings not initialized.")
    
    # 1. Pre-process: extract DUTY-FOCUSED text for embedding
    text_to_embed = preprocess_duties_for_embedding(user_text)
        
    response = openai_client.embeddings.create(
        model="text-embedding-3-small",
        input=[text_to_embed]
    )
    user_vector = np.array(response.data[0].embedding)
    
    # 2. Vectorized cosine similarity against pre-computed matrix
    similarities = _NOC_EMB_MATRIX @ user_vector  # (516,) dot products in one operation
    
    # 3. Get top_k indices using argpartition (faster than full sort for large arrays)
    top_indices = np.argpartition(similarities, -top_k)[-top_k:]
    top_indices = top_indices[np.argsort(similarities[top_indices])[::-1]]  # Sort the top_k
    
    # 4. Build a subset of NOC_INDEX using the pre-ordered key list
    top_nocs_dict = {}
    for idx in top_indices:
        idx_key = _NOC_EMB_KEYS[idx]
        entry = NOC_INDEX.get(idx_key)
        if entry and "code" in entry:
            top_nocs_dict[entry["code"]] = entry
    
    # 5. Duty-level reranking: re-score candidates using per-duty embedding similarity
    #    This catches cases where the correct NOC has high individual-duty overlap
    #    but lower whole-document similarity (e.g., "Collection clerks" for debt collectors)
    if _DUTY_EMB_MATRIX is not None and len(top_nocs_dict) > 0:
        top_nocs_dict = _duty_level_rerank(user_vector, top_nocs_dict, user_text=user_text)
    
    return top_nocs_dict


def _duty_level_rerank(user_vector: np.ndarray, candidates: dict, final_k: int = 10, user_text: str = "") -> dict:
    """Re-score and rerank NOC candidates using per-duty embedding similarity,
    blended with lead statement similarity for industry/employer-type alignment.
    
    For each candidate NOC:
    1. Compute cosine similarity between user text and EACH individual duty
    2. duty_score = avg_similarity × duty_coverage
    3. Compute lead statement similarity (captures employer type, e.g., 'employed by
       collection agencies') using pre-computed lead statement embeddings
    4. composite = duty_score × 0.6 + lead_sim × 0.4
    5. Rerank by composite score
    6. Inject duty coverage stats into each entry so the AI sees them in the prompt
    
    Cost: ZERO additional API calls. Pure numpy on pre-computed vectors.
    """
    COVERAGE_THRESHOLD = 0.30  # Minimum similarity for a duty to count as "matched"
    DUTY_WEIGHT = 0.6          # Weight for duty-level score
    LEAD_WEIGHT = 0.4          # Weight for lead statement similarity
    
    # Compute similarity between user vector and ALL duty vectors at once
    all_duty_sims = _DUTY_EMB_MATRIX @ user_vector  # (N_duties,) — fast vectorized
    
    # Compute lead statement similarities using a FOCUSED duty+employer embedding.
    # The full user_vector includes boilerplate that dilutes the lead statement signal.
    # We extract duty-related lines AND employer context, then embed them together.
    # The employer context is critical: the employer's name/industry distinguishes
    # occupations that share duties but operate in different industries (the same duties
    # at different employer types can map to different NOCs).
    # Cost: 1 extra embedding call (~$0.0001).
    lead_sims = {}
    if _LEAD_EMB_MATRIX is not None and openai_client:
        import re
        lines = user_text.split('\n') if user_text else []
        
        # Extract employer context: company name, RE: line, first few content lines
        employer_lines = []
        duty_lines = []
        for line in lines:
            stripped = line.strip()
            if not stripped or len(stripped) < 5:
                continue
            # Employer context: company names, RE: lines, "attention" lines
            if re.match(r'(?i)(RE:|ATTN|attention|dear|to whom|letter of employment)', stripped):
                employer_lines.append(stripped)
            elif re.match(r'^[A-Z][A-Z\s&.,]+$', stripped) and len(stripped) < 60:
                # ALL-CAPS lines are often company names
                employer_lines.append(stripped)
            # Duty lines
            # Occupation-neutral action verbs + structural cues that mark a line as a
            # duty line. Kept deliberately general — domain meaning is captured by the
            # embeddings, so we avoid biasing toward any single occupation's vocabulary.
            duty_keywords = re.compile(
                r'(?i)(responsible|duties|functions|tasks|include|perform|manage|develop|'
                r'maintain|coordinate|review|prepare|process|conduct|provide|assist|'
                r'monitor|analyze|ensure|create|implement|administer|negotiate|'
                r'contact|resolve|recommend|advise|report|inspect|test|operate|install|'
                r'client|customer|'
                r'has been employed|position of|role of|job title)',
            )
            if len(stripped) > 15 and duty_keywords.search(stripped):
                duty_lines.append(stripped)
        
        # Combine: employer context FIRST (gives industry signal), then duties
        context_text = ' '.join(employer_lines[:5]) + ' | ' + ' '.join(duty_lines)
        context_text = context_text[:4000]
        
        if context_text.strip():
            try:
                ctx_resp = openai_client.embeddings.create(
                    model="text-embedding-3-small",
                    input=[context_text]
                )
                ctx_vector = np.array(ctx_resp.data[0].embedding)
                all_lead_sims = _LEAD_EMB_MATRIX @ ctx_vector
                for code in candidates:
                    idx = _LEAD_CODE_TO_IDX.get(code)
                    if idx is not None:
                        lead_sims[code] = float(all_lead_sims[idx])
            except Exception as e:
                print(f"[Lead Matching] Context embedding failed: {e}. Falling back.")
                all_lead_sims = _LEAD_EMB_MATRIX @ user_vector
                for code in candidates:
                    idx = _LEAD_CODE_TO_IDX.get(code)
                    if idx is not None:
                        lead_sims[code] = float(all_lead_sims[idx])
        else:
            all_lead_sims = _LEAD_EMB_MATRIX @ user_vector
            for code in candidates:
                idx = _LEAD_CODE_TO_IDX.get(code)
                if idx is not None:
                    lead_sims[code] = float(all_lead_sims[idx])
    
    scored = []
    for code, entry in candidates.items():
        if code not in _DUTY_RANGES:
            scored.append((0.0, 0, 0, code, entry))
            continue
        
        start, end = _DUTY_RANGES[code]
        duty_sims = all_duty_sims[start:end]
        total = len(duty_sims)
        
        if total == 0:
            scored.append((0.0, 0, 0, code, entry))
            continue
        
        matched = int(np.sum(duty_sims >= COVERAGE_THRESHOLD))
        coverage = matched / total
        avg_sim = float(np.mean(duty_sims))
        
        # Duty-level score: rewards NOCs where MOST duties match, weighted by similarity
        duty_score = avg_sim * coverage
        
        # Lead statement similarity (captures employer/industry alignment)
        lead_sim = lead_sims.get(code, 0.0)
        composite = (duty_score * DUTY_WEIGHT) + (lead_sim * LEAD_WEIGHT)
        
        scored.append((composite, matched, total, code, entry))
    
    # Sort by composite score descending
    scored.sort(key=lambda x: x[0], reverse=True)
    
    # Log the reranking for debugging
    top5 = [(code, f"{score:.4f}") for score, _, _, code, _ in scored[:5]]
    print(f"[Duty Rerank] Top 5: {top5}")
    
    # Inject duty coverage AND lead statement alignment into each entry so the AI sees them
    reranked = {}
    for rank, (score, matched, total, code, entry) in enumerate(scored[:final_k]):
        enriched = dict(entry)  # shallow copy to avoid mutating the original
        coverage_pct = round((matched / total) * 100) if total > 0 else 0
        enriched["_duty_match_rank"] = rank + 1
        enriched["_pre_computed_duty_coverage_pct"] = coverage_pct
        enriched["_pre_computed_duties_matched"] = matched
        enriched["_pre_computed_duties_total"] = total
        # Inject lead statement match score so the AI can see employer/industry alignment
        lead_sim = lead_sims.get(code, 0.0)
        enriched["_lead_statement_match"] = round(lead_sim, 4)
        reranked[code] = enriched
    
    # The reranked dict (top `final_k`, each carrying `_duty_match_rank`) is the
    # ranking signal used by find_noc_with_openai's best-of-3 consensus. It is
    # returned to the caller and threaded through explicitly rather than stashed in
    # module-global state, which was not safe under concurrent requests.
    return reranked


def semantic_duty_precision(applicant_duties: list, code: str) -> tuple:
    """Applicant-side semantic precision of a NOC match.

    For EACH applicant duty, find the maximum cosine similarity to ANY official duty of
    `code`, then return (mean_of_those_maxes, [per_duty_max, ...]).

    This answers "does each thing the applicant actually does map into this NOC?" — which is
    the right question for classification. Crucially it is robust to MULTI-TITLE NOC groups
    (e.g. 51114 Translators / terminologists / interpreters): an interpreter's duties all map
    strongly to the interpreter sub-duties, and the translator/terminologist duties the person
    does NOT do never enter the denominator. Contrast with recall-style coverage
    (matched / total_NOC_duties), which deflates every match and collapses on multi-title groups.

    Uses the pre-computed per-duty embedding matrix; one batch embedding call for the applicant
    duties (~$0.00002). Returns (0.0, []) if embeddings/inputs are unavailable.
    """
    if _DUTY_EMB_MATRIX is None or code not in _DUTY_RANGES or not openai_client:
        return (0.0, [])
    duties = [d.strip() for d in (applicant_duties or []) if d and d.strip()][:15]
    if not duties:
        return (0.0, [])
    try:
        resp = openai_client.embeddings.create(model="text-embedding-3-small", input=duties)
    except Exception as e:
        print(f"[semantic_duty_precision] embedding failed: {e}")
        return (0.0, [])
    A = np.array([d.embedding for d in resp.data], dtype=np.float32)
    A /= (np.linalg.norm(A, axis=1, keepdims=True) + 1e-9)
    start, end = _DUTY_RANGES[code]
    D = _DUTY_EMB_MATRIX[start:end].astype(np.float32)
    D /= (np.linalg.norm(D, axis=1, keepdims=True) + 1e-9)
    per_app_max = (A @ D.T).max(axis=1)   # best NOC duty for each applicant duty
    return (float(per_app_max.mean()), [float(x) for x in per_app_max])


# ── NOC match confidence — SINGLE SOURCE OF TRUTH (used by BOTH Finder and Auditor) ──
# Calibration maps mean applicant-side per-duty similarity → 0-100%. Both tools call this one
# function so the "NOC match confidence" they display for the same letter+code is identical.
# Correct matches land ~0.41-0.71, wrong pairings ~0.24-0.34 (clean separation).
SEM_FIT_FLOOR = 0.25    # mean similarity at/below this → 0% (clearly not this occupation)
SEM_FIT_CEIL = 0.60     # mean similarity at/above this → 100%
STRONG_DUTY_SIM = 0.40  # an individual applicant duty counts as "aligned" for the display fraction


def noc_match_confidence(applicant_duties: list, code: str):
    """Calibrated NOC-match confidence (0-100) + (aligned, total) display fraction for `code`,
    based on applicant-side semantic precision. Returns None if the embedding signal is
    unavailable, so callers can fall back to their own heuristic.

    This is the ONE definition of 'how well do these duties fit this NOC', shared by the NOC
    Finder and the Letter Auditor to keep their headline NOC confidence consistent.
    """
    mean_max, per = semantic_duty_precision(applicant_duties, code)
    if not per:
        return None
    pct = (mean_max - SEM_FIT_FLOOR) / (SEM_FIT_CEIL - SEM_FIT_FLOOR)
    confidence = int(min(100, max(0, round(pct * 100))))
    aligned = sum(1 for x in per if x >= STRONG_DUTY_SIM)
    return confidence, aligned, len(per)


def scoped_duty_coverage(applicant_duties: list, code: str, llm_duties_match: list = None, role_hint: str = "") -> dict | None:
    """IRCC duty coverage scoped to the applicant's sub-occupation, for MULTI-TITLE NOC groups.

    Many NOC codes pack several distinct occupations under one code (e.g. 51114 = translators /
    terminologists / interpreters). Coverage against the COMBINED duty list unfairly penalizes an
    applicant for not performing the OTHER occupations' duties. The official NOC data already stores
    the split in `duty_groups` (each with a `sub_title`), so we deterministically:
      1. pick the sub-group the applicant best fits (applicant-side precision), and
      2. measure coverage against ONLY that sub-group's duties.

    Numerator respects the auditor LLM's strong/partial evidence judgment where a duty can be matched
    by text; otherwise falls back to semantic similarity. Returns None for single-title NOCs (so the
    auditor's normal coverage is left completely untouched) or when the signal is unavailable.
    """
    entry = NOC_CODE_TO_ENTRY.get(code)
    if not entry:
        return None
    groups = entry.get("duty_groups") or []
    if len(groups) <= 1:                      # single-title — do NOT touch normal coverage
        return None
    if code not in _DUTY_RANGES or _DUTY_EMB_MATRIX is None or not openai_client:
        return None
    duties = [d.strip() for d in (applicant_duties or []) if d and d.strip()][:15]
    if not duties:
        return None

    start, end = _DUTY_RANGES[code]
    if (end - start) != sum(len(g.get("duties", [])) for g in groups):
        return None                            # structure/embedding mismatch — bail safely

    try:
        resp = openai_client.embeddings.create(model="text-embedding-3-small", input=duties)
    except Exception as e:
        print(f"[scoped_duty_coverage] embedding failed: {e}")
        return None
    A = np.array([d.embedding for d in resp.data], dtype=np.float32)
    A /= (np.linalg.norm(A, axis=1, keepdims=True) + 1e-9)
    Dall = _DUTY_EMB_MATRIX[start:end].astype(np.float32)
    Dall /= (np.linalg.norm(Dall, axis=1, keepdims=True) + 1e-9)

    # ── Choose the sub-occupation ────────────────────────────────────────────
    # The applicant's JOB TITLE is the most reliable signal for which sub-occupation applies
    # (an "Interpreter" -> the Interpreters sub-group). Embedding precision alone is unreliable here
    # because translator/terminologist/interpreter duties sit extremely close in vector space, so we
    # use the title as the PRIMARY signal and fall back to precision only when the title is uninformative.
    def _stems(text):
        out = set()
        for t in re.findall(r"[a-z]+", (text or "").lower()):
            if len(t) < 4 or t in ("and", "the", "for", "with"):
                continue
            out.add(t[:-1] if t.endswith("s") else t)   # crude singularize (interpreters -> interpreter)
        return out
    role_stems = _stems(role_hint)

    candidates = []
    off = 0
    for g in groups:
        gd = [x for x in g.get("duties", []) if x and x.strip()]
        n = len(g.get("duties", []))
        if not gd:
            off += n
            continue
        Dg = Dall[off:off + len(gd)]
        off += n
        precision = float((A @ Dg.T).max(axis=1).mean())   # how well applicant fits this sub-group
        sub_stems = _stems(g.get("sub_title", ""))
        # fraction of the sub-title's OWN words present in the role title (1.0 = "Interpreters" fully in "... Interpreter")
        title_score = (len(sub_stems & role_stems) / len(sub_stems)) if sub_stems else 0.0
        candidates.append({"precision": precision, "title_score": title_score,
                           "sub_title": g.get("sub_title", ""), "duties": gd, "Dg": Dg})
    if not candidates:
        return None

    best_title = max(c["title_score"] for c in candidates)
    if best_title > 0:
        # A sub-title matches the role — pick it (break ties by precision).
        best = max((c for c in candidates if c["title_score"] == best_title), key=lambda c: c["precision"])
    else:
        best = max(candidates, key=lambda c: c["precision"])

    # Numerator: count evidenced duties in the chosen sub-group.
    gd, Dg = best["duties"], best["Dg"]
    sim_app = (Dg @ A.T).max(axis=1)
    lookup = {}
    for m in (llm_duties_match or []):
        nd = (m.get("noc_duty", "") or "").strip().lower()
        if nd:
            lookup[nd] = m.get("match_strength", "")
    pairs = []
    for i, dt in enumerate(gd):
        strength = lookup.get(dt.strip().lower())
        if strength in ("strong", "partial", "weak", "missing"):
            pairs.append((dt, strength))
        elif float(sim_app[i]) >= STRONG_DUTY_SIM:
            pairs.append((dt, "strong"))        # no LLM entry to match by text — strong semantic match
        else:
            pairs.append((dt, "missing"))
    # A duty the applicant CLEARLY demonstrates must never be dropped just because it belongs to a
    # different sub-title of a multi-title NOC (e.g. a procurement officer's "respond to customer
    # inquiries" duty). Add any such evidenced official duty from the OTHER sub-groups to the applicable
    # set so it stays in the table and the coverage %. (Non-evidenced other-sub-title duties are still
    # excluded — that is the whole point of scoping.)
    gd_lower = {d.strip().lower() for d in gd}
    for g in groups:
        for d in g.get("duties", []):
            dl = (d or "").strip().lower()
            if dl and dl not in gd_lower and lookup.get(dl) in ("strong", "partial"):
                pairs.append((d, lookup[dl]))
                gd_lower.add(dl)
    coverage = coverage_pct(pairs)
    essential = [(d, s) for (d, s) in pairs if _is_essential_duty(d)] or pairs
    covered = sum(1 for (_d, s) in essential if s in ("strong", "partial"))
    return {"coverage": coverage, "sub_title": best["sub_title"],
            "group_size": len(essential), "covered": covered,
            "applicable_duties": [d for (d, _s) in pairs]}


def _duty_stems(text):
    """Crude word stems (singularized, stop-words/short tokens dropped) for title/sub-title matching."""
    out = set()
    for t in re.findall(r"[a-z]+", (text or "").lower()):
        if len(t) < 4 or t in ("and", "the", "for", "with"):
            continue
        out.add(t[:-1] if t.endswith("s") else t)
    return out


def finder_noc_analysis(applicant_duties: list, code: str, role_hint: str = "") -> dict | None:
    """NOC-side duty analysis for the Finder — the SAME basis as the Employment Letter Auditor.

    Coverage = (official NOC duties demonstrated) / (total official NOC duties), scoped to the
    applicant's sub-occupation for multi-title NOC groups (e.g. 51114 -> Interpreters only). The
    per-applicant-duty breakdown and the uncovered official duties (gaps) are computed against the
    SAME scoped duty set, so the gauge, gaps, and breakdown can never contradict each other.

    Returns {coverage, sub_title, covered, total, duties_breakdown:[{letter_duty,noc_duty,match}],
    key_gaps:[official duty,...]} or None when the embedding signal is unavailable.
    """
    entry = get_noc_entry(code)
    if not entry:
        return None
    appd = [d.strip() for d in (applicant_duties or []) if d and d.strip()][:15]
    if not appd or code not in _DUTY_RANGES or _DUTY_EMB_MATRIX is None or not openai_client:
        return None
    groups = entry.get("duty_groups") or []
    # Flattened official duties, aligned with the embedding-matrix rows for this code.
    flat = [d.strip() for g in groups for d in (g.get("duties") or []) if d and d.strip()] if groups else []
    if not flat:
        flat = [d.strip() for d in (entry.get("duties") or []) if d and d.strip()]
    start, end = _DUTY_RANGES[code]
    if (end - start) != len(flat):
        return None  # structure/embedding mismatch — bail safely
    try:
        resp = openai_client.embeddings.create(model="text-embedding-3-small", input=appd)
    except Exception as e:
        print(f"[finder_noc_analysis] embedding failed: {e}")
        return None
    A = np.array([d.embedding for d in resp.data], dtype=np.float32)
    A /= (np.linalg.norm(A, axis=1, keepdims=True) + 1e-9)
    Dall = _DUTY_EMB_MATRIX[start:end].astype(np.float32)
    Dall /= (np.linalg.norm(Dall, axis=1, keepdims=True) + 1e-9)

    # ── Scope to the applicant's sub-occupation for MULTI-TITLE NOCs ─────────────────────────
    # Title is the primary signal (an "Interpreter" -> the Interpreters sub-group); embedding
    # precision is the tiebreaker/fallback — identical logic to scoped_duty_coverage.
    if len(groups) > 1:
        role_stems = _duty_stems(role_hint)
        cands, off = [], 0
        for g in groups:
            gd = [x.strip() for x in (g.get("duties") or []) if x and x.strip()]
            n = len(g.get("duties") or [])
            if not gd:
                off += n
                continue
            Dg = Dall[off:off + len(gd)]
            off += n
            precision = float((A @ Dg.T).max(axis=1).mean())
            sub_stems = _duty_stems(g.get("sub_title", ""))
            title_score = (len(sub_stems & role_stems) / len(sub_stems)) if sub_stems else 0.0
            cands.append({"precision": precision, "title_score": title_score,
                          "sub_title": g.get("sub_title", ""), "duties": gd, "Dg": Dg})
        if not cands:
            return None
        best_t = max(c["title_score"] for c in cands)
        chosen = (max((c for c in cands if c["title_score"] == best_t), key=lambda c: c["precision"])
                  if best_t > 0 else max(cands, key=lambda c: c["precision"]))
        set_duties, Dset, sub_title = chosen["duties"], chosen["Dg"], chosen["sub_title"]
    else:
        set_duties, Dset, sub_title = flat, Dall, ""

    # NOC-side, NOC-centric breakdown (same orientation as the Auditor): for each OFFICIAL duty,
    # find the applicant's best evidence and grade it strong/partial/weak/missing by similarity.
    # Coverage = (strong + partial) / total — identical formula to the Auditor. This is the SEMANTIC
    # fallback used only when the LLM grader is unavailable.
    noc_to_app = (Dset @ A.T)  # (n_set, n_app)
    total = len(set_duties)
    breakdown, covered, key_gaps = [], 0, []
    for i, nd in enumerate(set_duties):
        j = int(noc_to_app[i].argmax())
        s = float(noc_to_app[i][j])
        if s >= STRONG_DUTY_SIM:      # 0.40
            m = "strong"
        elif s >= 0.30:
            m = "partial"
        elif s >= 0.22:
            m = "weak"
        else:
            m = "missing"
        if m in ("strong", "partial"):
            covered += 1
        else:
            key_gaps.append(nd)
        breakdown.append({"noc_duty": nd, "letter_evidence": appd[j] if m != "missing" else "", "match": m})
    coverage = coverage_pct([(b["noc_duty"], b["match"]) for b in breakdown]) if total else 0

    return {"coverage": coverage, "sub_title": sub_title, "covered": covered, "total": total,
            "set_duties": set_duties, "duties_breakdown": breakdown, "key_gaps": key_gaps}


def grade_scoped_duties_llm(letter_text: str, set_duties: list) -> list | None:
    """LLM evidence-grade each official NOC duty against the letter — the SAME methodology as the
    Employment Letter Auditor, so the Finder's coverage matches it. Returns one entry per official
    duty: [{noc_duty, letter_evidence, match}] with match in strong/partial/weak/missing, or None."""
    if not openai_client or not set_duties or not letter_text:
        return None
    duties_block = "\n".join(f"{i + 1}. {d}" for i, d in enumerate(set_duties))
    # Used for TEXT input (typed duties), where the full Auditor rejects the input as "not a letter".
    # Mirrors the Auditor's persona + rubric (ai_service _build_prompt_text) as closely as possible and
    # uses gpt-4o so the coverage converges with what the Auditor would produce for the same content.
    system = (
        "You are a STRICT, SKEPTICAL, and FAIR Canadian Immigration Officer auditing work experience "
        "under Express Entry (Canadian Experience Class). Verify, do NOT trust — the burden of proof is "
        "on the applicant, and you do NOT assume a duty was performed unless the text shows it.\n\n"
        "For EACH numbered official NOC main duty, set match strength against the applicant's text:\n"
        "- strong: clear semantic alignment, with specific evidence in the text\n"
        "- partial: related language but vague or incomplete\n"
        "- weak: only tangentially related\n"
        "- missing: no evidence at all\n"
        "PARTIAL = RISK: if the evidence is vague, generic, or merely implied/plausible, it is 'weak' or "
        "'missing', never 'partial'. Match on the SEMANTIC FUNCTION performed, not on shared keywords.\n"
        'Return JSON: {"duties":[{"i":<duty number>,"match":"strong|partial|weak|missing",'
        '"evidence":"<short quote from the text, or empty if missing>"}]} — exactly one entry per duty.'
    )
    user = f"=== OFFICIAL NOC MAIN DUTIES ===\n{duties_block}\n\n=== APPLICANT'S WORK EXPERIENCE ===\n{letter_text}"
    try:
        comp = openai_client.chat.completions.create(
            model="gpt-4o", temperature=0.0, seed=42,
            response_format={"type": "json_object"},
            messages=[{"role": "system", "content": system}, {"role": "user", "content": user}],
        )
        items = json.loads(comp.choices[0].message.content).get("duties", [])
    except Exception as e:
        print(f"[grade_scoped_duties_llm] failed: {e}")
        return None
    by_i = {}
    for it in items:
        if isinstance(it, dict):
            try:
                by_i[int(it.get("i", 0))] = it
            except (TypeError, ValueError):
                pass
    out = []
    for idx, nd in enumerate(set_duties, start=1):
        it = by_i.get(idx, {})
        m = str(it.get("match", "missing")).lower()
        if m not in ("strong", "partial", "weak", "missing"):
            m = "missing"
        out.append({"noc_duty": nd, "letter_evidence": (it.get("evidence") or "") if m != "missing" else "", "match": m})
    return out


# When the NOC Finder runs the full Auditor internally (document path), we stash the complete audit
# here so that clicking "Audit my letter" on the Finder result reuses it instead of paying for a second
# identical audit. Keyed by (finder evaluation/stored_file_id, NOC code); in-memory, one-shot.
_FINDER_AUDIT_CACHE = {}
_FINDER_AUDIT_CACHE_MAX = 200

def cache_finder_audit(file_id: str, code: str, audit: dict) -> None:
    if not file_id or not code or not audit:
        return
    if len(_FINDER_AUDIT_CACHE) >= _FINDER_AUDIT_CACHE_MAX:
        _FINDER_AUDIT_CACHE.pop(next(iter(_FINDER_AUDIT_CACHE)), None)
    _FINDER_AUDIT_CACHE[(file_id, str(code))] = audit

def pop_finder_audit(file_id: str, code: str) -> dict | None:
    return _FINDER_AUDIT_CACHE.pop((file_id, str(code)), None)


def audit_duty_coverage(letter_text: str, code: str, page_images: list = None,
                        model_tier: str = "standard") -> dict | None:
    """Run the FULL Employment Letter Auditor against a fixed NOC code and return its duty coverage +
    breakdown — the EXACT same computation/prompt the Auditor uses, so the Finder converges with it.
    page_images: when provided (document uploads), the audit sees the pages too — making the result
    byte-identical in kind to a direct /analyze audit, so it can be REUSED as the real audit.
    Returns {coverage, sub_title, duties_breakdown:[{noc_duty,letter_evidence,match}], key_gaps} or None
    (None when there is no model, the code is unknown, or the Auditor rejects the input as not a letter)."""
    if not openai_client or not letter_text or not code or code not in NOC_CODE_TO_ENTRY:
        return None
    try:
        top = semantic_search_nocs(letter_text)
    except Exception:
        top = {}
    top[code] = NOC_CODE_TO_ENTRY[code]                       # ensure the locked code is in the reference
    system = _build_prompt_text(json.dumps(top, ensure_ascii=False), code)
    try:
        res = audit_document_with_openai(system, letter_text, page_images, auto_detected_noc=code,
                                         model_tier=model_tier)
    except Exception as e:
        print(f"[audit_duty_coverage] auditor failed: {e}")
        return None
    na = res.get("noc_analysis") or {}
    if not na.get("applicable", True):
        return None                                          # Auditor rejected the input — let caller fall back
    breakdown = []
    for d in (na.get("duties_match") or []):
        ms = str(d.get("match_strength", "missing")).lower()
        if ms not in ("strong", "partial", "weak", "missing"):
            ms = "missing"
        breakdown.append({
            "noc_duty": d.get("noc_duty", "") or "",
            "letter_evidence": (d.get("letter_evidence", "") or "") if ms != "missing" else "",
            "match": ms,
        })
    if not breakdown:
        return None
    cov = na.get("duty_coverage_percentage")
    if cov is None:
        cov = coverage_pct([(b["noc_duty"], b["match"]) for b in breakdown])
    gaps = na.get("missing_critical_duties") or [b["noc_duty"] for b in breakdown if b["match"] in ("weak", "missing")]
    return {"coverage": int(round(cov)), "sub_title": na.get("coverage_subtitle", "") or "",
            "duties_breakdown": breakdown, "key_gaps": [g for g in gaps if g],
            "_full_audit": res}  # complete Auditor result, for "Audit my letter" reuse


def _call_openai_structured(system_prompt: str, user_content: str, page_images: list[tuple[bytes, str]], response_format, label: str = "AI", seed: int = 42) -> dict:
    """Unified OpenAI structured output call. Used by both NOC Finder and Auditor.
    
    Args:
        system_prompt: The system instruction prompt.
        user_content: Extracted text content from the document.
        page_images: Optional list of (image_bytes, mime_type) tuples for vision.
        response_format: Pydantic model class for structured output.
        label: Log label for debugging.
        seed: Seed for reproducibility (default 42). Use different seeds for retries.
    """
    if not openai_client:
        raise ValueError("OPENAI_API_KEY is not set. Please configure it to use GPT-4o-mini.")
    
    messages = [
        {"role": "system", "content": system_prompt}
    ]
    
    user_message_content = []
    if user_content:
        user_message_content.append({"type": "text", "text": user_content})
        
    if page_images:
        for img_bytes, mime_type in page_images:
            base64_img = base64.b64encode(img_bytes).decode('utf-8')
            user_message_content.append({
                "type": "image_url",
                "image_url": {
                    "url": f"data:{mime_type};base64,{base64_img}",
                    "detail": "auto"
                }
            })
            
    messages.append({"role": "user", "content": user_message_content})
    
    print(f"Calling OpenAI gpt-4o-mini for {label}...")
    completion = openai_client.beta.chat.completions.parse(
        model="gpt-4o-mini",
        messages=messages,
        response_format=response_format,
        temperature=0.0,
        seed=seed
    )
    
    result = json.loads(completion.choices[0].message.content)
    return _sanitize_noc_response(result)


def find_noc_with_openai(system_prompt: str, user_content: str, page_images: list[tuple[bytes, str]] = None,
                         ranked_candidates: dict = None) -> dict:
    """NOC Finder: returns NOCFinderResponseSchema.

    Strips any employer-stated NOC codes from user_content before sending to the AI
    to prevent the model from deferring to often-incorrect employer classifications.

    ranked_candidates: the reranked dict returned by semantic_search_nocs for THIS
    request (each entry carries `_duty_match_rank`). Used to break best-of-3 voting
    ties deterministically. Passing it per-call avoids shared global state.
    """
    cleaned_content, stripped = strip_employer_noc_references(user_content)
    if stripped:
        print(f"[NOC Finder] Stripped employer NOC references from input to prevent bias")
    
    from models import NOCFinderResponseSchema
    # Don't send page_images to the NOC Finder — the extracted text is sufficient
    # for duty matching, and images introduce visual bias (employer NOC claims,
    # company logos, etc. that shouldn't influence NOC selection).
    # Images ARE still sent for the auditor where visual verification matters.
    result = _call_openai_structured(system_prompt, cleaned_content, None, NOCFinderResponseSchema, "NOC Finder")
    
    # POST-PROCESSING: Best-of-3 consensus for determinism.
    # gpt-4o-mini is non-deterministic even with temperature=0 and seed=42.
    # Run 3 calls with different seeds, collect all unique NOC picks, and select
    # the one with the best duty+lead rank. This is deterministic because the
    # composite rankings are pre-computed and stable.
    # Cost: 2 extra gpt-4o-mini calls (~$0.002).
    result2 = _call_openai_structured(system_prompt, cleaned_content, None, NOCFinderResponseSchema, "NOC Finder (vote 2)", seed=123)
    result3 = _call_openai_structured(system_prompt, cleaned_content, None, NOCFinderResponseSchema, "NOC Finder (vote 3)", seed=7)
    
    votes = [
        (result, result.get("recommended_noc", {}).get("code")),
        (result2, result2.get("recommended_noc", {}).get("code")),
        (result3, result3.get("recommended_noc", {}).get("code")),
    ]
    codes = [c for _, c in votes if c]
    unique_codes = list(set(codes))
    
    ranked_candidates = ranked_candidates or {}
    if len(unique_codes) > 1 and ranked_candidates:
        # Votes disagree — pick the result whose NOC has the best duty+lead rank
        def get_rank(code):
            entry = ranked_candidates.get(code)
            return entry.get("_duty_match_rank", 999) if entry else 999
        
        best_code = min(unique_codes, key=get_rank)
        best_result = next(r for r, c in votes if c == best_code)
        ranks = {c: get_rank(c) for c in unique_codes}
        print(f"[NOC Finder] CONSENSUS: votes={codes}, ranks={ranks} → picking {best_code} (rank #{get_rank(best_code)})")
        result = best_result
    else:
        print(f"[NOC Finder] CONSENSUS: all votes agree on {unique_codes[0] if unique_codes else '?'}")
    
    return result


def auto_detect_noc(user_content: str, page_images: list[tuple[bytes, str]] = None,
                    from_document: bool = False, model_tier: str = "standard",
                    content_key: str = None) -> str | None:
    """Run the NOC Finder (v2) to auto-detect the best NOC code for a document.

    Returns the detected NOC code string, or None if detection fails.
    Side effects: caches the full v2 confidence on `auto_detect_noc.last_confidence` so the Auditor
    can display the SAME NOC-match confidence the Finder would, for free; and stashes the complete
    internal audit (when v2 ran one — document path) on `auto_detect_noc.last_audit` so /analyze can
    REUSE it instead of paying for a second identical audit.
    content_key: stable source hash so the Auditor and the NOC Finder share one cached v2 result for
    the same letter (identical NOC/coverage/alternatives/flags across both tools).
    """
    auto_detect_noc.last_confidence = None
    auto_detect_noc.last_audit = None
    try:
        import noc_finder_v2
        result = noc_finder_v2.run_noc_finder_v2(user_content, page_images,
                                                 from_document=from_document, model_tier=model_tier,
                                                 content_key=content_key)

        rec = result.get("recommended_noc", {})
        detected_code = rec.get("code")
        detected_title = rec.get("title", "?")
        confidence = rec.get("confidence", 0)

        if detected_code and detected_code != "00000":
            print(f"[Auto-Detect NOC] Detected: {detected_code} ({detected_title}) — {confidence}% confidence")
            auto_detect_noc.last_confidence = confidence
            audit = result.get("_audit_full")
            if isinstance(audit, dict):
                import copy as _copy
                auto_detect_noc.last_audit = _copy.deepcopy(audit)
            return detected_code
        else:
            print("[Auto-Detect NOC] Failed to detect NOC code from result")
            return None
    except Exception as e:
        print(f"[Auto-Detect NOC] Error: {e}")
        return None


def _call_claude_structured(system_prompt: str, user_content: str, page_images: list[tuple[bytes, str]],
                            response_format, label: str = "AI",
                            model: str = "claude-haiku-4-5-20251001") -> dict:
    """Claude structured-output call mirroring _call_openai_structured: the schema is enforced by
    forcing a single tool call whose input_schema is the Pydantic model's JSON schema. Raises on
    any failure so callers can fall back to the OpenAI path."""
    import noc_agents
    client = noc_agents._get_anthropic_client()
    if client is None:
        raise ValueError("ANTHROPIC_API_KEY not configured")

    content = []
    if user_content:
        content.append({"type": "text", "text": user_content})
    for img_bytes, mime_type in (page_images or []):
        content.append({
            "type": "image",
            "source": {"type": "base64", "media_type": mime_type,
                       "data": base64.b64encode(img_bytes).decode("utf-8")},
        })

    print(f"Calling Claude {model} for {label}...")
    resp = noc_agents._anthropic_create(
        client,
        model=model,
        max_tokens=8000,   # AnalysisResponse with a full duty-by-duty table is large
        temperature=0.0,   # dropped automatically for models that deprecate it (e.g. Sonnet 5)
        system=system_prompt,
        messages=[{"role": "user", "content": content}],
        tools=[{
            "name": "emit_analysis",
            "description": ("Emit the complete analysis strictly conforming to the schema. "
                            "Include EVERY field — use empty arrays/strings for fields with nothing to report."),
            "input_schema": response_format.model_json_schema(),
        }],
        tool_choice={"type": "tool", "name": "emit_analysis"},
    )
    tool_input = next((b.input for b in resp.content if getattr(b, "type", "") == "tool_use"), None)
    if tool_input is None:
        raise ValueError("Claude returned no tool_use block")
    _backfill_empty_fields(tool_input, response_format)
    result = response_format.model_validate(tool_input).model_dump()
    return _sanitize_noc_response(result)


def _backfill_empty_fields(data: dict, response_format) -> None:
    """Claude's tool-use omits required-but-empty container fields at ANY depth (e.g. top-level
    refusal_reasons: [] or nested noc_analysis.alternative_nocs: []). Recursively backfill missing
    list/dict fields with empties so Pydantic validation passes; anything else still missing fails
    validation and triggers the model fallback."""
    from typing import get_origin, get_args
    from pydantic import BaseModel
    if not isinstance(data, dict) or not hasattr(response_format, "model_fields"):
        return
    for name, field in response_format.model_fields.items():
        ann = field.annotation
        # Unwrap Optional[...] / Union[..., None] to the first concrete arg.
        args = [a for a in get_args(ann) if a is not type(None)]
        core = args[0] if (get_origin(ann) is not None and args and get_origin(ann) is not list
                           and get_origin(ann) is not dict) else ann
        origin = get_origin(core) or core
        if name not in data:
            if origin is list:
                data[name] = []
            elif origin is dict:
                data[name] = {}
        elif isinstance(core, type) and issubclass(core, BaseModel) and isinstance(data.get(name), dict):
            _backfill_empty_fields(data[name], core)  # recurse into nested models (e.g. noc_analysis)


# Objective contact-coordinate patterns. A hit is unambiguous evidence, so it can only ADD a present
# element, never remove one.
_EMAIL_RE = re.compile(r"[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}")
_PHONE_RE = re.compile(r"(?:\+?\d[\d\s().\-]{7,}\d)")
_URL_RE = re.compile(r"(?:https?://|www\.)[A-Za-z0-9.\-]+\.[A-Za-z]{2,}", re.I)


def _apply_contact_backstop(result: dict, text: str) -> None:
    """False->True only, using objective evidence, for the presence elements the model misjudges.

    - contact_information: an email OR a phone number OR a website is, by definition, contact info.
    - company_letterhead: a reference letter that embeds the employer's own email / phone / website
      is in practice printed on letterhead (that header block IS the letterhead). We treat the same
      objective coordinates as letterhead evidence — this only ever corrects a false negative on a
      real company header, and reference letters that are genuinely letterhead-less almost never carry
      an embedded employer email/phone.
    Keeps compliance.missing_element lists honest by dropping anything we just flipped to present.
    """
    mr = result.get("mandatory_requirements")
    if not isinstance(mr, dict) or not text:
        return
    has_email = bool(_EMAIL_RE.search(text))
    has_phone = bool(_PHONE_RE.search(text))
    has_url = bool(_URL_RE.search(text))
    has_contact_coord = has_email or has_phone or has_url

    flipped = []
    if has_contact_coord and mr.get("contact_information") is not True:
        mr["contact_information"] = True
        flipped.append(("contact_information", "contact"))
    if has_contact_coord and mr.get("company_letterhead") is not True:
        mr["company_letterhead"] = True
        flipped.append(("company_letterhead", "letterhead"))

    if flipped:
        kinds = "+".join(k for _, k in flipped)
        print(f"[Auditor] Contact backstop: found email={has_email} phone={has_phone} url={has_url} "
              f"-> forced {kinds} = True")
        # Scrub the flipped elements from any 'missing' list so the narrative doesn't contradict the flags.
        comp = result.get("compliance")
        if isinstance(comp, dict):
            for key in ("missing_elements", "missing"):
                lst = comp.get(key)
                if isinstance(lst, list):
                    comp[key] = [m for m in lst if not _mentions_flipped(m, flipped)]


def _mentions_flipped(text: str, flipped) -> bool:
    low = (text or "").lower()
    for _, kind in flipped:
        if kind == "contact" and ("contact" in low or "email" in low or "phone" in low or "telephone" in low):
            return True
        if kind == "letterhead" and "letterhead" in low:
            return True
    return False


def audit_document_with_openai(system_prompt: str, user_content: str, page_images: list[tuple[bytes, str]] = None, auto_detected_noc: str = None, forced_noc: str = None, model_tier: str = "standard") -> dict:
    """Employment Auditor: returns AnalysisResponse.

    Args:
        auto_detected_noc: If set, this NOC was auto-detected (not user-specified).
            Employer NOC references will be stripped from the text to prevent
            the model from overriding the detected target.
        forced_noc: If set, the user EXPLICITLY requested an audit against this NOC
            (e.g. a re-evaluation against a clicked alternative or a manually typed code).
            Like auto_detected_noc, employer NOC references are stripped and the result's
            detected_code is hard-locked to this code — but the NOC-match confidence is
            computed fresh (we did not run auto-detection, so there is no cached score).
        model_tier: which model grades the audit. gpt-4o-mini is NOT used here — it systematically
            under-grades duty coverage (measured: 57% where the accurate answer is 85%) AND is
            non-deterministic, which broke Finder<->Auditor consistency. Both tiers use a strong,
            stable model so the shared numbers (coverage, flags, alternatives) are accurate and
            identical for every user:
              "standard" -> Claude Haiku 4.5 (accurate + deterministic)
              "premium"  -> Claude Sonnet 4.6 (our most capable model; paid tiers/credits)
            Chain falls back standard<-premium<-gpt-4o-mini so an audit always returns something.
    """
    # The NOC the result must end up locked to (explicit user request wins over auto-detection).
    target_noc = forced_noc or auto_detected_noc
    if target_noc:
        # Strip employer NOC claims so the model can't anchor to the employer's stated NOC and
        # ignore the target. Example: employer writes "NOC 42201" but we are auditing against 41321.
        cleaned_content, stripped = strip_employer_noc_references(user_content)
        if stripped:
            print(f"[Auditor] Stripped employer NOC references (target NOC: {target_noc})")
        user_content = cleaned_content

    from models import AnalysisResponse
    result = None
    # Every letter is graded on Claude Haiku 4.5 — accurate and deterministic. (Sonnet was dropped:
    # it added cost/latency for paid users without an accuracy gain over Haiku on this task.)
    attempts = [(AUDIT_STANDARD_MODEL, "haiku-4-5")]
    for model_id, tag in attempts:
        try:
            result = _call_claude_structured(system_prompt, user_content, page_images, AnalysisResponse,
                                             f"Auditor ({tag})", model=model_id)
            result["_audit_model"] = tag
            break
        except Exception as e:
            print(f"[Auditor] {tag} call failed: {e}")
            result = None
    if result is None:
        print("[Auditor] All Claude audits failed — falling back to gpt-4o-mini (accuracy-degraded).")
        result = _call_openai_structured(system_prompt, user_content, page_images, AnalysisResponse, "Auditor")
        result["_audit_model"] = "gpt-4o-mini"

    # Post-process: if a target NOC was determined (auto-detected OR explicitly requested) but the
    # model returned a different detected_code (e.g. the employer's NOC claim leaked via page images,
    # or the model "second-guessed" the requested code), force-correct it to the intended target.
    if target_noc and result.get("noc_analysis", {}).get("detected_code") != target_noc:
        noc_analysis = result.get("noc_analysis", {})
        model_code = noc_analysis.get("detected_code")
        model_title = noc_analysis.get("detected_title")
        print(f"[Auditor] Correcting detected_code: model said {model_code}, target is {target_noc}")

        # Get the correct title from the index
        target_entry = NOC_CODE_TO_ENTRY.get(target_noc)
        target_title = target_entry.get("title", "") if target_entry else target_noc

        noc_analysis["detected_code"] = target_noc
        noc_analysis["detected_title"] = target_title
        
        # Preserve the model's original pick as an alternative so nothing is lost
        alts = noc_analysis.get("alternative_nocs", [])
        if model_code and not any(a.get("noc_code") == model_code for a in alts):
            reason = (f"This audit was explicitly run against NOC {target_noc} ({target_title}) at your "
                      f"request. The model's own best guess was NOC {model_code} — kept here as an alternative."
                      if forced_noc else
                      f"Employer's stated NOC. Overridden by duty-level analysis which found {target_noc} "
                      f"({target_title}) as a better match.")
            alts.append({
                "noc_code": model_code,
                "noc_title": model_title or model_code,
                "fit_assessment": "moderate",
                "reason": reason,
            })
            noc_analysis["alternative_nocs"] = alts
        
        result["noc_analysis"] = noc_analysis
    
    # Deterministic backstop for the three "presence" mandatory elements the model judges least
    # reliably (letterhead / contact info / signatory). OCR and PDF extraction routinely drop or
    # reorder header blocks, so gpt-4o-mini sometimes marks an obvious letterhead absent. Objective
    # textual evidence (a real email, phone, or website) can only turn these False->True — never the
    # reverse — so it removes false negatives without risking false positives.
    _apply_contact_backstop(result, user_content)

    # Completeness safety net: if the model still omitted any official main duty of the target NOC,
    # add it so the duty-by-duty table and coverage denominator are never silently short a duty.
    _backfill_official_duties(result, target_noc or result.get("noc_analysis", {}).get("detected_code"))

    # Post-process: Fix math errors from the LLM
    # LLMs are notoriously bad at math (e.g., outputting 8 instead of 100 for 8/8 requirements).
    # We recalculate these percentages natively to ensure 100% accuracy.

    # 1. Compliance Score (out of 8 mandatory requirements)
    if "mandatory_requirements" in result and "compliance" in result:
        mand_reqs = result.get("mandatory_requirements", {})
        true_count = sum(1 for v in mand_reqs.values() if v is True)
        result["compliance"]["score"] = int((true_count / 8.0) * 100)
        print(f"[Auditor] Math Fix: Recalculated compliance score to {result['compliance']['score']}% ({true_count}/8)")
        
    # 2. Duty Coverage Percentage
    if "noc_analysis" in result and "duties_match" in result["noc_analysis"]:
        duties = result["noc_analysis"].get("duties_match", [])
        if duties:
            pairs = [(d.get("noc_duty"), d.get("match_strength")) for d in duties]
            new_pct = coverage_pct(pairs)  # binary, essential duties only ("May …" excluded)
            n_ess = len([1 for (dd, _s) in pairs if _is_essential_duty(dd)]) or len(pairs)
            covered = sum(1 for (dd, s) in pairs if _is_essential_duty(dd) and s in ("strong", "partial"))
            old_pct = result["noc_analysis"].get("duty_coverage_percentage")
            if new_pct != old_pct:
                print(f"[Auditor] Duty coverage {old_pct}% -> {new_pct}% "
                      f"({covered}/{n_ess} essential duties evidenced)")
            result["noc_analysis"]["duty_coverage_percentage"] = new_pct

    # ── NOC-match confidence: keep it IDENTICAL to what the NOC Finder shows ──
    # Two distinct numbers, clearly separated:
    #   • noc_match_confidence  = "is this the right NOC?" (semantic precision — same metric & value
    #                             as the Finder, so a user sees the same % in both tools)
    #   • duty_coverage_percentage = "does the letter EVIDENCE enough duties for IRCC?" (recall;
    #                             drives the ACCEPT/PFL/REFUSE decision — auditor-specific)
    na = result.get("noc_analysis")
    if isinstance(na, dict):
        detected = na.get("detected_code", "")
        # Applicant duties: prefer the letter_evidence quotes the auditor mapped; fall back to the
        # official-duty wording it matched. Used for both confidence and sub-group scoping.
        applicant_duties = [(d.get("letter_evidence") or d.get("noc_duty") or "").strip()
                            for d in na.get("duties_match", [])]
        applicant_duties = [d for d in applicant_duties if d and "NOT FOUND" not in d.upper()]

        # (a) NOC-match confidence — keep IDENTICAL to the NOC Finder.
        noc_conf = None
        if auto_detected_noc and detected == auto_detected_noc:
            noc_conf = getattr(auto_detect_noc, "last_confidence", None)  # exact same number as Finder
        if noc_conf is None and detected:
            scored = noc_match_confidence(applicant_duties, detected)
            if scored is not None:
                noc_conf = scored[0]
        if noc_conf is not None:
            na["noc_match_confidence"] = noc_conf

        # (b) Duty coverage — for MULTI-TITLE NOC groups, scope to the applicant's sub-occupation
        # so a perfect (e.g.) interpreter isn't penalized for not doing translator/terminologist duties.
        scoped = scoped_duty_coverage(applicant_duties, detected, na.get("duties_match"),
                                      role_hint=(result.get("role_name") or na.get("detected_title") or ""))
        if scoped is not None:
            old_cov = na.get("duty_coverage_percentage")
            na["duty_coverage_percentage"] = scoped["coverage"]
            na["coverage_subtitle"] = scoped["sub_title"]
            print(f"[Auditor] Multi-title NOC {detected}: scoped duty coverage to "
                  f"'{scoped['sub_title']}' sub-group ({scoped['covered']}/{scoped['group_size']}): "
                  f"{old_cov}% -> {scoped['coverage']}%")
            # Trim the duty-by-duty table + missing list to the applicable sub-group (UI clarity).
            applicable = {d.strip().lower() for d in scoped["applicable_duties"]}
            dm = [m for m in na.get("duties_match", [])
                  if (m.get("noc_duty", "") or "").strip().lower() in applicable]
            if dm:
                na["duties_match"] = dm
            na["missing_critical_duties"] = [
                m.get("noc_duty", "") for m in na.get("duties_match", [])
                if m.get("match_strength") in ("weak", "missing")
            ]
            # Re-derive the coverage-driven verdict — but only ever RELAX it (multi-title deflation
            # can only have made the verdict too harsh). Compliance-driven severity is left intact.
            cov = scoped["coverage"]
            band = "ACCEPT" if cov >= 75 else ("PFL_RISK" if cov >= 50 else "REFUSE")
            rank = {"REFUSE": 0, "PFL_RISK": 1, "ACCEPT": 2}
            if rank.get(band, 0) > rank.get(result.get("decision", "REFUSE"), 0):
                result["decision"] = band
                ra = result.get("risk_assessment")
                if isinstance(ra, dict):
                    ra["overall_risk"] = {"ACCEPT": "low", "PFL_RISK": "moderate", "REFUSE": "high"}[band]
                    ra["pfl_likelihood"] = {"ACCEPT": "low", "PFL_RISK": "medium", "REFUSE": "high"}[band]
                print(f"[Auditor] Verdict relaxed to {band} after multi-title coverage correction.")

    return result


def _resolve_noc_code(code: str, model_title: str) -> str:
    """Resolve the correct NOC code, trusting the model's title over its code when they disagree.
    
    The model is much better at remembering titles than 5-digit codes.
    When code and title disagree, the title is almost always the true intent.
    Returns the resolved code (may be the original or a corrected one).
    """
    if code in NOC_LOOKUP:
        db_title = NOC_LOOKUP[code]
        # Code exists in DB — check if the model's title matches
        if model_title and model_title.lower().strip() != db_title.lower().strip():
            # Mismatch: model said one title but code points to a different one
            resolved = NOC_TITLE_TO_CODE.get(model_title.lower().strip())
            if resolved and resolved != code:
                print(f"[Sanitizer] CODE-TITLE MISMATCH FIX: model said code={code} "
                      f"({db_title}) but title='{model_title}'. "
                      f"Resolved to {resolved} via reverse title lookup.")
                return resolved
        return code
    
    # Code is completely fake — try to resolve from title alone
    if model_title:
        resolved = NOC_TITLE_TO_CODE.get(model_title.lower().strip())
        if resolved:
            print(f"[Sanitizer] FAKE CODE FIX: {code} not in DB, "
                  f"resolved to {resolved} from title '{model_title}'")
            return resolved
    
    # Neither code nor title could be resolved
    print(f"[Sanitizer] UNRESOLVABLE: code={code}, title='{model_title}' — not in DB")
    return code


def _sanitize_noc_response(result: dict, recompute_confidence: bool = True) -> dict:
    """Post-process AI response to fix hallucinated NOC codes/titles.

    gpt-4o-mini sometimes invents NOC codes or pairs real codes with wrong titles.
    This function:
    - Resolves code/title mismatches by trusting the model's title (reverse lookup)
    - Removes alternative NOCs with completely fake codes
    - Validates confidence against reported duty counts

    recompute_confidence:
        True  (default, the v1 LLM path): re-derive confidence as duties_matched / duties_total
              grounded against the official duty count — necessary because the LLM self-reports
              these and can hallucinate.
        False (the v2 path): the caller already computed a calibrated confidence (applicant-side
              semantic precision) and the recall-style matched/total recompute would CLOBBER it —
              and collapse multi-title NOC groups like 51114. Only code/title resolution runs.
    """
    # --- Fix NOC Finder response format ---
    rec = result.get("recommended_noc")
    if rec and isinstance(rec, dict):
        model_title = rec.get("title", "")
        code = rec.get("code", "")

        # Resolve code using title-to-code reverse lookup (always — anti-hallucination)
        resolved_code = _resolve_noc_code(code, model_title)
        rec["code"] = resolved_code

        # Always set title from DB (source of truth)
        if resolved_code in NOC_LOOKUP:
            rec["title"] = NOC_LOOKUP[resolved_code]

        if recompute_confidence:
            # Get actual duties total from database index if possible to prevent LLM hallucinations
            actual_entry = NOC_CODE_TO_ENTRY.get(resolved_code)
            if actual_entry:
                actual_total = len(actual_entry.get("duties", []))
                if actual_total > 0:
                    rec["duties_total"] = actual_total

            # Cap duties_matched to prevent division > 100%
            total = rec.get("duties_total", 0)
            matched = rec.get("duties_matched", 0)
            if matched > total:
                matched = total
                rec["duties_matched"] = matched

            # Calculate true confidence from database-grounded counts
            if total > 0:
                calculated = round((matched / total) * 100)
                calculated = min(100, max(0, calculated))
                rec["confidence"] = calculated

            # Sync root result_type and confidence_level with the calculated confidence
            conf = rec.get("confidence", 0)
            if conf >= 70:
                result["result_type"] = "STRONG_MATCH"
                result["confidence_level"] = "high"
            elif conf >= 45:
                result["result_type"] = "MODERATE_MATCH"
                result["confidence_level"] = "medium"
            else:
                result["result_type"] = "NO_MATCH"
                result["confidence_level"] = "low"
    
    if "alternatives" in result and isinstance(result["alternatives"], list):
        cleaned = []
        for alt in result["alternatives"]:
            model_title = alt.get("title", "")
            code = alt.get("code", "")
            resolved_code = _resolve_noc_code(code, model_title)
            if resolved_code in NOC_LOOKUP:
                alt["code"] = resolved_code
                alt["title"] = NOC_LOOKUP[resolved_code]
                cleaned.append(alt)
            else:
                print(f"[Sanitizer] Removed hallucinated alternative NOC: {code} - {model_title}")
        result["alternatives"] = cleaned
    
    # --- Fix Auditor response format ---
    noc_analysis = result.get("noc_analysis")
    if noc_analysis and isinstance(noc_analysis, dict):
        model_title = noc_analysis.get("detected_title", "")
        code = noc_analysis.get("detected_code", "")
        resolved_code = _resolve_noc_code(code, model_title)
        if resolved_code in NOC_LOOKUP:
            noc_analysis["detected_code"] = resolved_code
            noc_analysis["detected_title"] = NOC_LOOKUP[resolved_code]
        
        if "alternative_nocs" in noc_analysis and isinstance(noc_analysis["alternative_nocs"], list):
            cleaned = []
            for alt in noc_analysis["alternative_nocs"]:
                model_title = alt.get("noc_title", "")
                code = alt.get("noc_code", "")
                resolved_code = _resolve_noc_code(code, model_title)
                if resolved_code in NOC_LOOKUP:
                    alt["noc_code"] = resolved_code
                    alt["noc_title"] = NOC_LOOKUP[resolved_code]
                    cleaned.append(alt)
                else:
                    print(f"[Sanitizer] Removed hallucinated alternative NOC: {code} - {model_title}")
            noc_analysis["alternative_nocs"] = cleaned
    
    return result

# ── Letter Builder Functions ──

def get_noc_details(noc_code: str) -> dict | None:
    """Look up a single NOC code from the loaded index. Returns the entry or None."""
    return NOC_CODE_TO_ENTRY.get(noc_code)


def build_duty_analysis_prompt(noc_entry: dict, user_duty: str) -> str:
    """Builds a lightweight prompt to evaluate ONE duty against ONE NOC's official duties."""
    
    noc_code = noc_entry.get("code", "")
    noc_title = noc_entry.get("title", "")
    lead_statement = noc_entry.get("lead_statement", "")
    duties = noc_entry.get("duties", [])
    duties_text = "\n".join(f"  {i+1}. {d}" for i, d in enumerate(duties))
    
    return f"""You are an IRCC employment letter compliance assistant specializing in NOC 2021.

Your ONLY task: evaluate whether a user-written duty statement aligns with the official duties of NOC {noc_code} ({noc_title}).

=== OFFICIAL NOC {noc_code} INFORMATION ===
Title: {noc_title}
Lead Statement: {lead_statement}
Main Duties:
{duties_text}

=== USER'S DUTY STATEMENT ===
"{user_duty}"

=== YOUR TASK ===

1. Determine which official duty (if any) this user statement most closely aligns with.
2. Rate the alignment: strong (clear match), partial (related but vague), weak (tangential), none (no alignment).
3. If the duty is vague or incomplete, provide up to 3 specific coaching questions to help them make it stronger and more IRCC-compliant.
4. Determine if this duty is "IRCC ready" — meaning it is specific enough, action-oriented, and clearly aligns with the NOC to be included in a formal employment letter.

=== STRICT RULES ===
- Do NOT rewrite the duty for them
- Do NOT invent information the user didn't provide
- Do NOT suggest duties the user didn't mention performing
- If the duty has NO alignment with any official NOC duty, set alignment to "none" and explain clearly
- Coaching questions should ask about SPECIFICS: tools, frequency, scope, outcomes, who they report to
- Be encouraging but honest

Output your analysis strictly conforming to the requested JSON schema.
"""


def analyze_single_duty(user_duty: str, noc_code: str) -> dict:
    """Analyze a single user-written duty against a specific NOC code using AI."""
    from letter_builder_models import DutyAnalysisResponse
    
    noc_entry = get_noc_details(noc_code)
    if not noc_entry:
        raise ValueError(f"NOC code {noc_code} not found in index")
    
    prompt = build_duty_analysis_prompt(noc_entry, user_duty)
    
    print(f"Analyzing duty against NOC {noc_code}: '{user_duty[:60]}...'")
    
    response = gemini_client.models.generate_content(
        model='gemini-2.5-flash',
        contents=prompt,
        config=types.GenerateContentConfig(
            response_mime_type="application/json",
            response_schema=DutyAnalysisResponse,
            temperature=0.0,
        ),
    )
    
    return json.loads(response.text)


def assemble_letter_text(employment_details: dict, noc_code: str, noc_title: str, approved_duties: list) -> dict:
    """Assemble a complete IRCC-compliant employment letter from user-approved data.
    
    This is pure Python template rendering — NO AI call needed.
    Every word in the output was provided or approved by the user.
    """
    import datetime as _dt
    
    d = employment_details
    today = _dt.date.today().strftime("%B %d, %Y")
    
    # Determine employment status phrasing
    is_ongoing = d.get("end_date", "").lower() in ("ongoing", "present", "current", "to date", "")
    if is_ongoing:
        period_phrase = f"since {d['start_date']}"
        status_phrase = f"is currently employed"
        end_display = "Present"
    else:
        period_phrase = f"from {d['start_date']} to {d['end_date']}"
        status_phrase = f"was employed"
        end_display = d['end_date']
    
    # Salary formatting
    salary_str = f"{d['salary_amount']} {d['salary_currency']} {d['salary_period']}"
    
    # Build duties bullet list
    duties_bullets = "\n".join(f"    • {duty['text']}" for duty in approved_duties)
    
    # Assemble intro paragraph
    intro = (
        f"This letter is to confirm that {d['applicant_name']} "
        f"{status_phrase} by {d['company_name']} as a {d['job_title']} "
        f"{period_phrase}."
    )
    
    # Employment details paragraph
    emp_details = (
        f"{d['applicant_name']} works {d['hours_per_week']} hours per week on a "
        f"{d['employment_type']} basis. "
        f"{'Their' if is_ongoing else 'Their'} compensation is {salary_str}. "
        f"The position is based in {d['work_city']}, {d['work_country']}."
    )
    
    # Duties section
    duties_section = (
        f"During {'their' if is_ongoing else 'their'} employment, "
        f"{d['applicant_name']}'s main duties and responsibilities include:\n\n"
        f"{duties_bullets}"
    )
    
    # Closing
    closing = (
        f"Should you require any further information regarding "
        f"{d['applicant_name']}'s employment, please do not hesitate "
        f"to contact the undersigned at {d['supervisor_contact']}."
    )
    
    # Supervisor block
    supervisor_block = (
        f"Sincerely,\n\n"
        f"{d['supervisor_name']}\n"
        f"{d['supervisor_title']}\n"
        f"{d['company_name']}\n"
        f"{d.get('company_address', '')}"
    )
    
    # Full letter
    full_text = (
        f"[COMPANY LETTERHEAD]\n\n"
        f"{today}\n\n"
        f"To Whom It May Concern,\n\n"
        f"{intro}\n\n"
        f"{emp_details}\n\n"
        f"{duties_section}\n\n"
        f"{closing}\n\n"
        f"{supervisor_block}\n\n"
        f"[SIGNATURE]"
    )
    
    # Warnings
    warnings = []
    if len(approved_duties) < 4:
        warnings.append(f"Only {len(approved_duties)} duties provided. IRCC typically expects at least 4 meaningful duties.")
    
    strong_count = sum(1 for d in approved_duties if d.get("alignment") == "strong")
    if strong_count < 2:
        warnings.append("Fewer than 2 duties have strong NOC alignment. Consider strengthening your duty statements.")
    
    return {
        "status": "APPROVED" if len(approved_duties) >= 4 else "INCOMPLETE",
        "noc_code": noc_code,
        "noc_title": noc_title,
        "letter_sections": {
            "header_placeholder": "[COMPANY LETTERHEAD]",
            "date": today,
            "addressee": "To Whom It May Concern,",
            "intro_paragraph": intro,
            "employment_details_paragraph": emp_details,
            "duties_section": duties_section,
            "closing_paragraph": closing,
            "supervisor_block": supervisor_block,
            "signature_placeholder": "[SIGNATURE]",
        },
        "letter_full_text": full_text,
        "warnings": warnings,
    }


# ── ITA Strategy Report Generation ──

def get_draw_context_string() -> str:
    """Loads the draw results JSON from the frontend and returns a context string for the AI prompt."""
    try:
        # Prefer the backend-local copy (present in the deployed Render tree); fall back to the
        # monorepo frontend source for local dev. The frontend file is the source of truth — keep
        # backend/draw_results.json synced when draws are added.
        here = os.path.dirname(__file__)
        candidates = [
            os.path.join(here, 'draw_results.json'),
            os.path.join(here, '..', 'frontend', 'src', 'data', 'draw_results.json'),
        ]
        json_path = next((p for p in candidates if os.path.exists(p)), candidates[0])
        with open(json_path, 'r', encoding='utf-8') as f:
            draws = json.load(f)

        cec_draws = [d for d in draws if d.get('drawType') == 'CEC']
        if not cec_draws:
            return "Latest Express Entry general draw cutoff is approximately 520-540 points."
            
        latest_cec = cec_draws[0]['crsScore']
        recent_cec = cec_draws[:15]
        avg_cec = sum(d['crsScore'] for d in recent_cec) // len(recent_cec)
        
        trend = "stable"
        if len(recent_cec) >= 6:
            older_avg = sum(d['crsScore'] for d in recent_cec[3:6]) / 3
            newer_avg = sum(d['crsScore'] for d in recent_cec[:3]) / 3
            if newer_avg > older_avg + 5:
                trend = "rising"
            elif newer_avg < older_avg - 5:
                trend = "falling"

        french_draws = [d for d in draws if d.get('drawType') == 'French']
        latest_french = french_draws[0]['crsScore'] if french_draws else "Unknown"

        pnp_draws = [d for d in draws if d.get('drawType') == 'PNP']
        latest_pnp = pnp_draws[0]['crsScore'] if pnp_draws else "Unknown"
        
        return (f"CURRENT EXPRESS ENTRY CLIMATE (Provide objective advice based on these numbers, do NOT state statistical probabilities or guarantees):\n"
                f"- Latest CEC (General) Cutoff: {latest_cec}\n"
                f"- 6-Month CEC Average: {avg_cec} (Trend: {trend})\n"
                f"- Latest French Category Cutoff: {latest_french}\n"
                f"- Latest PNP Cutoff: {latest_pnp}\n\n"
                f"You MUST use this data to perform a 'gap analysis'. Tell the user exactly how many points they are away from the 6-month CEC Average, and acknowledge recent trends. "
                f"If their score is very low, mention alternative pathways like French language (NCLC 7) if applicable.")
    except Exception as e:
        print(f"Warning: Could not load draw context: {e}")
        return "The latest Express Entry general draw cutoff is approximately 520-540 points. Category-based draws can have lower cutoffs."

def generate_ita_strategy(raw_inputs: dict, score: dict, breakdown: dict) -> dict:
    """
    Generate a personalized ITA (Invitation to Apply) strategy report
    based on the user's exact CRS profile inputs and calculated score.
    """
    
    # Build a human-readable profile summary for the AI
    profile_lines = []
    profile_lines.append(f"Age: {raw_inputs.get('age', 'Unknown')}")
    profile_lines.append(f"Marital Status: {raw_inputs.get('maritalStatus', 'Unknown')}")
    if raw_inputs.get('maritalStatus') in ['Married', 'Common-Law']:
        profile_lines.append(f"  Spouse is Canadian PR/Citizen: {raw_inputs.get('spouseIsPR', 'Unknown')}")
        profile_lines.append(f"  Spouse accompanying: {raw_inputs.get('spouseAccompanying', 'Unknown')}")
    
    education_labels = {
        'none': 'None or less than secondary',
        'secondary': 'Secondary diploma (high school)',
        'one-year': 'One-year post-secondary',
        'two-year': 'Two-year post-secondary',
        'bachelors': "Bachelor's degree",
        'two-or-more': 'Two or more certificates/degrees',
        'masters': "Master's degree or professional degree",
        'doctoral': 'Doctoral (PhD)',
    }
    profile_lines.append(f"Education: {education_labels.get(raw_inputs.get('education', ''), raw_inputs.get('education', 'Unknown'))}")
    profile_lines.append(f"Canadian Education: {raw_inputs.get('hasCanadianEducation', 'No')}")
    if raw_inputs.get('hasCanadianEducation') == 'Yes':
        can_edu_labels = {'one-two': '1-2 year credential', 'three-plus': '3+ years / Masters / PhD'}
        profile_lines.append(f"  Canadian Credential: {can_edu_labels.get(raw_inputs.get('canadianEducation', ''), 'Unknown')}")
    
    profile_lines.append(f"Primary Language Test: {raw_inputs.get('lang1Test', 'None')}")
    profile_lines.append(f"  Listening: {raw_inputs.get('lang1L', 'N/A')}, Speaking: {raw_inputs.get('lang1S', 'N/A')}, Reading: {raw_inputs.get('lang1R', 'N/A')}, Writing: {raw_inputs.get('lang1W', 'N/A')}")
    
    lang2 = raw_inputs.get('lang2Test', 'None / Not Applicable')
    if lang2 != 'None / Not Applicable':
        profile_lines.append(f"Second Language Test: {lang2}")
        profile_lines.append(f"  Listening: {raw_inputs.get('lang2L', 'N/A')}, Speaking: {raw_inputs.get('lang2S', 'N/A')}, Reading: {raw_inputs.get('lang2R', 'N/A')}, Writing: {raw_inputs.get('lang2W', 'N/A')}")
    else:
        profile_lines.append("Second Language Test: None")
    
    profile_lines.append(f"Canadian Work Experience: {raw_inputs.get('canadianWork', 'None')}")
    profile_lines.append(f"Foreign Work Experience: {raw_inputs.get('foreignWork', 'None')}")
    profile_lines.append(f"Provincial Nomination: {raw_inputs.get('provincialNom', 'No')}")
    profile_lines.append(f"Sibling in Canada: {raw_inputs.get('siblingInCanada', 'No')}")
    profile_lines.append(f"Certificate of Qualification: {raw_inputs.get('certOfQualification', 'No')}")
    
    # Spouse factors
    if raw_inputs.get('spouseAccompanying') == 'Yes' and raw_inputs.get('spouseIsPR') == 'No':
        profile_lines.append(f"Spouse Education: {education_labels.get(raw_inputs.get('spouseEducation', ''), 'Unknown')}")
        profile_lines.append(f"Spouse Language Test: {raw_inputs.get('spLangTest', 'None')}")
        if raw_inputs.get('spLangTest', 'None / Not Applicable') != 'None / Not Applicable':
            profile_lines.append(f"  Spouse L: {raw_inputs.get('spL', 'N/A')}, S: {raw_inputs.get('spS', 'N/A')}, R: {raw_inputs.get('spR', 'N/A')}, W: {raw_inputs.get('spW', 'N/A')}")
        profile_lines.append(f"Spouse Canadian Work: {raw_inputs.get('spouseCanadianWork', 'None')}")
    
    profile_summary = "\n".join(profile_lines)
    
    score_summary = f"""
CRS Score Breakdown:
  Total Score: {score.get('total', 0)} / 1200
  Core/Human Capital: {score.get('core', 0)}
  Spouse Factors: {score.get('spouse', 0)}
  Skill Transferability: {score.get('transferability', 0)}
  Additional Points: {score.get('additional', 0)}

Detailed Core Breakdown:
  Age Points: {breakdown.get('core', {}).get('age', 0)}
  Education Points: {breakdown.get('core', {}).get('education', 0)}
  Official Languages: {breakdown.get('core', {}).get('officialLanguages', 0)}
    First Official Language: {breakdown.get('core', {}).get('firstOfficialLanguage', 0)}
    Second Official Language: {breakdown.get('core', {}).get('secondOfficialLanguage', 0)}
  Canadian Work Experience: {breakdown.get('core', {}).get('canadianWorkExperience', 0)}

Skill Transferability Breakdown:
  Education + Language: {breakdown.get('transferability', {}).get('education', {}).get('languageAndEducation', 0)}
  Education + Canadian Work: {breakdown.get('transferability', {}).get('education', {}).get('canadianWorkAndEducation', 0)}
  Foreign Work + Language: {breakdown.get('transferability', {}).get('foreignWork', {}).get('languageAndForeignWork', 0)}
  Foreign Work + Canadian Work: {breakdown.get('transferability', {}).get('foreignWork', {}).get('canadianAndForeignWork', 0)}
  Certificate of Qualification: {breakdown.get('transferability', {}).get('certificateOfQualification', 0)}

Additional Points Breakdown:
  Provincial Nomination: {breakdown.get('additional', {}).get('provincialNomination', 0)}
  Study in Canada: {breakdown.get('additional', {}).get('studyInCanada', 0)}
  Sibling in Canada: {breakdown.get('additional', {}).get('siblingInCanada', 0)}
  French Language Skills: {breakdown.get('additional', {}).get('frenchLanguageSkills', 0)}
"""

    prompt = f"""You are an expert Canadian immigration consultant specializing in Express Entry and the Comprehensive Ranking System (CRS). You have deep knowledge of all pathways to improve CRS scores, including Provincial Nominee Programs (PNPs), language testing strategies, education credential assessment, and Canadian work experience optimization.

A user has completed their CRS calculation and their EXACT profile is below. Your job is to create a highly personalized, actionable strategy report to help them receive an Invitation to Apply (ITA).

═══════════════════════════════════════
USER PROFILE:
{profile_summary}

{score_summary}
═══════════════════════════════════════

{get_draw_context_string()}

═══════════════════════════════════════
OFFICIAL CRS SCORING REFERENCE
(You MUST use these tables for ALL point calculations. Do NOT guess point values.)
═══════════════════════════════════════

A. CORE / HUMAN CAPITAL FACTORS:

Age (with spouse / without spouse):
  17 or under: 0/0, 18: 90/99, 19: 95/105, 20-29: 100/110,
  30: 95/105, 31: 90/99, 32: 85/94, 33: 80/88, 34: 75/83,
  35: 70/77, 36: 65/72, 37: 60/66, 38: 55/61, 39: 50/55,
  40: 45/50, 41: 35/39, 42: 25/28, 43: 15/17, 44: 5/6, 45+: 0/0

Education (with spouse / without spouse):
  None: 0/0, Secondary: 28/30, One-year post-sec: 84/90,
  Two-year post-sec: 91/98, Bachelor's: 112/120,
  Two or more credentials: 119/128, Master's: 126/135, Doctoral: 140/150

First Official Language PER ABILITY (with spouse / without spouse):
  CLB < 4: 0/0, CLB 4-5: 6/6, CLB 6: 8/9, CLB 7: 16/17,
  CLB 8: 22/23, CLB 9: 29/31, CLB 10+: 32/34
  (Multiply by 4 abilities for total. Max = 128/136)

Second Official Language PER ABILITY (with spouse / without spouse):
  CLB < 5: 0/0, CLB 5-6: 1/1, CLB 7-8: 3/3, CLB 9+: 5/6
  (Multiply by 4 abilities for total. Max = 20/24)

Canadian Work Experience (with spouse / without spouse):
  None: 0/0, 1yr: 35/40, 2yr: 46/53, 3yr: 56/64, 4yr: 63/72, 5+yr: 70/80

B. SPOUSE FACTORS (only when spouse is accompanying and NOT a PR/citizen):

Spouse Education:
  None: 0, Secondary: 2, One-year: 6, Two-year: 7,
  Bachelor's: 8, Two or more: 9, Master's: 10, Doctoral: 10

Spouse Language PER ABILITY:
  CLB < 5: 0, CLB 5-6: 1, CLB 7-8: 3, CLB 9+: 5
  (Multiply by 4 abilities for total. Max = 20)

Spouse Canadian Work:
  None: 0, 1yr: 5, 2yr: 7, 3yr: 8, 4yr: 9, 5+yr: 10

C. ADDITIONAL POINTS:
  Provincial Nomination: +600
  Sibling in Canada (PR/citizen): +15
  Canadian education (1-2yr credential): +15, (3+yr or grad): +30
  French language proficiency: 
    - IF user has NCLC 7+ in all 4 French abilities AND English is CLB 4 or lower (or no English test): +25 points
    - IF user has NCLC 7+ in all 4 French abilities AND English is CLB 5 or higher in all 4 abilities: +50 points
    (Note: You MUST check the user's English CLB levels before applying this bonus. If they have strong English, use 50.)

D. SKILL TRANSFERABILITY FACTORS (Max 100 points total combination limit)
  Education + Language:
    1-yr/2-yr credential + CLB 7/8 = 13 pts. With CLB 9+ = 25 pts.
    Two or more / Masters / PhD + CLB 7/8 = 25 pts. With CLB 9+ = 50 pts.
  Education + Canadian Work Experience:
    1-yr/2-yr credential + 1 yr Cdn Work = 13 pts. With 2+ yr Cdn Work = 25 pts.
    Two or more / Masters / PhD + 1 yr Cdn Work = 25 pts. With 2+ yr Cdn Work = 50 pts.
  Foreign Work + Language:
    1-2 yrs Exp + CLB 7/8 = 13 pts. With CLB 9+ = 25 pts.
    3+ yrs Exp + CLB 7/8 = 25 pts. With CLB 9+ = 50 pts.
  Foreign Work + Canadian Work Experience:
    1-2 yrs Foreign Exp + 1 yr Cdn = 13 pts. With 2+ yr Cdn = 25 pts.
    3+ yrs Foreign Exp + 1 yr Cdn = 25 pts. With 2+ yr Cdn = 50 pts.
  Certificate of Qualification (Trades) + Language:
    With Cert + CLB 5/6 = 25 pts. With CLB 7+ = 50 pts.

═══════════════════════════════════════
STRATEGIC DIRECTIVES & WARNINGS:
1. AGE DECAY WARNING: Check the user's age. If they are exactly 29, emphasize they will start losing 5 points every birthday. If they are >29, warn them that processing delays cost 5-6 points per year so speed is critical.
2. UNACCOMPANYING SPOUSE STRATEGY: If the user's accompanying spouse has low education (no degree) AND low language scores (CLB < 7), you MUST run a simulation and explicitly recommend considering declaring the spouse as 'Non-Accompanying' to leverage the higher single-applicant scoring grid.
3. 100-POINT TRANSFERABILITY CAP: The absolute maximum for ALL Skill Transferability points (Section D) combined is 100 points. Do NOT project point gains that push someone over this cap.

CRITICAL: When recommending actions, you MUST calculate the EXACT point difference
between the user's current level and the target level using these tables.
For example: Spouse going from no language test (0 pts) to CLB 9 all abilities = 5×4 = +20 points, NOT a rough estimate.
═══════════════════════════════════════

Generate a comprehensive, personalized ITA strategy report. You MUST return ONLY a valid JSON object with this exact structure (no markdown, no code fences, just raw JSON):

{{
  "current_score": {score.get('total', 0)},
  "estimated_cutoff": 530,
  "gap": {max(530 - score.get('total', 0), 0)},
  "overall_assessment": "YOUR_ASSESSMENT_HERE",
  "category_based_eligibility": [
    {{
      "category": "CATEGORY_NAME",
      "eligible": true,
      "note": "EXPLANATION"
    }}
  ],
  "actions": [
    {{
      "rank": 1,
      "title": "ACTION_TITLE",
      "description": "DETAILED_DESCRIPTION",
      "potential_points": "+XX to +YY",
      "effort_level": "Low or Medium or High",
      "estimated_timeline": "X-Y months",
      "estimated_cost": "$XXX CAD",
      "priority": "Critical or High or Medium or Low",
      "specific_targets": "SPECIFIC_TARGETS"
    }}
  ],
  "language_optimization": {{
    "current_first_language_points": {breakdown.get('core', {}).get('firstOfficialLanguage', 0)},
    "max_first_language_points": 136,
    "improvement_possible": 0,
    "specific_targets": "TARGET_SCORES",
    "second_language_recommendation": "RECOMMENDATION"
  }},
  "pnp_recommendations": [
    {{
      "province": "PROVINCE",
      "stream": "STREAM_NAME",
      "why_suitable": "REASON",
      "points_impact": "+600",
      "requirements_summary": "REQUIREMENTS"
    }}
  ],
  "timeline_summary": "TIMELINE_PARAGRAPH",
  "disclaimer": "This report provides general guidance based on publicly available CRS criteria and immigration program information. It is not legal advice. For personalized legal advice, consult a Regulated Canadian Immigration Consultant (RCIC) or immigration lawyer. Immigration policies and cutoff scores change frequently — always verify with IRCC's official website."
}}

═══════════════════════════════════════
OFFICIAL LANGUAGE TEST → CLB CONVERSION TABLES
(You MUST use these tables. Do NOT guess conversions.)
═══════════════════════════════════════

CELPIP-General → CLB:
  CELPIP scores map DIRECTLY 1:1 to CLB levels.
  CELPIP 4 = CLB 4, CELPIP 5 = CLB 5, CELPIP 6 = CLB 6,
  CELPIP 7 = CLB 7, CELPIP 8 = CLB 8, CELPIP 9 = CLB 9,
  CELPIP 10 = CLB 10, CELPIP 11 = CLB 11, CELPIP 12 = CLB 12.
  CELPIP uses WHOLE NUMBERS only (no decimals like 8.0 or 7.5).

IELTS General Training → CLB:
  Listening: 4.5=CLB4, 5.0=CLB5, 5.5=CLB6, 6.0=CLB7, 7.5=CLB8, 8.0=CLB9, 8.5=CLB10
  Reading:   3.5=CLB4, 4.0=CLB5, 5.0=CLB6, 6.0=CLB7, 6.5=CLB8, 7.0=CLB9, 8.0=CLB10
  Writing:   4.0=CLB4, 5.0=CLB5, 5.5=CLB6, 6.0=CLB7, 6.5=CLB8, 7.0=CLB9, 7.5=CLB10
  Speaking:  4.0=CLB4, 5.0=CLB5, 5.5=CLB6, 6.0=CLB7, 6.5=CLB8, 7.0=CLB9, 7.5=CLB10

TEF Canada → NCLC:
  Listening: 145-216=NCLC4, 217-248=NCLC5, 249-279=NCLC6, 280-297=NCLC7, 298-315=NCLC8, 316-333=NCLC9, 334-360=NCLC10
  Reading:   121-150=NCLC4, 151-180=NCLC5, 181-206=NCLC6, 207-232=NCLC7, 233-247=NCLC8, 248-262=NCLC9, 263-300=NCLC10
  Writing:   181-225=NCLC4, 226-270=NCLC5, 271-309=NCLC6, 310-348=NCLC7, 349-370=NCLC8, 371-392=NCLC9, 393-450=NCLC10
  Speaking:  181-225=NCLC4, 226-270=NCLC5, 271-309=NCLC6, 310-348=NCLC7, 349-370=NCLC8, 371-392=NCLC9, 393-450=NCLC10

TCF Canada → NCLC:
  Listening: 331-368=NCLC4, 369-397=NCLC5, 398-457=NCLC6, 458-502=NCLC7, 503-522=NCLC8, 523-548=NCLC9, 549-699=NCLC10
  Reading:   342-374=NCLC4, 375-405=NCLC5, 406-452=NCLC6, 453-498=NCLC7, 499-523=NCLC8, 524-548=NCLC9, 549-699=NCLC10
  Writing:   4=NCLC4, 6=NCLC5, 6=NCLC6, 10=NCLC7, 12=NCLC8, 14=NCLC9, 16=NCLC10
  Speaking:  4=NCLC4, 6=NCLC5, 6=NCLC6, 10=NCLC7, 12=NCLC8, 14=NCLC9, 16=NCLC10

═══════════════════════════════════════

RULES:
1. Be SPECIFIC to this user's exact profile. Reference their actual scores, not generic advice.
2. Rank actions from highest impact/easiest to lowest impact/hardest.
3. If the user already has maximum points in a category, acknowledge it and skip it.
4. Be honest — if the gap is very large, say so and suggest realistic pathways.
5. Include at least 4-6 actionable recommendations.
6. For language optimization, use ONLY the conversion tables above. NEVER mix up CELPIP and IELTS scoring. If the user took CELPIP, all targets must be in CELPIP whole numbers. If the user took IELTS, all targets must be in IELTS band scores. State both the test-specific score AND the CLB level.
7. For PNP, suggest 2-3 specific provincial programs they may qualify for based on their profile.
8. All costs should be in CAD.
9. NEVER use decimal scores (e.g., "8.0") for CELPIP — CELPIP uses whole numbers only.
10. NEVER confuse test scoring systems. Double-check every conversion against the tables above.
"""

    try:
        response = gemini_client.models.generate_content(
            model="gemini-2.5-flash",
            contents=[prompt],
            config=types.GenerateContentConfig(
                temperature=0.1,
                max_output_tokens=16384,
                response_mime_type="application/json",
            )
        )
        
        raw_text = response.text.strip()
        
        # Strip markdown code fences if present
        if raw_text.startswith("```"):
            lines = raw_text.split("\n")
            lines = lines[1:]  # Remove first ```json line
            if lines and lines[-1].strip() == "```":
                lines = lines[:-1]
            raw_text = "\n".join(lines)
        
        strategy = json.loads(raw_text)
        return strategy
        
    except json.JSONDecodeError as e:
        print(f"Failed to parse ITA strategy JSON: {e}")
        print(f"Raw response: {raw_text[:1000]}")
        raise ValueError(f"AI returned invalid JSON for ITA strategy: {str(e)}")
    except Exception as e:
        print(f"ITA Strategy generation failed: {e}")
        raise

