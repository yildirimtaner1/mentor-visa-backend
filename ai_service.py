import os
import json
import io
import datetime
import base64
import numpy as np
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from google import genai
from google.genai import types
from models import AnalysisResponse
from dotenv import load_dotenv
from openai import OpenAI

load_dotenv()

# --- Gemini client (used by Letter Builder + ITA Strategy) ---
gemini_client = genai.Client()

# --- OpenAI client (used by NOC Finder + Auditor + OCR) ---
_openai_api_key = os.getenv("OPENAI_API_KEY", "")
openai_client = OpenAI(api_key=_openai_api_key) if _openai_api_key else None

# Load the NOC index once at startup
_noc_index_path = os.path.join(os.path.dirname(__file__), "noc_index.json")
with open(_noc_index_path, "r", encoding="utf-8") as f:
    NOC_INDEX = json.load(f)
print(f"Loaded NOC index: {len(NOC_INDEX)} unit groups")

# Pre-build fast lookup: NOC code -> title (used by sanitizer)
NOC_LOOKUP = {}
for _entry in NOC_INDEX.values():
    _code = _entry.get("code", "")
    if _code:
        NOC_LOOKUP[_code] = _entry.get("title", "")
print(f"Built NOC_LOOKUP: {len(NOC_LOOKUP)} codes")

# Pre-build reverse lookup: title (lowercase) -> code (used by sanitizer for hallucination fix)
NOC_TITLE_TO_CODE = {}
for _code, _title in NOC_LOOKUP.items():
    NOC_TITLE_TO_CODE[_title.lower().strip()] = _code
print(f"Built NOC_TITLE_TO_CODE: {len(NOC_TITLE_TO_CODE)} titles")

# Load NOC embeddings for RAG and pre-compute numpy matrix
NOC_EMBEDDINGS = {}
_NOC_EMB_MATRIX = None  # Pre-computed numpy matrix for fast similarity
_NOC_EMB_KEYS = []      # Ordered list of index keys matching matrix rows
_embeddings_path = os.path.join(os.path.dirname(__file__), "noc_embeddings.json")
if os.path.exists(_embeddings_path):
    with open(_embeddings_path, "r", encoding="utf-8") as f:
        NOC_EMBEDDINGS = json.load(f)
    # Pre-compute numpy matrix: one np.array call instead of 516 per request
    _NOC_EMB_KEYS = list(NOC_EMBEDDINGS.keys())
    _NOC_EMB_MATRIX = np.array([NOC_EMBEDDINGS[k] for k in _NOC_EMB_KEYS])
    print(f"Loaded NOC embeddings: {len(NOC_EMBEDDINGS)} vectors, matrix shape: {_NOC_EMB_MATRIX.shape}")
else:
    print("WARNING: noc_embeddings.json not found. RAG NOC Finder will fail.")

# Load per-duty embeddings for duty-level reranking (zero-cost at query time)
_DUTY_EMB_MATRIX = None   # (N_duties, 1536) pre-computed duty vectors
_DUTY_RANGES = {}         # NOC code -> (start_row, end_row) in the matrix
_duty_emb_path = os.path.join(os.path.dirname(__file__), "noc_duty_embeddings.npz")
_duty_idx_path = os.path.join(os.path.dirname(__file__), "noc_duty_index.json")
if os.path.exists(_duty_emb_path) and os.path.exists(_duty_idx_path):
    _DUTY_EMB_MATRIX = np.load(_duty_emb_path)["embeddings"]
    with open(_duty_idx_path, "r", encoding="utf-8") as f:
        _duty_idx = json.load(f)
    _DUTY_RANGES = {code: tuple(rng) for code, rng in _duty_idx["ranges"].items()}
    print(f"Loaded per-duty embeddings: {_DUTY_EMB_MATRIX.shape[0]} duties across {len(_DUTY_RANGES)} NOCs")
else:
    print("WARNING: Per-duty embeddings not found. Duty-level reranking disabled.")

# Load pre-computed lead statement embeddings for industry/employer-type matching
_LEAD_EMB_MATRIX = None   # (516, 1536) lead statement vectors
_LEAD_EMB_CODES = []      # NOC codes in matrix order
_lead_emb_path = os.path.join(os.path.dirname(__file__), "noc_lead_embeddings.npz")
_lead_idx_path = os.path.join(os.path.dirname(__file__), "noc_lead_index.json")
if os.path.exists(_lead_emb_path) and os.path.exists(_lead_idx_path):
    _LEAD_EMB_MATRIX = np.load(_lead_emb_path)["embeddings"]
    with open(_lead_idx_path, "r", encoding="utf-8") as f:
        _LEAD_EMB_CODES = json.load(f)["codes"]
    # Build code → row index mapping for fast lookup
    _LEAD_CODE_TO_IDX = {code: idx for idx, code in enumerate(_LEAD_EMB_CODES)}
    print(f"Loaded lead statement embeddings: {_LEAD_EMB_MATRIX.shape[0]} NOCs")
else:
    _LEAD_CODE_TO_IDX = {}
    print("WARNING: Lead statement embeddings not found. Lead-weighted reranking disabled.")

# Module-level state for duty-rank enforcement (populated by _duty_level_rerank)
_LAST_DUTY_TOP5 = {}  # Top-5 duty-ranked NOCs from the last search

# MIME type mapping for images
IMAGE_MIME_TYPES = {
    '.jpg': 'image/jpeg',
    '.jpeg': 'image/jpeg',
    '.png': 'image/png',
    '.bmp': 'image/bmp',
    '.tiff': 'image/tiff',
    '.tif': 'image/tiff',
    '.webp': 'image/webp',
}

def extract_text_from_pdf(pdf_bytes: bytes) -> str:
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    text = ""
    for page in doc:
        text += page.get_text() + "\n"
    return text

def pdf_pages_to_images(pdf_bytes: bytes, max_pages: int = 5) -> list[tuple[bytes, str]]:
    """Convert each page of a PDF to a PNG image. Returns list of (image_bytes, mime_type). Limits to max_pages."""
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    images = []
    
    # Only process up to max_pages to prevent API rate limits / huge costs
    pages_to_process = min(len(doc), max_pages)
    
    for i in range(pages_to_process):
        page = doc[i]
        # Render at 2x resolution for better OCR quality
        pix = page.get_pixmap(dpi=200)
        img_bytes = pix.tobytes("png")
        images.append((img_bytes, "image/png"))
        
    print(f"Converted {len(images)} PDF page(s) out of {len(doc)} to images for vision processing")
    return images

def extract_text_from_docx(docx_bytes: bytes) -> str:
    doc = DocxDocument(io.BytesIO(docx_bytes))
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                paragraphs.append(" | ".join(row_text))
    return "\n".join(paragraphs)

def _build_prompt_text(noc_reference: str, target_noc: str = None) -> str:
    """Builds the system instruction for the skeptical immigration officer auditor prompt."""
    
    task_noc_matching = """
=== TASK 1 - NOC DETECTION & ALIGNMENT ===

Read the employment letter carefully. Compare the duties described in the letter against the "duties" arrays in the NOC 2021 database.

NOC MATCHING STRATEGY:
Step 1: Identify TOP 3 candidate NOCs based on duties
Step 2: Select BEST-FIT NOC (highest duty coverage)
Step 3: Provide alternatives with lower confidence

=== CRITICAL: SEMANTIC MATCHING, NOT KEYWORD MATCHING ===
Do NOT match based on surface-level keywords. Focus on the SEMANTIC MEANING of what the person actually does day-to-day.
For example, a "Fire Watch Team Member" who inspects work areas for hazards, documents conditions, and recommends controls is performing SAFETY SPECIALIST duties — not firefighter duties — even though the word "fire" appears frequently. Always ask: "What is this person's core function?" rather than "What keywords appear most often?"

=== NO OVER-RELIANCE ON JOB TITLES ===
Job titles are NOT determinative. Duties override titles. A "Manager" who only does clerical work is NOT performing management duties.

=== TIE-BREAKING RULE ===
If two or more NOC codes score within 5 percentage points of each other, select the NOC whose LEAD STATEMENT most accurately describes the person's primary role. List the close runner-up as the first alternative NOC and explicitly note in the explanation that these two codes are close matches.

- In `noc_analysis.detected_code`, return the 5-digit code.
- In `noc_analysis.detected_title`, return the exact title from the database.
- In `noc_analysis.match_score`, evaluate objective duty coverage out of 100.
- In `noc_analysis.confidence`, rate how confident you are in this selection (0-100). Low match_score + high confidence = sure it's the right NOC but the letter is weak.
- In `noc_analysis.alternative_nocs`, list secondary matches >= 50%.
  **IMPORTANT: Score each alternative NOC as if you were doing a deep, dedicated evaluation against that specific NOC. Do NOT give rough estimates.**
  **CRITICAL: You may ONLY list alternative NOCs whose codes appear in the provided database. NEVER invent or recall NOC codes from memory.**
"""

    if target_noc:
        task_noc_matching = f"""
=== TASK 1 - TARGETED NOC EVALUATION ===
The user explicitly requested to evaluate this document against NOC {target_noc}.
You MUST lock the primary match to NOC {target_noc}. Evaluate how strongly the applicant's duties align specifically with NOC {target_noc}.
- Set `noc_analysis.detected_code` strictly to "{target_noc}".
- Fetch and set `noc_analysis.detected_title` to the exact title for {target_noc} from the database.
- In `noc_analysis.match_score`, evaluate the objective percentage out of 100 showing how well their duties align with {target_noc}. (It is okay if it is low, be honest).
- In `noc_analysis.confidence`, rate how confident you are in this selection (0-100).
- You may still list other better fits in `alternative_nocs`.
- Map their duties strictly against the official duties of NOC {target_noc}. Explain where they overlap and note glaring gaps if any.
"""

    return f"""You are an advanced AI system acting as a STRICT, SKEPTICAL, and FAIR Canadian Immigration Officer auditing employment letters under Express Entry - Canadian Experience Class (CEC).

**IMPORTANT: Today's date is {datetime.date.today().strftime('%B %d, %Y')}. Use this as the current date for any date-related analysis. Do NOT hallucinate or guess a different date.**

PRIMARY OBJECTIVE:
Determine whether the applicant's work experience, as described in the employment letter, would likely be:
1) ACCEPTED
2) FLAGGED FOR PROCEDURAL FAIRNESS LETTER (PFL)
3) REFUSED

You must think like an IRCC officer whose role is to VERIFY, NOT TRUST.
Do NOT assume the applicant qualifies. The burden of proof is on the applicant.

You have been given:
1. A PRE-FILTERED subset of the NOC 2021 database containing the most relevant unit groups for this document. Each entry has an official code, title, lead statement, and main duties.
2. A document uploaded by the user - either as extracted text or as an image.

**CRITICAL RULE: You may ONLY reference NOC codes and titles that appear in the provided database below. Do NOT invent, guess, or recall NOC codes from memory. If a NOC code is not in the database provided, you MUST NOT use it anywhere in your response — not as the primary match, not as an alternative, and not in any narrative text.**

---

=== CORE PRINCIPLES ===

1. BURDEN OF PROOF - Only consider what is explicitly stated or strongly implied in the letter. Do NOT assume missing duties were performed.
2. SKEPTICAL REVIEW - Actively look for weaknesses, ambiguity, and gaps that could justify refusal.
3. DUTY-BASED MATCHING - Match duties semantically, not by keywords. The applicant must demonstrate the MAJORITY (>=70%) of the MAIN DUTIES of a single NOC.
4. PARTIAL MATCH = RISK - If duties are vague, generic, or incomplete, increase refusal risk.

---

=== VALIDATION PIPELINE (PRE-CHECKS - Do this FIRST) ===

Classify each check as HARD_FAIL (stop analysis) or SOFT_FAIL (continue with warnings) or PASS.

CHECK 1 - READABLE CONTENT (HARD_FAIL if blank/corrupted):
- If the document is blank, empty, corrupted, or contains no readable text whatsoever, REJECT.
- Set `document_type` to "Blank / Unreadable Document".

CHECK 2 - DOCUMENT TYPE (HARD_FAIL if wrong type):
Only ACCEPTABLE:
  Employment / Reference / Experience letter issued BY an employer
  Job offer letter that includes duties/responsibilities
  A single letter covering multiple roles at the SAME company

NOT ACCEPTABLE (REJECT with clear explanation):
  Payslips, T4/T4A slips, ROE, tax/payroll documents
  Resumes or CVs (self-authored)
  Cover letters (applicant-authored)
  Job postings / advertisements
  LinkedIn profile screenshots
  Contracts without duties section
  Bank statements, invoices, receipts, ID documents
  Business cards, certificates, diplomas
  Unrelated photos or images

CHECK 3 - MULTIPLE EMPLOYERS (HARD_FAIL):
- If letters from TWO OR MORE DIFFERENT employers/companies are merged into one file, REJECT.
- A single letter covering multiple ROLES at the SAME company is perfectly valid.

CHECK 4 - LANGUAGE (SOFT_FAIL):
- If NOT in English or French, add a medium-severity risk. Still attempt analysis if duties are discernible.

CHECK 5 - DUTIES QUALITY (SOFT_FAIL):
- If fewer than 2 duties, flag as high-severity risk.
- If duties appear copy-pasted VERBATIM from the NOC website, flag as high-severity risk.

CHECK 6 - HEAVY REDACTION (SOFT_FAIL):
- If significant portions are redacted, add a warning noting which elements could not be verified.

=== REJECTION OUTPUT FORMAT ===
If ANY of Checks 1-3 result in HARD_FAIL:
- Set `decision` to "REFUSE", `confidence_score` to 95+.
- Set `noc_analysis.applicable` to false, `noc_analysis.detected_code` to "", `noc_analysis.match_score` to 0.
- Write a clear `officer_narrative` explaining what was detected and what to upload instead.
- Populate `refusal_reasons` with specific grounds.
- STOP. Do NOT attempt NOC matching or compliance auditing.

If all checks pass (or only SOFT_FAIL), proceed with the full analysis.

---

{task_noc_matching}

---

=== TASK 2 - DUTY EVIDENCE MAPPING (CRITICAL - This drives the decision) ===

You MUST map EVERY main duty from the selected NOC against the letter's content:

For `duties_match`, include ALL main duties from the NOC database - not just the ones that match.
For each duty, set `match_strength`:
  - "strong" - clear semantic alignment, specific evidence quoted from the letter
  - "partial" - related language but vague or incomplete
  - "weak" - only tangentially related
  - "missing" - no evidence in the letter at all

For `missing_critical_duties`, list every NOC duty that received "missing" or "weak" match_strength.
For `duty_coverage_percentage`, calculate: (count of "strong" + "partial") / (total NOC main duties) x 100

For `lead_statement_*`:
  - Quote the official lead statement from the NOC database
  - Quote the most relevant evidence from the letter
  - Explain how they align (or don't)

---

=== TASK 3 - IRCC COMPLIANCE AUDIT ===

Evaluate the employment letter against the official IRCC requirements for CEC reference letters.

MANDATORY ELEMENTS (populate `mandatory_requirements` booleans - set TRUE only if verifiably present):
1. OFFICIAL COMPANY LETTERHEAD
2. APPLICANT'S FULL NAME
3. COMPANY CONTACT INFORMATION - Address, telephone, email
4. JOB TITLE(S)
5. DATES OF EMPLOYMENT - Specific start and end dates
6. HOURS WORKED PER WEEK - Must prove full-time (30+) or state part-time hours
7. SALARY / COMPENSATION - Any format acceptable (hourly, weekly, monthly, annual)
8. SIGNATORY - Name, title, and signature of supervisor OR HR officer (both valid)

For `compliance.score`: (count of true mandatory_requirements / 8) x 100
For `compliance.missing_elements`: List mandatory elements completely absent.
For `compliance.warnings`: List elements present but insufficient/ambiguous.

=== DO NOT FLAG (IMPORTANT - These are NOT issues) ===
- Letter dated AFTER employment end date (normal for reference letters)
- Salary in any format (hourly, monthly, biweekly - all acceptable)
- HR signatory instead of supervisor (both valid per IRCC)
- Ongoing employment as "currently employed", "to date", "present" (all acceptable)
- Job title not matching NOC title exactly (duties override titles)
- Minor wording differences from official NOC descriptions
- Minor formatting inconsistencies

---

=== TASK 4 - RISK ASSESSMENT ===

`overall_risk`: low (strong case) / moderate (gaps but defensible) / high (likely refusal)
`pfl_likelihood`: low / medium / high - how likely is IRCC to issue a Procedural Fairness Letter?
`key_risks`: Every identified risk, ordered by severity (highest first). Each must be specific and actionable.

---

=== TASK 5 - DECISION (You MUST choose ONE. No neutrality allowed.) ===

ACCEPT if:
- >=75% duty coverage
- Clear, specific, verifiable duties
- All critical compliance elements present

PFL_RISK if:
- 50-75% duty coverage OR
- Ambiguity / vague duties OR
- Missing supporting clarity
- You MUST populate `refusal_reasons` with the specific grounds that would trigger a PFL.

REFUSE if:
- <50% duty coverage OR
- Missing key duties entirely OR
- Insufficient evidence to establish NOC alignment
- You MUST populate `refusal_reasons`.

---

=== TASK 6 - OFFICER NARRATIVE ===

Write `officer_narrative` in a realistic IRCC officer tone:
- Formal, concise, evidence-based, slightly skeptical
- 3-5 sentences referencing specific duties and gaps
- Example: "The duties described are insufficiently detailed to establish alignment with the claimed NOC. While the applicant references supervisory responsibilities, the letter lacks specificity regarding..."

---

=== TASK 7 - ACTION PLAN & SUGGESTED WORDING ===

`action_plan`: Priority-ordered, specific, actionable fixes. Most critical first. Each item must be directly tied to an identified risk or gap. Not generic advice.
`suggested_wording`: If duties are weak or missing, provide sample sentences the applicant can give to their employer to strengthen the letter's alignment with the NOC.

---

=== TASK 8 - LOCATION OF EXPERIENCE ===

Analyze the company address, letterhead, and geographic references:
- Canadian address/postal code/province -> "canada"
- Non-Canadian address -> "outside_canada"
- Cannot determine -> "unknown"

---

=== NOC 2021 DATABASE (Pre-filtered subset — ONLY use codes listed here) ===
{noc_reference}

---

=== FINAL RULE ===
If evidence is weak or missing, you MUST penalize heavily.
You MUST NOT "help" the applicant pass.
Your role is to PROTECT the integrity of the immigration system while giving fair, actionable feedback.

Output your analysis strictly conforming to the requested JSON schema. Be precise and evidence-based.
"""

def build_noc_finder_prompt(noc_reference: str, target_noc: str = None) -> str:
    """Builds the NOC Finder prompt v2 — fast, reliable, IRCC-consistent NOC suggestion."""

    today = datetime.date.today().strftime('%B %d, %Y')

    # --- Task 1 block varies for auto vs targeted evaluation ---
    if target_noc:
        task_1 = f"""
=== TASK 1 — TARGETED NOC EVALUATION ===

The user explicitly requested evaluation against NOC {target_noc}.
You MUST lock the primary match to NOC {target_noc}.

- Set `recommended_noc.code` strictly to "{target_noc}".
- Set `recommended_noc.title` to the exact title for {target_noc} from the database.

=== CONFIDENCE CALCULATION (MUST FOLLOW THIS EXACTLY) ===

To compute `recommended_noc.confidence`, you MUST:
1. Look up ALL main duties for NOC {target_noc} from the database.
2. For EACH main duty, classify the applicant's evidence:
   - "strong" = clear semantic alignment with specific evidence from the input
   - "partial" = related language but vague or incomplete
   - "missing" = no evidence in the input at all
3. Set `recommended_noc.duties_total` = total number of main duties for the NOC.
4. Set `recommended_noc.duties_matched` = count of "strong" + "partial".
5. Set `recommended_noc.confidence` = (duties_matched / duties_total) × 100.

These three numbers MUST be mathematically consistent. Do NOT estimate — count the duties.

- Classify `result_type` based on confidence: STRONG_MATCH (≥75%), MODERATE_MATCH (60-74%), NO_MATCH (<60%).
- `key_matches`: List the duties classified as "strong" (up to 5)
- `key_gaps`: List the duties classified as "missing" (up to 3)
- You may still list better fits in `alternatives`.
"""
    else:
        task_1 = """
=== TASK 1 — NOC MATCHING ===

=== PRE-COMPUTED DUTY COVERAGE (USE THIS DATA) ===

Each NOC entry in the database below includes machine-computed scores:
- `_duty_match_rank`: Position in duty-level matching (1 = strongest duty overlap)
- `_pre_computed_duty_coverage_pct`: % of this NOC's duties that semantically match the input
- `_pre_computed_duties_matched` / `_pre_computed_duties_total`: raw duty counts
- `_lead_statement_match`: How well the NOC's lead statement (employer type/industry) aligns
  with the user's described work (0.0-1.0). Higher = better industry match.
  CRITICAL: If two NOCs have similar duty scores but different lead_statement_match scores,
  ALWAYS prefer the one with the higher _lead_statement_match. This prevents misclassifying
  e.g. a collection agency worker as a bank teller.

These scores were computed by comparing EACH individual NOC duty against the user's text 
using embedding similarity. They are OBJECTIVE and should be your PRIMARY signal for 
NOC selection. Do NOT override these scores based on job title alone.

START your evaluation with the top-ranked candidates (rank 1-5) and verify their 
alignment by reading the duties yourself. Only select a lower-ranked candidate if the
top candidates' duties genuinely do not match the user's described work AND their
_lead_statement_match is low.

Steps:

1. Read the user's duties carefully.
2. Start with the NOCs ranked #1-3 by `_duty_match_rank` — these have the highest
   objective duty overlap with the input.
3. For your top candidate, VERIFY the machine scores by reading each duty yourself.
4. Compare duties against candidate NOCs in the database:
   - Lead statement alignment
   - Main duties overlap (SEMANTIC, not keyword matching)
3. For your top candidate NOC, compute DUTY COVERAGE using the method below.
4. Select the BEST NOC based on priority order:
   a) Duty coverage %
   b) Lead statement alignment
   c) Specificity of duties
5. Also evaluate up to 2 ALTERNATIVE NOCs with honest scores.

=== CRITICAL: SEMANTIC MATCHING, NOT KEYWORD MATCHING ===
Focus on the SEMANTIC MEANING of what the person actually does day-to-day.
Example: A "Fire Watch Team Member" who inspects work areas for hazards and recommends controls
is performing SAFETY SPECIALIST duties — not firefighter duties — even though the word "fire"
appears frequently. Always ask: "What is this person's core function?"

=== EMPLOYER INDUSTRY CROSS-CHECK (MANDATORY — DO THIS BEFORE FINALIZING) ===

BEFORE finalizing your top NOC, verify the EMPLOYER'S INDUSTRY against the NOC's lead statement:
1. Identify what the employer's business actually is (e.g., collection agency, bank, software company,
   restaurant, staffing firm). Use clues from the letter: company name, address, nature of work described.
2. Read the lead statement of your selected NOC — it lists the types of employers where this occupation
   is typically found (e.g., "employed by banks, trust companies, credit unions").
3. If the employer type DOES NOT appear in the lead statement, search for a NOC whose lead statement
   explicitly includes that employer type.

This is critical: two NOCs can have overlapping duty keywords but serve completely different industries.
A person handling "delinquent accounts" at a COLLECTION AGENCY is a collection clerk, not a bank teller.
Always prefer the NOC whose lead statement matches the actual employer.

=== TIE-BREAKING RULE ===
If two NOC codes score within 5 points, select the NOC whose LEAD STATEMENT most accurately
describes the person's primary role. List the runner-up as the first alternative.

=== CONFIDENCE CALCULATION (MUST FOLLOW THIS EXACTLY) ===

To compute `recommended_noc.confidence`, you MUST:
1. Look up ALL main duties for the selected NOC from the database.
2. For EACH main duty, classify the applicant's evidence:
   - "strong" = clear semantic alignment with specific evidence from the input
   - "partial" = related language but vague or incomplete
   - "missing" = no evidence in the input at all
3. Set `recommended_noc.duties_total` = total number of main duties for the NOC.
4. Set `recommended_noc.duties_matched` = count of "strong" + "partial".
5. Set `recommended_noc.confidence` = (duties_matched / duties_total) × 100.

These three numbers MUST be mathematically consistent. Do NOT estimate — count the duties.

=== MATCH CLASSIFICATION ===
Based on computed confidence:
- ≥75% → result_type = "STRONG_MATCH"
- 60–74% → result_type = "MODERATE_MATCH"
- <60% → result_type = "NO_MATCH"

If NO NOC reaches ~60%, still return the best candidate but set result_type to NO_MATCH
and explain in `why_this_noc` that alignment is weak.

=== OUTPUT MAPPING ===
- `recommended_noc.code`: Best-matching 5-digit NOC code
- `recommended_noc.title`: Exact title from the database
- `recommended_noc.confidence`: Computed duty coverage percentage (0-100)
- `recommended_noc.duties_total`: Total main duties for this NOC in the database
- `recommended_noc.duties_matched`: How many duties have strong or partial evidence
- `why_this_noc`: 1-2 sentence explanation of selection + any key concerns
- `key_matches`: Duties classified as "strong" (up to 5, short strings)
- `key_gaps`: Duties classified as "missing" (up to 3, short strings)
- `alternatives`: Up to 2 alternative NOCs with their computed confidence scores.
  **CRITICAL: You may ONLY suggest alternative NOCs whose codes appear in the provided database. NEVER invent or recall NOC codes from memory.**
"""

    return f"""You are a Canadian immigration NOC (National Occupational Classification) expert specializing in NOC 2021.

Your role is to IDENTIFY the most likely NOC for a user's work experience, while ensuring consistency with IRCC evaluation standards.

---

PRIMARY OBJECTIVE:
Recommend the MOST LIKELY NOC based on the user's duties, BUT ONLY if the duties appear to demonstrate meaningful alignment with that NOC.

You are NOT allowed to confidently recommend a NOC if alignment is weak. Be honest.

---

IMPORTANT CONTEXT:
- Today's date: {today}
- The applicant must demonstrate a MAJORITY (~60-70%) of the main duties of a NOC to be considered a strong match.
- Your output must be CONSISTENT with a future audit. Do NOT recommend a NOC that would likely fail a detailed audit.

---

=== INPUT VALIDATION (RUN FIRST) ===

IF DOCUMENT INPUT:

1. READABILITY:
   - If blank, corrupted, or no readable text → Set document_valid=false, rejection_reason="Document is blank or unreadable."

2. DOCUMENT TYPE:
   ACCEPT: Employment/reference/experience letters, Job offer letters WITH duties, Single document with multiple roles at SAME company
   ALLOW WITH LOWER CONFIDENCE: Resume/CV, Job descriptions (set applicable=true but note lower reliability in `notes`)
   REJECT: Payslips, T4s, ID documents, contracts without duties

3. MULTIPLE EMPLOYERS:
   - If multiple DIFFERENT employers → REJECT

4. DUTIES:
   - If fewer than 2 meaningful duties → REJECT

IF MANUAL INPUT:
   - If no clear job title OR fewer than 2 concrete duties → REJECT

VALIDATION RESULT:
- If PASSES: set document_valid=true, rejection_reason=""
- If FAILS: set document_valid=false, write clear rejection_reason. Do NOT populate noc_analysis.

{task_1}

=== TASK 2 — LOCATION OF EXPERIENCE ===
- If the input clearly mentions a Canadian address/province → "canada"
- If it clearly mentions a location outside Canada → "outside_canada"
- If unclear → "unknown"

=== STRICT RULES ===
- NEVER assume missing duties were performed
- NEVER rely on job title alone — always evaluate actual duties
- IGNORE any NOC codes written in the document by the employer. Employers frequently choose the
  wrong NOC. Your job is to independently determine the best match based on DUTIES, not to
  confirm the employer's claim.
- Your selection MUST come from the top 5 candidates by `_duty_match_rank` unless NONE of
  them have any relevant duties. If you pick a candidate ranked #6 or lower, you must explicitly
  explain why all top-5 candidates were rejected.
- ALWAYS prefer accuracy over completeness
- KEEP notes/explanations short and precise (1-2 sentences)
- ENSURE consistency with IRCC-style evaluation logic
- Score alternative NOCs honestly as if doing a full dedicated evaluation against each

Your goal is to provide a FAST, RELIABLE, and TRUSTWORTHY NOC suggestion — not a full audit.

=== NOC 2021 DATABASE (Pre-filtered subset — ONLY use codes listed here) ===
{noc_reference}

Output your analysis strictly conforming to the requested JSON schema.
"""


def ocr_from_page_images(page_images: list[tuple[bytes, str]], max_pages: int = 2) -> str:
    """Use GPT-4o-mini vision to OCR text from scanned PDF page images.
    
    This is a lightweight fallback for scanned PDFs where pdfminer returns no text.
    We only OCR the first few pages to keep costs low — just enough for the RAG search.
    """
    if not openai_client or not page_images:
        return ""
    
    content = [{"type": "text", "text": "Extract ALL text visible in this document image. Return ONLY the raw text, no commentary."}]
    for img_bytes, mime_type in page_images[:max_pages]:
        b64 = base64.b64encode(img_bytes).decode('utf-8')
        content.append({
            "type": "image_url",
            "image_url": {"url": f"data:{mime_type};base64,{b64}", "detail": "auto"}
        })
    
    try:
        resp = openai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": content}],
            temperature=0.0,
            max_tokens=2000
        )
        extracted = resp.choices[0].message.content or ""
        print(f"[OCR] Extracted {len(extracted)} chars from {min(len(page_images), max_pages)} page image(s)")
        return extracted
    except Exception as e:
        print(f"[OCR] Failed: {e}")
        return ""


def extract_document_content(doc_bytes: bytes, ext: str, is_image: bool) -> tuple[str, list]:
    """Extract text and page images from any supported document format.
    
    Returns (user_content, page_images) tuple ready for RAG search and AI processing.
    Handles scanned PDF OCR fallback automatically.
    """
    page_images = []
    user_content = ""
    
    if is_image:
        mime_type = IMAGE_MIME_TYPES.get(ext, 'image/jpeg')
        page_images.append((doc_bytes, mime_type))
        # OCR the image so RAG gets real text
        user_content = ocr_from_page_images(page_images)
        if not user_content.strip():
            user_content = "The user uploaded an image of their employment letter. Extract the job title and duties."
    elif ext == '.pdf':
        page_images = pdf_pages_to_images(doc_bytes)
        extracted_text = extract_text_from_pdf(doc_bytes)
        # Scanned PDF fallback: OCR page images if pdfminer returned nothing
        if len(extracted_text.strip()) < 50 and page_images:
            print("[Scanned PDF] Text extraction returned <50 chars, running OCR...")
            extracted_text = ocr_from_page_images(page_images)
        user_content = f"=== EXTRACTED PDF TEXT ===\n{extracted_text}"
    elif ext in ('.docx', '.doc'):
        user_content = f"=== EXTRACTED WORD TEXT ===\n{extract_text_from_docx(doc_bytes)}"
    else:
        user_content = f"=== EXTRACTED TEXT ===\n{doc_bytes.decode('utf-8', errors='replace')}"
    
    return user_content, page_images


def semantic_search_nocs(user_text: str, top_k: int = 40) -> dict:
    """Embed the user's text and find the top_k closest NOC codes.
    
    Uses pre-computed numpy embedding matrix for fast vectorized similarity.
    Default top_k=40 provides good coverage for edge cases where the correct NOC
    ranks just outside a narrower window (e.g., collection clerks at rank #33).
    """
    if not openai_client or _NOC_EMB_MATRIX is None:
        raise ValueError("OpenAI client or NOC embeddings not initialized.")
    
    # 1. Pre-process: extract DUTY-FOCUSED text for embedding.
    #    Employment letters contain boilerplate (company name, addresses, employee bios,
    #    "To Whom It May Concern", signatory blocks) that dilutes the embedding signal.
    #    E.g., "Amandeep Agro Chemicals" repeated in letterhead pushes the embedding toward
    #    agricultural NOCs even though the actual duties are retail supervision.
    #    Strategy: extract the duty section if possible, else strip boilerplate aggressively.
    import re
    text_to_embed = user_text
    # Remove the "=== EXTRACTED ... ===" header
    text_to_embed = re.sub(r'^===.*?===\s*', '', text_to_embed)
    
    # --- Attempt 1: Extract just the duty section from employment letters ---
    # Look for common duty section markers in employment letters
    duty_section = None
    duty_markers = [
        # Allow up to 40 chars between "duties" and "included/are/were" to handle
        # patterns like "duties as a Supervisor included the following"
        r'(?i)(?:duties|responsibilities|job duties|main duties|key duties|'
        r'principal duties|role and responsibilities|scope of work|'
        r'duties and responsibilities).{0,40}?(?:included?|are|were|as follows|'
        r'but (?:were|are) not limited to|:)',
    ]
    for marker in duty_markers:
        match = re.search(marker, text_to_embed)
        if match:
            # Extract from the marker to the end, then trim at common ending markers
            section = text_to_embed[match.start():]
            # Trim at signatory/closing markers
            end_match = re.search(
                r'(?i)(?:^|\n)\s*(?:if you (?:require|need|have)|sincerely|regards|'
                r'yours truly|please (?:do not hesitate|feel free)|for verification|'
                r'should you (?:require|need)|we wish|authorized signatory|'
                r'managing director|human resource|HR manager)',
                section
            )
            if end_match:
                section = section[:end_match.start()]
            duty_section = section.strip()
            break
    
    # Also extract job title line if present (important context for embedding)
    title_line = ""
    title_match = re.search(
        r'(?i)(?:job title|position|capacity|role|designation)\s*(?:of|as|:|-|–)?\s*(.+)',
        text_to_embed
    )
    if title_match:
        title_line = title_match.group(0).strip()[:100]
    
    if duty_section and len(duty_section) > 50:
        # Use the focused duty section with the job title
        text_to_embed = f"{title_line}\n{duty_section}" if title_line else duty_section
        print(f"[RAG Preprocess] Extracted duty section: {len(duty_section)} chars "
              f"(from {len(user_text)} total)")
    else:
        # --- Attempt 2: Aggressive boilerplate stripping ---
        # Remove phone/email/fax lines
        text_to_embed = re.sub(r'(?m)^.*?(PHONE|TOLL FREE|FAX|www\.|http|@).*$', '', text_to_embed)
        text_to_embed = re.sub(r'(?m)^.*?\d{3}[- ]\d{3}[- ]\d{4}.*$', '', text_to_embed)
        # Remove postal/zip codes (Canadian and Indian PIN codes)
        text_to_embed = re.sub(r'(?m)^.*?[A-Z]\d[A-Z]\s*\d[A-Z]\d.*$', '', text_to_embed)
        text_to_embed = re.sub(r'(?m)^.*?\d{6}.*$', '', text_to_embed)  # 6-digit PIN codes
        # Remove GSTIN/tax ID lines
        text_to_embed = re.sub(r'(?m)^.*?(?:GSTIN|GST|PAN|TIN|EIN)\s*[:#]?\s*\w+.*$', '', text_to_embed)
        # Remove common letter boilerplate
        text_to_embed = re.sub(r'(?im)^.*(?:to whom it may concern|this is to certify|'
                               r'ref\.?\s*no|dated\s*:|sincerely|regards|yours truly|'
                               r'managing director|authorized signatory|'
                               r'if you require any|please feel free|'
                               r'for verification purposes).*$', '', text_to_embed)
        # Remove date lines (DD/MM/YYYY, MM.DD.YYYY, etc.)
        text_to_embed = re.sub(r'(?m)^.*?\d{1,2}[./\-]\d{1,2}[./\-]\d{2,4}.*$', '', text_to_embed)
        # Remove employee bio fluff (dedication, sincerity, etc.)
        text_to_embed = re.sub(r'(?i)(?:with )?dedication,?\s*determination\s*and\s*sincerity[^.]*\.?', '', text_to_embed)
        text_to_embed = re.sub(r'(?i)(?:we found|she was|he was) (?:her|him|them) (?:active|professional|hard)[^.]*\.?', '', text_to_embed)
        text_to_embed = re.sub(r'(?i)(?:we are gratified|we wish)[^.]*\.?', '', text_to_embed)
        print(f"[RAG Preprocess] No duty section found, used aggressive stripping")
    
    text_to_embed = re.sub(r'\s+', ' ', text_to_embed).strip()
    
    text_to_embed = text_to_embed[:8000]  # Safe limit for embedding model
    if not text_to_embed.strip():
        text_to_embed = "General professional duties"
        
    response = openai_client.embeddings.create(
        model="text-embedding-3-small",
        input=[text_to_embed]
    )
    user_vector = np.array(response.data[0].embedding)
    
    # 2. Vectorized cosine similarity against pre-computed matrix
    similarities = _NOC_EMB_MATRIX @ user_vector  # (516,) dot products in one operation
    
    # 3. Get top_k indices using argpartition (faster than full sort for large arrays)
    top_indices = np.argpartition(similarities, -top_k)[-top_k:]
    top_indices = top_indices[np.argsort(similarities[top_indices])[::-1]]  # Sort the top_k
    
    # 4. Build a subset of NOC_INDEX using the pre-ordered key list
    top_nocs_dict = {}
    for idx in top_indices:
        idx_key = _NOC_EMB_KEYS[idx]
        entry = NOC_INDEX.get(idx_key)
        if entry and "code" in entry:
            top_nocs_dict[entry["code"]] = entry
    
    # 5. Duty-level reranking: re-score candidates using per-duty embedding similarity
    #    This catches cases where the correct NOC has high individual-duty overlap
    #    but lower whole-document similarity (e.g., "Collection clerks" for debt collectors)
    if _DUTY_EMB_MATRIX is not None and len(top_nocs_dict) > 0:
        top_nocs_dict = _duty_level_rerank(user_vector, top_nocs_dict, user_text=user_text)
    
    return top_nocs_dict


def _duty_level_rerank(user_vector: np.ndarray, candidates: dict, final_k: int = 10, user_text: str = "") -> dict:
    """Re-score and rerank NOC candidates using per-duty embedding similarity,
    blended with lead statement similarity for industry/employer-type alignment.
    
    For each candidate NOC:
    1. Compute cosine similarity between user text and EACH individual duty
    2. duty_score = avg_similarity × duty_coverage
    3. Compute lead statement similarity (captures employer type, e.g., 'employed by
       collection agencies') using pre-computed lead statement embeddings
    4. composite = duty_score × 0.6 + lead_sim × 0.4
    5. Rerank by composite score
    6. Inject duty coverage stats into each entry so the AI sees them in the prompt
    
    Cost: ZERO additional API calls. Pure numpy on pre-computed vectors.
    """
    COVERAGE_THRESHOLD = 0.30  # Minimum similarity for a duty to count as "matched"
    DUTY_WEIGHT = 0.6          # Weight for duty-level score
    LEAD_WEIGHT = 0.4          # Weight for lead statement similarity
    
    # Compute similarity between user vector and ALL duty vectors at once
    all_duty_sims = _DUTY_EMB_MATRIX @ user_vector  # (N_duties,) — fast vectorized
    
    # Compute lead statement similarities using a FOCUSED duty+employer embedding.
    # The full user_vector includes boilerplate that dilutes the lead statement signal.
    # We extract duty-related lines AND employer context, then embed them together.
    # The employer context is critical: "Gatestone & Co." is a collection agency, not a bank,
    # and this differentiates 14202 (Collection clerks) from 64400 (Bank tellers).
    # Cost: 1 extra embedding call (~$0.0001).
    lead_sims = {}
    if _LEAD_EMB_MATRIX is not None and openai_client:
        import re
        lines = user_text.split('\n') if user_text else []
        
        # Extract employer context: company name, RE: line, first few content lines
        employer_lines = []
        duty_lines = []
        for line in lines:
            stripped = line.strip()
            if not stripped or len(stripped) < 5:
                continue
            # Employer context: company names, RE: lines, "attention" lines
            if re.match(r'(?i)(RE:|ATTN|attention|dear|to whom|letter of employment)', stripped):
                employer_lines.append(stripped)
            elif re.match(r'^[A-Z][A-Z\s&.,]+$', stripped) and len(stripped) < 60:
                # ALL-CAPS lines are often company names
                employer_lines.append(stripped)
            # Duty lines
            duty_keywords = re.compile(
                r'(?i)(responsible|duties|functions|tasks|include|perform|manage|develop|'
                r'maintain|coordinate|review|prepare|process|conduct|provide|assist|'
                r'monitor|analyze|ensure|create|implement|administer|negotiate|'
                r'contact|collect|resolve|recommend|advise|report|overdue|account|'
                r'client|customer|payment|delinquent|credit|invoice|repayment|'
                r'has been employed|position of|role of|job title)',
            )
            if len(stripped) > 15 and duty_keywords.search(stripped):
                duty_lines.append(stripped)
        
        # Combine: employer context FIRST (gives industry signal), then duties
        context_text = ' '.join(employer_lines[:5]) + ' | ' + ' '.join(duty_lines)
        context_text = context_text[:4000]
        
        if context_text.strip():
            try:
                ctx_resp = openai_client.embeddings.create(
                    model="text-embedding-3-small",
                    input=[context_text]
                )
                ctx_vector = np.array(ctx_resp.data[0].embedding)
                all_lead_sims = _LEAD_EMB_MATRIX @ ctx_vector
                for code in candidates:
                    idx = _LEAD_CODE_TO_IDX.get(code)
                    if idx is not None:
                        lead_sims[code] = float(all_lead_sims[idx])
            except Exception as e:
                print(f"[Lead Matching] Context embedding failed: {e}. Falling back.")
                all_lead_sims = _LEAD_EMB_MATRIX @ user_vector
                for code in candidates:
                    idx = _LEAD_CODE_TO_IDX.get(code)
                    if idx is not None:
                        lead_sims[code] = float(all_lead_sims[idx])
        else:
            all_lead_sims = _LEAD_EMB_MATRIX @ user_vector
            for code in candidates:
                idx = _LEAD_CODE_TO_IDX.get(code)
                if idx is not None:
                    lead_sims[code] = float(all_lead_sims[idx])
    
    scored = []
    for code, entry in candidates.items():
        if code not in _DUTY_RANGES:
            scored.append((0.0, 0, 0, code, entry))
            continue
        
        start, end = _DUTY_RANGES[code]
        duty_sims = all_duty_sims[start:end]
        total = len(duty_sims)
        
        if total == 0:
            scored.append((0.0, 0, 0, code, entry))
            continue
        
        matched = int(np.sum(duty_sims >= COVERAGE_THRESHOLD))
        coverage = matched / total
        avg_sim = float(np.mean(duty_sims))
        
        # Duty-level score: rewards NOCs where MOST duties match, weighted by similarity
        duty_score = avg_sim * coverage
        
        # Lead statement similarity (captures employer/industry alignment)
        lead_sim = lead_sims.get(code, 0.0)
        composite = (duty_score * DUTY_WEIGHT) + (lead_sim * LEAD_WEIGHT)
        
        scored.append((composite, matched, total, code, entry))
    
    # Sort by composite score descending
    scored.sort(key=lambda x: x[0], reverse=True)
    
    # Log the reranking for debugging
    top5 = [(code, f"{score:.4f}") for score, _, _, code, _ in scored[:5]]
    print(f"[Duty Rerank] Top 5: {top5}")
    
    # Inject duty coverage AND lead statement alignment into each entry so the AI sees them
    reranked = {}
    for rank, (score, matched, total, code, entry) in enumerate(scored[:final_k]):
        enriched = dict(entry)  # shallow copy to avoid mutating the original
        coverage_pct = round((matched / total) * 100) if total > 0 else 0
        enriched["_duty_match_rank"] = rank + 1
        enriched["_pre_computed_duty_coverage_pct"] = coverage_pct
        enriched["_pre_computed_duties_matched"] = matched
        enriched["_pre_computed_duties_total"] = total
        # Inject lead statement match score so the AI can see employer/industry alignment
        lead_sim = lead_sims.get(code, 0.0)
        enriched["_lead_statement_match"] = round(lead_sim, 4)
        reranked[code] = enriched
    
    # Store top-10 for post-processing enforcement in find_noc_with_openai
    # (Named _LAST_DUTY_TOP5 for legacy reasons, but stores top-10 for a wider safety net)
    global _LAST_DUTY_TOP5
    _LAST_DUTY_TOP5 = {code: reranked[code] for code in list(reranked.keys())[:10]}
    
    return reranked


def _call_openai_structured(system_prompt: str, user_content: str, page_images: list[tuple[bytes, str]], response_format, label: str = "AI", seed: int = 42) -> dict:
    """Unified OpenAI structured output call. Used by both NOC Finder and Auditor.
    
    Args:
        system_prompt: The system instruction prompt.
        user_content: Extracted text content from the document.
        page_images: Optional list of (image_bytes, mime_type) tuples for vision.
        response_format: Pydantic model class for structured output.
        label: Log label for debugging.
        seed: Seed for reproducibility (default 42). Use different seeds for retries.
    """
    if not openai_client:
        raise ValueError("OPENAI_API_KEY is not set. Please configure it to use GPT-4o-mini.")
    
    messages = [
        {"role": "system", "content": system_prompt}
    ]
    
    user_message_content = []
    if user_content:
        user_message_content.append({"type": "text", "text": user_content})
        
    if page_images:
        for img_bytes, mime_type in page_images:
            base64_img = base64.b64encode(img_bytes).decode('utf-8')
            user_message_content.append({
                "type": "image_url",
                "image_url": {
                    "url": f"data:{mime_type};base64,{base64_img}",
                    "detail": "auto"
                }
            })
            
    messages.append({"role": "user", "content": user_message_content})
    
    print(f"Calling OpenAI gpt-4o-mini for {label}...")
    completion = openai_client.beta.chat.completions.parse(
        model="gpt-4o-mini",
        messages=messages,
        response_format=response_format,
        temperature=0.0,
        seed=seed
    )
    
    result = json.loads(completion.choices[0].message.content)
    return _sanitize_noc_response(result)


def find_noc_with_openai(system_prompt: str, user_content: str, page_images: list[tuple[bytes, str]] = None) -> dict:
    """NOC Finder: returns NOCFinderResponseSchema.
    
    Strips any employer-stated NOC codes from user_content before sending to the AI
    to prevent the model from deferring to often-incorrect employer classifications.
    """
    import re
    # Strip employer NOC code claims — these bias the model toward wrong codes
    # Matches patterns like: "NOC 42201", "(NOC 42201)", "NOC: 42201", "NOC Code: 42201"
    # Also matches the title pattern: "NOC 42201 - Social and community service workers"
    cleaned_content = re.sub(
        r'(?i)\(?\s*NOC\s*[:#]?\s*\d{4,5}\s*[-–—]?\s*[^\n)]{0,80}\)?',
        '[EMPLOYER NOC REFERENCE REDACTED]',
        user_content
    )
    if cleaned_content != user_content:
        print(f"[NOC Finder] Stripped employer NOC references from input to prevent bias")
    
    from models import NOCFinderResponseSchema
    # Don't send page_images to the NOC Finder — the extracted text is sufficient
    # for duty matching, and images introduce visual bias (employer NOC claims,
    # company logos, etc. that shouldn't influence NOC selection).
    # Images ARE still sent for the auditor where visual verification matters.
    result = _call_openai_structured(system_prompt, cleaned_content, None, NOCFinderResponseSchema, "NOC Finder")
    
    # POST-PROCESSING: Best-of-3 consensus for determinism.
    # gpt-4o-mini is non-deterministic even with temperature=0 and seed=42.
    # Run 3 calls with different seeds, collect all unique NOC picks, and select
    # the one with the best duty+lead rank. This is deterministic because the
    # composite rankings are pre-computed and stable.
    # Cost: 2 extra gpt-4o-mini calls (~$0.002).
    result2 = _call_openai_structured(system_prompt, cleaned_content, None, NOCFinderResponseSchema, "NOC Finder (vote 2)", seed=123)
    result3 = _call_openai_structured(system_prompt, cleaned_content, None, NOCFinderResponseSchema, "NOC Finder (vote 3)", seed=7)
    
    votes = [
        (result, result.get("recommended_noc", {}).get("code")),
        (result2, result2.get("recommended_noc", {}).get("code")),
        (result3, result3.get("recommended_noc", {}).get("code")),
    ]
    codes = [c for _, c in votes if c]
    unique_codes = list(set(codes))
    
    if len(unique_codes) > 1 and _LAST_DUTY_TOP5:
        # Votes disagree — pick the result whose NOC has the best duty+lead rank
        def get_rank(code):
            entry = _LAST_DUTY_TOP5.get(code)
            return entry.get("_duty_match_rank", 999) if entry else 999
        
        best_code = min(unique_codes, key=get_rank)
        best_result = next(r for r, c in votes if c == best_code)
        ranks = {c: get_rank(c) for c in unique_codes}
        print(f"[NOC Finder] CONSENSUS: votes={codes}, ranks={ranks} → picking {best_code} (rank #{get_rank(best_code)})")
        result = best_result
    else:
        print(f"[NOC Finder] CONSENSUS: all votes agree on {unique_codes[0] if unique_codes else '?'}")
    
    return result


def auto_detect_noc(user_content: str, page_images: list[tuple[bytes, str]] = None) -> str | None:
    """Run the NOC Finder pipeline to auto-detect the best NOC code for a document.
    
    Used as a pre-processing step by the auditor when no target_noc is provided.
    This guarantees the auditor uses the EXACT same NOC detection logic as the
    NOC Finder — same duty reranking, same employer NOC stripping, same image
    exclusion, same prompt constraints.
    
    Returns the detected NOC code string, or None if detection fails.
    Cost: ~$0.001-0.003 (one gpt-4o-mini call, text-only).
    """
    try:
        top_nocs = semantic_search_nocs(user_content)
        noc_reference = json.dumps(top_nocs, ensure_ascii=False)
        system_prompt = build_noc_finder_prompt(noc_reference)
        
        result = find_noc_with_openai(
            system_prompt=system_prompt,
            user_content=f"=== USER INPUT ===\n{user_content}",
            page_images=page_images  # find_noc_with_openai internally strips images
        )
        
        detected_code = result.get("recommended_noc", {}).get("code")
        detected_title = result.get("recommended_noc", {}).get("title", "?")
        confidence = result.get("recommended_noc", {}).get("confidence", 0)
        
        if detected_code:
            print(f"[Auto-Detect NOC] Detected: {detected_code} ({detected_title}) — {confidence}% confidence")
            return detected_code
        else:
            print("[Auto-Detect NOC] Failed to detect NOC code from result")
            return None
    except Exception as e:
        print(f"[Auto-Detect NOC] Error: {e}")
        return None


def audit_document_with_openai(system_prompt: str, user_content: str, page_images: list[tuple[bytes, str]] = None, auto_detected_noc: str = None) -> dict:
    """Employment Auditor: returns AnalysisResponse.
    
    Args:
        auto_detected_noc: If set, this NOC was auto-detected (not user-specified).
            Employer NOC references will be stripped from the text to prevent
            the model from overriding the detected target.
    """
    import re
    if auto_detected_noc:
        # Strip employer NOC claims to prevent the model from ignoring the auto-detected target.
        # Example: employer writes "NOC 42201" but auto-detect found 41321 is the correct match.
        cleaned_content = re.sub(
            r'(?i)\(?\s*NOC\s*[:#]?\s*\d{4,5}\s*[-–—]?\s*[^\n)]{0,80}\)?',
            '[EMPLOYER NOC REFERENCE REDACTED]',
            user_content
        )
        if cleaned_content != user_content:
            print(f"[Auditor] Stripped employer NOC references (auto-detected target: {auto_detected_noc})")
        user_content = cleaned_content
    
    from models import AnalysisResponse
    result = _call_openai_structured(system_prompt, user_content, page_images, AnalysisResponse, "Auditor")
    
    # Post-process: If we auto-detected a target NOC but the model overrode it
    # (e.g., because the employer's NOC claim is visible in page images),
    # force-correct the detected_code to match the auto-detected target.
    if auto_detected_noc and result.get("noc_analysis", {}).get("detected_code") != auto_detected_noc:
        noc_analysis = result.get("noc_analysis", {})
        model_code = noc_analysis.get("detected_code")
        model_title = noc_analysis.get("detected_title")
        print(f"[Auditor] Correcting detected_code: model said {model_code}, auto-detect said {auto_detected_noc}")
        
        # Get the correct title from the index
        target_entry = next((v for v in NOC_INDEX.values() if v.get("code") == auto_detected_noc), None)
        target_title = target_entry.get("title", "") if target_entry else auto_detected_noc
        
        noc_analysis["detected_code"] = auto_detected_noc
        noc_analysis["detected_title"] = target_title
        
        # Preserve the model's original pick as an alternative so nothing is lost
        alts = noc_analysis.get("alternative_nocs", [])
        if model_code and not any(a.get("noc_code") == model_code for a in alts):
            alts.append({
                "noc_code": model_code,
                "noc_title": model_title or model_code,
                "fit_assessment": "moderate",
                "reason": f"Employer's stated NOC. Overridden by duty-level analysis which found {auto_detected_noc} ({target_title}) as a better match."
            })
            noc_analysis["alternative_nocs"] = alts
        
        result["noc_analysis"] = noc_analysis
    
    return result


def _resolve_noc_code(code: str, model_title: str) -> str:
    """Resolve the correct NOC code, trusting the model's title over its code when they disagree.
    
    The model is much better at remembering titles than 5-digit codes.
    When code and title disagree, the title is almost always the true intent.
    Returns the resolved code (may be the original or a corrected one).
    """
    if code in NOC_LOOKUP:
        db_title = NOC_LOOKUP[code]
        # Code exists in DB — check if the model's title matches
        if model_title and model_title.lower().strip() != db_title.lower().strip():
            # Mismatch: model said one title but code points to a different one
            resolved = NOC_TITLE_TO_CODE.get(model_title.lower().strip())
            if resolved and resolved != code:
                print(f"[Sanitizer] CODE-TITLE MISMATCH FIX: model said code={code} "
                      f"({db_title}) but title='{model_title}'. "
                      f"Resolved to {resolved} via reverse title lookup.")
                return resolved
        return code
    
    # Code is completely fake — try to resolve from title alone
    if model_title:
        resolved = NOC_TITLE_TO_CODE.get(model_title.lower().strip())
        if resolved:
            print(f"[Sanitizer] FAKE CODE FIX: {code} not in DB, "
                  f"resolved to {resolved} from title '{model_title}'")
            return resolved
    
    # Neither code nor title could be resolved
    print(f"[Sanitizer] UNRESOLVABLE: code={code}, title='{model_title}' — not in DB")
    return code


def _sanitize_noc_response(result: dict) -> dict:
    """Post-process AI response to fix hallucinated NOC codes/titles.
    
    gpt-4o-mini sometimes invents NOC codes or pairs real codes with wrong titles.
    This function:
    - Resolves code/title mismatches by trusting the model's title (reverse lookup)
    - Removes alternative NOCs with completely fake codes
    - Validates confidence against reported duty counts
    """
    # --- Fix NOC Finder response format ---
    rec = result.get("recommended_noc")
    if rec and isinstance(rec, dict):
        model_title = rec.get("title", "")
        code = rec.get("code", "")
        
        # Resolve code using title-to-code reverse lookup
        resolved_code = _resolve_noc_code(code, model_title)
        rec["code"] = resolved_code
        
        # Always set title from DB (source of truth)
        if resolved_code in NOC_LOOKUP:
            rec["title"] = NOC_LOOKUP[resolved_code]
        
        # Validate confidence against duty counts (Fix 2b)
        total = rec.get("duties_total", 0)
        matched = rec.get("duties_matched", 0)
        if total > 0:
            calculated = round((matched / total) * 100)
            stated = rec.get("confidence", 0)
            if abs(calculated - stated) > 10:
                print(f"[Sanitizer] CONFIDENCE CORRECTION: stated={stated}%, "
                      f"calculated={calculated}% ({matched}/{total} duties). Using calculated.")
                rec["confidence"] = calculated
    
    if "alternatives" in result and isinstance(result["alternatives"], list):
        cleaned = []
        for alt in result["alternatives"]:
            model_title = alt.get("title", "")
            code = alt.get("code", "")
            resolved_code = _resolve_noc_code(code, model_title)
            if resolved_code in NOC_LOOKUP:
                alt["code"] = resolved_code
                alt["title"] = NOC_LOOKUP[resolved_code]
                cleaned.append(alt)
            else:
                print(f"[Sanitizer] Removed hallucinated alternative NOC: {code} - {model_title}")
        result["alternatives"] = cleaned
    
    # --- Fix Auditor response format ---
    noc_analysis = result.get("noc_analysis")
    if noc_analysis and isinstance(noc_analysis, dict):
        model_title = noc_analysis.get("detected_title", "")
        code = noc_analysis.get("detected_code", "")
        resolved_code = _resolve_noc_code(code, model_title)
        if resolved_code in NOC_LOOKUP:
            noc_analysis["detected_code"] = resolved_code
            noc_analysis["detected_title"] = NOC_LOOKUP[resolved_code]
        
        if "alternative_nocs" in noc_analysis and isinstance(noc_analysis["alternative_nocs"], list):
            cleaned = []
            for alt in noc_analysis["alternative_nocs"]:
                model_title = alt.get("noc_title", "")
                code = alt.get("noc_code", "")
                resolved_code = _resolve_noc_code(code, model_title)
                if resolved_code in NOC_LOOKUP:
                    alt["noc_code"] = resolved_code
                    alt["noc_title"] = NOC_LOOKUP[resolved_code]
                    cleaned.append(alt)
                else:
                    print(f"[Sanitizer] Removed hallucinated alternative NOC: {code} - {model_title}")
            noc_analysis["alternative_nocs"] = cleaned
    
    return result

# ── Letter Builder Functions ──

def get_noc_details(noc_code: str) -> dict | None:
    """Look up a single NOC code from the loaded index. Returns the entry or None."""
    for entry in NOC_INDEX.values():
        if entry.get("code") == noc_code:
            return entry
    return None


def build_duty_analysis_prompt(noc_entry: dict, user_duty: str) -> str:
    """Builds a lightweight prompt to evaluate ONE duty against ONE NOC's official duties."""
    
    noc_code = noc_entry.get("code", "")
    noc_title = noc_entry.get("title", "")
    lead_statement = noc_entry.get("lead_statement", "")
    duties = noc_entry.get("duties", [])
    duties_text = "\n".join(f"  {i+1}. {d}" for i, d in enumerate(duties))
    
    return f"""You are an IRCC employment letter compliance assistant specializing in NOC 2021.

Your ONLY task: evaluate whether a user-written duty statement aligns with the official duties of NOC {noc_code} ({noc_title}).

=== OFFICIAL NOC {noc_code} INFORMATION ===
Title: {noc_title}
Lead Statement: {lead_statement}
Main Duties:
{duties_text}

=== USER'S DUTY STATEMENT ===
"{user_duty}"

=== YOUR TASK ===

1. Determine which official duty (if any) this user statement most closely aligns with.
2. Rate the alignment: strong (clear match), partial (related but vague), weak (tangential), none (no alignment).
3. If the duty is vague or incomplete, provide up to 3 specific coaching questions to help them make it stronger and more IRCC-compliant.
4. Determine if this duty is "IRCC ready" — meaning it is specific enough, action-oriented, and clearly aligns with the NOC to be included in a formal employment letter.

=== STRICT RULES ===
- Do NOT rewrite the duty for them
- Do NOT invent information the user didn't provide
- Do NOT suggest duties the user didn't mention performing
- If the duty has NO alignment with any official NOC duty, set alignment to "none" and explain clearly
- Coaching questions should ask about SPECIFICS: tools, frequency, scope, outcomes, who they report to
- Be encouraging but honest

Output your analysis strictly conforming to the requested JSON schema.
"""


def analyze_single_duty(user_duty: str, noc_code: str) -> dict:
    """Analyze a single user-written duty against a specific NOC code using AI."""
    from letter_builder_models import DutyAnalysisResponse
    
    noc_entry = get_noc_details(noc_code)
    if not noc_entry:
        raise ValueError(f"NOC code {noc_code} not found in index")
    
    prompt = build_duty_analysis_prompt(noc_entry, user_duty)
    
    print(f"Analyzing duty against NOC {noc_code}: '{user_duty[:60]}...'")
    
    response = gemini_client.models.generate_content(
        model='gemini-2.5-flash',
        contents=prompt,
        config=types.GenerateContentConfig(
            response_mime_type="application/json",
            response_schema=DutyAnalysisResponse,
            temperature=0.0,
        ),
    )
    
    return json.loads(response.text)


def assemble_letter_text(employment_details: dict, noc_code: str, noc_title: str, approved_duties: list) -> dict:
    """Assemble a complete IRCC-compliant employment letter from user-approved data.
    
    This is pure Python template rendering — NO AI call needed.
    Every word in the output was provided or approved by the user.
    """
    import datetime as _dt
    
    d = employment_details
    today = _dt.date.today().strftime("%B %d, %Y")
    
    # Determine employment status phrasing
    is_ongoing = d.get("end_date", "").lower() in ("ongoing", "present", "current", "to date", "")
    if is_ongoing:
        period_phrase = f"since {d['start_date']}"
        status_phrase = f"is currently employed"
        end_display = "Present"
    else:
        period_phrase = f"from {d['start_date']} to {d['end_date']}"
        status_phrase = f"was employed"
        end_display = d['end_date']
    
    # Salary formatting
    salary_str = f"{d['salary_amount']} {d['salary_currency']} {d['salary_period']}"
    
    # Build duties bullet list
    duties_bullets = "\n".join(f"    • {duty['text']}" for duty in approved_duties)
    
    # Assemble intro paragraph
    intro = (
        f"This letter is to confirm that {d['applicant_name']} "
        f"{status_phrase} by {d['company_name']} as a {d['job_title']} "
        f"{period_phrase}."
    )
    
    # Employment details paragraph
    emp_details = (
        f"{d['applicant_name']} works {d['hours_per_week']} hours per week on a "
        f"{d['employment_type']} basis. "
        f"{'Their' if is_ongoing else 'Their'} compensation is {salary_str}. "
        f"The position is based in {d['work_city']}, {d['work_country']}."
    )
    
    # Duties section
    duties_section = (
        f"During {'their' if is_ongoing else 'their'} employment, "
        f"{d['applicant_name']}'s main duties and responsibilities include:\n\n"
        f"{duties_bullets}"
    )
    
    # Closing
    closing = (
        f"Should you require any further information regarding "
        f"{d['applicant_name']}'s employment, please do not hesitate "
        f"to contact the undersigned at {d['supervisor_contact']}."
    )
    
    # Supervisor block
    supervisor_block = (
        f"Sincerely,\n\n"
        f"{d['supervisor_name']}\n"
        f"{d['supervisor_title']}\n"
        f"{d['company_name']}\n"
        f"{d.get('company_address', '')}"
    )
    
    # Full letter
    full_text = (
        f"[COMPANY LETTERHEAD]\n\n"
        f"{today}\n\n"
        f"To Whom It May Concern,\n\n"
        f"{intro}\n\n"
        f"{emp_details}\n\n"
        f"{duties_section}\n\n"
        f"{closing}\n\n"
        f"{supervisor_block}\n\n"
        f"[SIGNATURE]"
    )
    
    # Warnings
    warnings = []
    if len(approved_duties) < 4:
        warnings.append(f"Only {len(approved_duties)} duties provided. IRCC typically expects at least 4 meaningful duties.")
    
    strong_count = sum(1 for d in approved_duties if d.get("alignment") == "strong")
    if strong_count < 2:
        warnings.append("Fewer than 2 duties have strong NOC alignment. Consider strengthening your duty statements.")
    
    return {
        "status": "APPROVED" if len(approved_duties) >= 4 else "INCOMPLETE",
        "noc_code": noc_code,
        "noc_title": noc_title,
        "letter_sections": {
            "header_placeholder": "[COMPANY LETTERHEAD]",
            "date": today,
            "addressee": "To Whom It May Concern,",
            "intro_paragraph": intro,
            "employment_details_paragraph": emp_details,
            "duties_section": duties_section,
            "closing_paragraph": closing,
            "supervisor_block": supervisor_block,
            "signature_placeholder": "[SIGNATURE]",
        },
        "letter_full_text": full_text,
        "warnings": warnings,
    }


# ── ITA Strategy Report Generation ──

def get_draw_context_string() -> str:
    """Loads the draw results JSON from the frontend and returns a context string for the AI prompt."""
    try:
        json_path = os.path.join(os.path.dirname(__file__), '..', 'frontend', 'src', 'data', 'draw_results.json')
        with open(json_path, 'r', encoding='utf-8') as f:
            draws = json.load(f)

        cec_draws = [d for d in draws if d.get('drawType') == 'CEC']
        if not cec_draws:
            return "Latest Express Entry general draw cutoff is approximately 520-540 points."
            
        latest_cec = cec_draws[0]['crsScore']
        recent_cec = cec_draws[:15]
        avg_cec = sum(d['crsScore'] for d in recent_cec) // len(recent_cec)
        
        trend = "stable"
        if len(recent_cec) >= 6:
            older_avg = sum(d['crsScore'] for d in recent_cec[3:6]) / 3
            newer_avg = sum(d['crsScore'] for d in recent_cec[:3]) / 3
            if newer_avg > older_avg + 5:
                trend = "rising"
            elif newer_avg < older_avg - 5:
                trend = "falling"

        french_draws = [d for d in draws if d.get('drawType') == 'French']
        latest_french = french_draws[0]['crsScore'] if french_draws else "Unknown"

        pnp_draws = [d for d in draws if d.get('drawType') == 'PNP']
        latest_pnp = pnp_draws[0]['crsScore'] if pnp_draws else "Unknown"
        
        return (f"CURRENT EXPRESS ENTRY CLIMATE (Provide objective advice based on these numbers, do NOT state statistical probabilities or guarantees):\n"
                f"- Latest CEC (General) Cutoff: {latest_cec}\n"
                f"- 6-Month CEC Average: {avg_cec} (Trend: {trend})\n"
                f"- Latest French Category Cutoff: {latest_french}\n"
                f"- Latest PNP Cutoff: {latest_pnp}\n\n"
                f"You MUST use this data to perform a 'gap analysis'. Tell the user exactly how many points they are away from the 6-month CEC Average, and acknowledge recent trends. "
                f"If their score is very low, mention alternative pathways like French language (NCLC 7) if applicable.")
    except Exception as e:
        print(f"Warning: Could not load draw context: {e}")
        return "The latest Express Entry general draw cutoff is approximately 520-540 points. Category-based draws can have lower cutoffs."

def generate_ita_strategy(raw_inputs: dict, score: dict, breakdown: dict) -> dict:
    """
    Generate a personalized ITA (Invitation to Apply) strategy report
    based on the user's exact CRS profile inputs and calculated score.
    """
    
    # Build a human-readable profile summary for the AI
    profile_lines = []
    profile_lines.append(f"Age: {raw_inputs.get('age', 'Unknown')}")
    profile_lines.append(f"Marital Status: {raw_inputs.get('maritalStatus', 'Unknown')}")
    if raw_inputs.get('maritalStatus') in ['Married', 'Common-Law']:
        profile_lines.append(f"  Spouse is Canadian PR/Citizen: {raw_inputs.get('spouseIsPR', 'Unknown')}")
        profile_lines.append(f"  Spouse accompanying: {raw_inputs.get('spouseAccompanying', 'Unknown')}")
    
    education_labels = {
        'none': 'None or less than secondary',
        'secondary': 'Secondary diploma (high school)',
        'one-year': 'One-year post-secondary',
        'two-year': 'Two-year post-secondary',
        'bachelors': "Bachelor's degree",
        'two-or-more': 'Two or more certificates/degrees',
        'masters': "Master's degree or professional degree",
        'doctoral': 'Doctoral (PhD)',
    }
    profile_lines.append(f"Education: {education_labels.get(raw_inputs.get('education', ''), raw_inputs.get('education', 'Unknown'))}")
    profile_lines.append(f"Canadian Education: {raw_inputs.get('hasCanadianEducation', 'No')}")
    if raw_inputs.get('hasCanadianEducation') == 'Yes':
        can_edu_labels = {'one-two': '1-2 year credential', 'three-plus': '3+ years / Masters / PhD'}
        profile_lines.append(f"  Canadian Credential: {can_edu_labels.get(raw_inputs.get('canadianEducation', ''), 'Unknown')}")
    
    profile_lines.append(f"Primary Language Test: {raw_inputs.get('lang1Test', 'None')}")
    profile_lines.append(f"  Listening: {raw_inputs.get('lang1L', 'N/A')}, Speaking: {raw_inputs.get('lang1S', 'N/A')}, Reading: {raw_inputs.get('lang1R', 'N/A')}, Writing: {raw_inputs.get('lang1W', 'N/A')}")
    
    lang2 = raw_inputs.get('lang2Test', 'None / Not Applicable')
    if lang2 != 'None / Not Applicable':
        profile_lines.append(f"Second Language Test: {lang2}")
        profile_lines.append(f"  Listening: {raw_inputs.get('lang2L', 'N/A')}, Speaking: {raw_inputs.get('lang2S', 'N/A')}, Reading: {raw_inputs.get('lang2R', 'N/A')}, Writing: {raw_inputs.get('lang2W', 'N/A')}")
    else:
        profile_lines.append("Second Language Test: None")
    
    profile_lines.append(f"Canadian Work Experience: {raw_inputs.get('canadianWork', 'None')}")
    profile_lines.append(f"Foreign Work Experience: {raw_inputs.get('foreignWork', 'None')}")
    profile_lines.append(f"Provincial Nomination: {raw_inputs.get('provincialNom', 'No')}")
    profile_lines.append(f"Sibling in Canada: {raw_inputs.get('siblingInCanada', 'No')}")
    profile_lines.append(f"Certificate of Qualification: {raw_inputs.get('certOfQualification', 'No')}")
    
    # Spouse factors
    if raw_inputs.get('spouseAccompanying') == 'Yes' and raw_inputs.get('spouseIsPR') == 'No':
        profile_lines.append(f"Spouse Education: {education_labels.get(raw_inputs.get('spouseEducation', ''), 'Unknown')}")
        profile_lines.append(f"Spouse Language Test: {raw_inputs.get('spLangTest', 'None')}")
        if raw_inputs.get('spLangTest', 'None / Not Applicable') != 'None / Not Applicable':
            profile_lines.append(f"  Spouse L: {raw_inputs.get('spL', 'N/A')}, S: {raw_inputs.get('spS', 'N/A')}, R: {raw_inputs.get('spR', 'N/A')}, W: {raw_inputs.get('spW', 'N/A')}")
        profile_lines.append(f"Spouse Canadian Work: {raw_inputs.get('spouseCanadianWork', 'None')}")
    
    profile_summary = "\n".join(profile_lines)
    
    score_summary = f"""
CRS Score Breakdown:
  Total Score: {score.get('total', 0)} / 1200
  Core/Human Capital: {score.get('core', 0)}
  Spouse Factors: {score.get('spouse', 0)}
  Skill Transferability: {score.get('transferability', 0)}
  Additional Points: {score.get('additional', 0)}

Detailed Core Breakdown:
  Age Points: {breakdown.get('core', {}).get('age', 0)}
  Education Points: {breakdown.get('core', {}).get('education', 0)}
  Official Languages: {breakdown.get('core', {}).get('officialLanguages', 0)}
    First Official Language: {breakdown.get('core', {}).get('firstOfficialLanguage', 0)}
    Second Official Language: {breakdown.get('core', {}).get('secondOfficialLanguage', 0)}
  Canadian Work Experience: {breakdown.get('core', {}).get('canadianWorkExperience', 0)}

Skill Transferability Breakdown:
  Education + Language: {breakdown.get('transferability', {}).get('education', {}).get('languageAndEducation', 0)}
  Education + Canadian Work: {breakdown.get('transferability', {}).get('education', {}).get('canadianWorkAndEducation', 0)}
  Foreign Work + Language: {breakdown.get('transferability', {}).get('foreignWork', {}).get('languageAndForeignWork', 0)}
  Foreign Work + Canadian Work: {breakdown.get('transferability', {}).get('foreignWork', {}).get('canadianAndForeignWork', 0)}
  Certificate of Qualification: {breakdown.get('transferability', {}).get('certificateOfQualification', 0)}

Additional Points Breakdown:
  Provincial Nomination: {breakdown.get('additional', {}).get('provincialNomination', 0)}
  Study in Canada: {breakdown.get('additional', {}).get('studyInCanada', 0)}
  Sibling in Canada: {breakdown.get('additional', {}).get('siblingInCanada', 0)}
  French Language Skills: {breakdown.get('additional', {}).get('frenchLanguageSkills', 0)}
"""

    prompt = f"""You are an expert Canadian immigration consultant specializing in Express Entry and the Comprehensive Ranking System (CRS). You have deep knowledge of all pathways to improve CRS scores, including Provincial Nominee Programs (PNPs), language testing strategies, education credential assessment, and Canadian work experience optimization.

A user has completed their CRS calculation and their EXACT profile is below. Your job is to create a highly personalized, actionable strategy report to help them receive an Invitation to Apply (ITA).

═══════════════════════════════════════
USER PROFILE:
{profile_summary}

{score_summary}
═══════════════════════════════════════

{get_draw_context_string()}

═══════════════════════════════════════
OFFICIAL CRS SCORING REFERENCE
(You MUST use these tables for ALL point calculations. Do NOT guess point values.)
═══════════════════════════════════════

A. CORE / HUMAN CAPITAL FACTORS:

Age (with spouse / without spouse):
  17 or under: 0/0, 18: 90/99, 19: 95/105, 20-29: 100/110,
  30: 95/105, 31: 90/99, 32: 85/94, 33: 80/88, 34: 75/83,
  35: 70/77, 36: 65/72, 37: 60/66, 38: 55/61, 39: 50/55,
  40: 45/50, 41: 35/39, 42: 25/28, 43: 15/17, 44: 5/6, 45+: 0/0

Education (with spouse / without spouse):
  None: 0/0, Secondary: 28/30, One-year post-sec: 84/90,
  Two-year post-sec: 91/98, Bachelor's: 112/120,
  Two or more credentials: 119/128, Master's: 126/135, Doctoral: 140/150

First Official Language PER ABILITY (with spouse / without spouse):
  CLB < 4: 0/0, CLB 4-5: 6/6, CLB 6: 8/9, CLB 7: 16/17,
  CLB 8: 22/23, CLB 9: 29/31, CLB 10+: 32/34
  (Multiply by 4 abilities for total. Max = 128/136)

Second Official Language PER ABILITY (with spouse / without spouse):
  CLB < 5: 0/0, CLB 5-6: 1/1, CLB 7-8: 3/3, CLB 9+: 5/6
  (Multiply by 4 abilities for total. Max = 20/24)

Canadian Work Experience (with spouse / without spouse):
  None: 0/0, 1yr: 35/40, 2yr: 46/53, 3yr: 56/64, 4yr: 63/72, 5+yr: 70/80

B. SPOUSE FACTORS (only when spouse is accompanying and NOT a PR/citizen):

Spouse Education:
  None: 0, Secondary: 2, One-year: 6, Two-year: 7,
  Bachelor's: 8, Two or more: 9, Master's: 10, Doctoral: 10

Spouse Language PER ABILITY:
  CLB < 5: 0, CLB 5-6: 1, CLB 7-8: 3, CLB 9+: 5
  (Multiply by 4 abilities for total. Max = 20)

Spouse Canadian Work:
  None: 0, 1yr: 5, 2yr: 7, 3yr: 8, 4yr: 9, 5+yr: 10

C. ADDITIONAL POINTS:
  Provincial Nomination: +600
  Sibling in Canada (PR/citizen): +15
  Canadian education (1-2yr credential): +15, (3+yr or grad): +30
  French language proficiency: 
    - IF user has NCLC 7+ in all 4 French abilities AND English is CLB 4 or lower (or no English test): +25 points
    - IF user has NCLC 7+ in all 4 French abilities AND English is CLB 5 or higher in all 4 abilities: +50 points
    (Note: You MUST check the user's English CLB levels before applying this bonus. If they have strong English, use 50.)

D. SKILL TRANSFERABILITY FACTORS (Max 100 points total combination limit)
  Education + Language:
    1-yr/2-yr credential + CLB 7/8 = 13 pts. With CLB 9+ = 25 pts.
    Two or more / Masters / PhD + CLB 7/8 = 25 pts. With CLB 9+ = 50 pts.
  Education + Canadian Work Experience:
    1-yr/2-yr credential + 1 yr Cdn Work = 13 pts. With 2+ yr Cdn Work = 25 pts.
    Two or more / Masters / PhD + 1 yr Cdn Work = 25 pts. With 2+ yr Cdn Work = 50 pts.
  Foreign Work + Language:
    1-2 yrs Exp + CLB 7/8 = 13 pts. With CLB 9+ = 25 pts.
    3+ yrs Exp + CLB 7/8 = 25 pts. With CLB 9+ = 50 pts.
  Foreign Work + Canadian Work Experience:
    1-2 yrs Foreign Exp + 1 yr Cdn = 13 pts. With 2+ yr Cdn = 25 pts.
    3+ yrs Foreign Exp + 1 yr Cdn = 25 pts. With 2+ yr Cdn = 50 pts.
  Certificate of Qualification (Trades) + Language:
    With Cert + CLB 5/6 = 25 pts. With CLB 7+ = 50 pts.

═══════════════════════════════════════
STRATEGIC DIRECTIVES & WARNINGS:
1. AGE DECAY WARNING: Check the user's age. If they are exactly 29, emphasize they will start losing 5 points every birthday. If they are >29, warn them that processing delays cost 5-6 points per year so speed is critical.
2. UNACCOMPANYING SPOUSE STRATEGY: If the user's accompanying spouse has low education (no degree) AND low language scores (CLB < 7), you MUST run a simulation and explicitly recommend considering declaring the spouse as 'Non-Accompanying' to leverage the higher single-applicant scoring grid.
3. 100-POINT TRANSFERABILITY CAP: The absolute maximum for ALL Skill Transferability points (Section D) combined is 100 points. Do NOT project point gains that push someone over this cap.

CRITICAL: When recommending actions, you MUST calculate the EXACT point difference
between the user's current level and the target level using these tables.
For example: Spouse going from no language test (0 pts) to CLB 9 all abilities = 5×4 = +20 points, NOT a rough estimate.
═══════════════════════════════════════

Generate a comprehensive, personalized ITA strategy report. You MUST return ONLY a valid JSON object with this exact structure (no markdown, no code fences, just raw JSON):

{{
  "current_score": {score.get('total', 0)},
  "estimated_cutoff": 530,
  "gap": {max(530 - score.get('total', 0), 0)},
  "overall_assessment": "YOUR_ASSESSMENT_HERE",
  "category_based_eligibility": [
    {{
      "category": "CATEGORY_NAME",
      "eligible": true,
      "note": "EXPLANATION"
    }}
  ],
  "actions": [
    {{
      "rank": 1,
      "title": "ACTION_TITLE",
      "description": "DETAILED_DESCRIPTION",
      "potential_points": "+XX to +YY",
      "effort_level": "Low or Medium or High",
      "estimated_timeline": "X-Y months",
      "estimated_cost": "$XXX CAD",
      "priority": "Critical or High or Medium or Low",
      "specific_targets": "SPECIFIC_TARGETS"
    }}
  ],
  "language_optimization": {{
    "current_first_language_points": {breakdown.get('core', {}).get('firstOfficialLanguage', 0)},
    "max_first_language_points": 136,
    "improvement_possible": 0,
    "specific_targets": "TARGET_SCORES",
    "second_language_recommendation": "RECOMMENDATION"
  }},
  "pnp_recommendations": [
    {{
      "province": "PROVINCE",
      "stream": "STREAM_NAME",
      "why_suitable": "REASON",
      "points_impact": "+600",
      "requirements_summary": "REQUIREMENTS"
    }}
  ],
  "timeline_summary": "TIMELINE_PARAGRAPH",
  "disclaimer": "This report provides general guidance based on publicly available CRS criteria and immigration program information. It is not legal advice. For personalized legal advice, consult a Regulated Canadian Immigration Consultant (RCIC) or immigration lawyer. Immigration policies and cutoff scores change frequently — always verify with IRCC's official website."
}}

═══════════════════════════════════════
OFFICIAL LANGUAGE TEST → CLB CONVERSION TABLES
(You MUST use these tables. Do NOT guess conversions.)
═══════════════════════════════════════

CELPIP-General → CLB:
  CELPIP scores map DIRECTLY 1:1 to CLB levels.
  CELPIP 4 = CLB 4, CELPIP 5 = CLB 5, CELPIP 6 = CLB 6,
  CELPIP 7 = CLB 7, CELPIP 8 = CLB 8, CELPIP 9 = CLB 9,
  CELPIP 10 = CLB 10, CELPIP 11 = CLB 11, CELPIP 12 = CLB 12.
  CELPIP uses WHOLE NUMBERS only (no decimals like 8.0 or 7.5).

IELTS General Training → CLB:
  Listening: 4.5=CLB4, 5.0=CLB5, 5.5=CLB6, 6.0=CLB7, 7.5=CLB8, 8.0=CLB9, 8.5=CLB10
  Reading:   3.5=CLB4, 4.0=CLB5, 5.0=CLB6, 6.0=CLB7, 6.5=CLB8, 7.0=CLB9, 8.0=CLB10
  Writing:   4.0=CLB4, 5.0=CLB5, 5.5=CLB6, 6.0=CLB7, 6.5=CLB8, 7.0=CLB9, 7.5=CLB10
  Speaking:  4.0=CLB4, 5.0=CLB5, 5.5=CLB6, 6.0=CLB7, 6.5=CLB8, 7.0=CLB9, 7.5=CLB10

TEF Canada → NCLC:
  Listening: 145-216=NCLC4, 217-248=NCLC5, 249-279=NCLC6, 280-297=NCLC7, 298-315=NCLC8, 316-333=NCLC9, 334-360=NCLC10
  Reading:   121-150=NCLC4, 151-180=NCLC5, 181-206=NCLC6, 207-232=NCLC7, 233-247=NCLC8, 248-262=NCLC9, 263-300=NCLC10
  Writing:   181-225=NCLC4, 226-270=NCLC5, 271-309=NCLC6, 310-348=NCLC7, 349-370=NCLC8, 371-392=NCLC9, 393-450=NCLC10
  Speaking:  181-225=NCLC4, 226-270=NCLC5, 271-309=NCLC6, 310-348=NCLC7, 349-370=NCLC8, 371-392=NCLC9, 393-450=NCLC10

TCF Canada → NCLC:
  Listening: 331-368=NCLC4, 369-397=NCLC5, 398-457=NCLC6, 458-502=NCLC7, 503-522=NCLC8, 523-548=NCLC9, 549-699=NCLC10
  Reading:   342-374=NCLC4, 375-405=NCLC5, 406-452=NCLC6, 453-498=NCLC7, 499-523=NCLC8, 524-548=NCLC9, 549-699=NCLC10
  Writing:   4=NCLC4, 6=NCLC5, 6=NCLC6, 10=NCLC7, 12=NCLC8, 14=NCLC9, 16=NCLC10
  Speaking:  4=NCLC4, 6=NCLC5, 6=NCLC6, 10=NCLC7, 12=NCLC8, 14=NCLC9, 16=NCLC10

═══════════════════════════════════════

RULES:
1. Be SPECIFIC to this user's exact profile. Reference their actual scores, not generic advice.
2. Rank actions from highest impact/easiest to lowest impact/hardest.
3. If the user already has maximum points in a category, acknowledge it and skip it.
4. Be honest — if the gap is very large, say so and suggest realistic pathways.
5. Include at least 4-6 actionable recommendations.
6. For language optimization, use ONLY the conversion tables above. NEVER mix up CELPIP and IELTS scoring. If the user took CELPIP, all targets must be in CELPIP whole numbers. If the user took IELTS, all targets must be in IELTS band scores. State both the test-specific score AND the CLB level.
7. For PNP, suggest 2-3 specific provincial programs they may qualify for based on their profile.
8. All costs should be in CAD.
9. NEVER use decimal scores (e.g., "8.0") for CELPIP — CELPIP uses whole numbers only.
10. NEVER confuse test scoring systems. Double-check every conversion against the tables above.
"""

    try:
        response = gemini_client.models.generate_content(
            model="gemini-2.5-flash",
            contents=[prompt],
            config=types.GenerateContentConfig(
                temperature=0.1,
                max_output_tokens=16384,
                response_mime_type="application/json",
            )
        )
        
        raw_text = response.text.strip()
        
        # Strip markdown code fences if present
        if raw_text.startswith("```"):
            lines = raw_text.split("\n")
            lines = lines[1:]  # Remove first ```json line
            if lines and lines[-1].strip() == "```":
                lines = lines[:-1]
            raw_text = "\n".join(lines)
        
        strategy = json.loads(raw_text)
        return strategy
        
    except json.JSONDecodeError as e:
        print(f"Failed to parse ITA strategy JSON: {e}")
        print(f"Raw response: {raw_text[:1000]}")
        raise ValueError(f"AI returned invalid JSON for ITA strategy: {str(e)}")
    except Exception as e:
        print(f"ITA Strategy generation failed: {e}")
        raise

